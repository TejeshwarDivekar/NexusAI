import os
import re
import hashlib
from typing import Dict, Any, List, Optional
from datetime import datetime

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from app.core.logging import logger
from app.services.document_generation.document_model import (
    StructuredResearchDocument, DocumentModelBuilder
)
from app.services.document_generation.citation_validator import CitationValidator


class IEEEDocumentGenerator:
    """
    Production IEEE-Style Microsoft Word Document (.docx) Generator.
    Adheres strictly to IEEE publication structure and verified research data:
    - Formal Research Title Block
    - 1-Minute Executive Research Summary Box
    - Abstract & Keywords (Index Terms)
    - Numbered Sections (I. Introduction, II. Methodology, III. Findings, etc.)
    - Evidence Analysis Table
    - Limitations & Scope
    - Conclusion
    - IEEE Formatted References (strictly from real retrieved sources)
    """

    @classmethod
    def set_cell_background(cls, cell, fill_hex: str):
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    @classmethod
    def set_cell_margins(cls, cell, top=80, bottom=80, left=120, right=120):
        tcPr = cell._element.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    @classmethod
    def generate_docx(
        cls,
        task_id: str,
        query: str,
        report_markdown: str,
        sources: List[Dict[str, Any]],
        evidence_matrix: List[Dict[str, Any]],
        claims: List[Dict[str, Any]],
        contradictions: List[Dict[str, Any]],
        summary: Optional[str] = None,
        retrieval_timestamp: Optional[str] = None,
        author_name: str = "Principal Researcher",
        output_dir: str = "generated_docs",
        version: int = 1,
        classification: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Generates an accurate, publication-grade IEEE Word Document."""
        os.makedirs(output_dir, exist_ok=True)

        safe_query_name = re.sub(r'[^a-zA-Z0-9]', '_', query[:40]).strip('_')
        filename = f"IEEE_Report_{safe_query_name}_v{version}_{task_id[:8]}.docx"
        file_path = os.path.join(output_dir, filename)

        # 1. Build Structured Document Model
        doc_model = DocumentModelBuilder.build_structured_document(
            task_id=task_id,
            query=query,
            report_markdown=report_markdown,
            sources=sources,
            evidence_matrix=evidence_matrix,
            claims=claims,
            contradictions=contradictions,
            summary=summary,
            retrieval_timestamp=retrieval_timestamp,
            author_name=author_name,
            classification=classification
        )

        # 2. Validate and Align Citations
        val_report = CitationValidator.validate_and_align_citations(doc_model)

        # 3. Create Word Document
        doc = docx.Document()

        # Set 0.75-inch standard IEEE margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

            # Header & Footer
            header = section.header
            header_p = header.paragraphs[0]
            header_p.text = "IEEE RESEARCH SYNTHESIS — VERIFIED EVIDENCE REPORT"
            header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            header_p.style.font.name = "Times New Roman"
            header_p.style.font.size = Pt(8.5)
            header_p.style.font.color.rgb = RGBColor(120, 120, 120)

            footer = section.footer
            footer_p = footer.paragraphs[0]
            footer_p.text = f"Report Version {version} | Data Retrieved: {doc_model.retrieval_timestamp} | Task ID: {task_id[:8]}"
            footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_p.style.font.name = "Times New Roman"
            footer_p.style.font.size = Pt(8.5)
            footer_p.style.font.color.rgb = RGBColor(120, 120, 120)

        # Base Normal Style
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(10)
        normal_style.font.color.rgb = RGBColor(15, 23, 42)

        # ----------------------------------------------------
        # 1. TITLE BLOCK
        # ----------------------------------------------------
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(12)
        title_p.paragraph_format.space_after = Pt(4)
        title_run = title_p.add_run(doc_model.formal_title)
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(15, 23, 42)

        if doc_model.subtitle:
            sub_p = doc.add_paragraph()
            sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_p.paragraph_format.space_after = Pt(8)
            sub_run = sub_p.add_run(doc_model.subtitle)
            sub_run.font.name = "Times New Roman"
            sub_run.font.size = Pt(10.5)
            sub_run.font.italic = True
            sub_run.font.color.rgb = RGBColor(71, 85, 105)

        # ----------------------------------------------------
        # 2. AUTHOR & INSTITUTION BLOCK
        # ----------------------------------------------------
        author_p = doc.add_paragraph()
        author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_p.paragraph_format.space_after = Pt(14)
        
        ar1 = author_p.add_run(f"Author: {doc_model.author_name}   |   Organization: {doc_model.organization}\n")
        ar1.font.bold = True
        ar1.font.size = Pt(9.5)
        
        ar2 = author_p.add_run(f"Date: {doc_model.generation_date}   |   Data Retrieved: {doc_model.retrieval_timestamp}")
        ar2.font.size = Pt(9.0)
        ar2.font.color.rgb = RGBColor(100, 116, 139)

        # ----------------------------------------------------
        # 3. 1-MINUTE EXECUTIVE RESEARCH SUMMARY BOX
        # ----------------------------------------------------
        summary_box = doc_model.research_summary
        t_box = doc.add_table(rows=1, cols=1)
        t_box.alignment = WD_TABLE_ALIGNMENT.CENTER
        t_box.autofit = False
        t_box.columns[0].width = Inches(7.0)
        cell = t_box.cell(0, 0)
        cls.set_cell_background(cell, "F8FAFC")
        cls.set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        
        bp = cell.paragraphs[0]
        bp.paragraph_format.space_after = Pt(4)
        r_box_title = bp.add_run("EXECUTIVE RESEARCH SUMMARY (1-MINUTE OVERVIEW)\n")
        r_box_title.bold = True
        r_box_title.font.size = Pt(9.5)
        r_box_title.font.color.rgb = RGBColor(0, 51, 102)

        r_box_target = bp.add_run(f"Target Topic: {summary_box.researched_topic}\n\nKey Findings:\n")
        r_box_target.font.size = Pt(9.0)
        for f in summary_box.core_findings[:3]:
            rf = bp.add_run(f"  • {f}\n")
            rf.font.size = Pt(9.0)

        bp.add_run("\nPrimary Evidence:\n").font.size = Pt(9.0)
        for e in summary_box.key_evidence_points[:2]:
            re_run = bp.add_run(f"  • {e}\n")
            re_run.font.size = Pt(9.0)

        rb_bot = bp.add_run(f"\nBottom Line: {summary_box.bottom_line_conclusion}")
        rb_bot.font.size = Pt(9.0)
        rb_bot.bold = True

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # ----------------------------------------------------
        # 4. ABSTRACT & KEYWORDS
        # ----------------------------------------------------
        abs_p = doc.add_paragraph()
        abs_p.paragraph_format.space_after = Pt(8)
        abs_p.paragraph_format.line_spacing = 1.15
        
        abs_head = abs_p.add_run("Abstract— ")
        abs_head.bold = True
        abs_head.italic = True
        abs_head.font.size = Pt(10)
        
        abs_text = abs_p.add_run(doc_model.abstract)
        abs_text.italic = True
        abs_text.font.size = Pt(9.5)

        if doc_model.keywords:
            kw_p = doc.add_paragraph()
            kw_p.paragraph_format.space_after = Pt(12)
            kw_head = kw_p.add_run("Index Terms— ")
            kw_head.bold = True
            kw_head.italic = True
            kw_head.font.size = Pt(9.5)
            
            kw_text = kw_p.add_run(", ".join(doc_model.keywords))
            kw_text.italic = True
            kw_text.font.size = Pt(9.5)

        # ----------------------------------------------------
        # 5. MAIN SECTIONS & TABLES
        # ----------------------------------------------------
        for sec in doc_model.sections:
            h_p = doc.add_paragraph()
            h_p.paragraph_format.space_before = Pt(14)
            h_p.paragraph_format.space_after = Pt(4)
            h_run = h_p.add_run(f"{sec.roman_number}.  {sec.title.upper()}")
            h_run.bold = True
            h_run.font.size = Pt(10.5)
            h_run.font.color.rgb = RGBColor(15, 23, 42)

            for p_text in sec.paragraphs:
                p_obj = doc.add_paragraph()
                p_obj.paragraph_format.space_after = Pt(5)
                p_obj.paragraph_format.line_spacing = 1.15
                
                # Split inline citations [X] to bold them
                parts = re.split(r'(\[\d+\])', p_text)
                for part in parts:
                    if re.match(r'^\[\d+\]$', part):
                        r = p_obj.add_run(part)
                        r.bold = True
                        r.font.size = Pt(9.5)
                    else:
                        r = p_obj.add_run(part)
                        r.font.size = Pt(10)

            # If section has a formatted Table
            if sec.table:
                t_obj = sec.table
                tp = doc.add_paragraph()
                tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                tp.paragraph_format.space_before = Pt(6)
                tp.paragraph_format.space_after = Pt(3)
                tr = tp.add_run(f"TABLE {t_obj.table_number}\n{t_obj.title.upper()}")
                tr.font.size = Pt(8.5)
                tr.bold = True
                
                t_word = doc.add_table(rows=len(t_obj.rows) + 1, cols=len(t_obj.headers))
                t_word.alignment = WD_TABLE_ALIGNMENT.CENTER
                t_word.autofit = True
                
                # Header row
                for i, h in enumerate(t_obj.headers):
                    c = t_word.rows[0].cells[i]
                    cls.set_cell_background(c, "0F172A")
                    cls.set_cell_margins(c, 60, 60, 80, 80)
                    p = c.paragraphs[0]
                    r = p.add_run(h)
                    r.bold = True
                    r.font.size = Pt(8.5)
                    r.font.color.rgb = RGBColor(255, 255, 255)

                # Data rows
                for r_idx, row_vals in enumerate(t_obj.rows):
                    row_el = t_word.rows[r_idx + 1]
                    bg = "FFFFFF" if r_idx % 2 == 0 else "F8FAFC"
                    for c_idx, val in enumerate(row_vals):
                        c = row_el.cells[c_idx]
                        cls.set_cell_background(c, bg)
                        cls.set_cell_margins(c, 40, 40, 60, 60)
                        p = c.paragraphs[0]
                        r = p.add_run(val)
                        r.font.size = Pt(8.5)
                        if c_idx == 0:
                            r.bold = True
                
                doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # ----------------------------------------------------
        # 6. LIMITATIONS
        # ----------------------------------------------------
        lim_h = doc.add_paragraph()
        lim_h.paragraph_format.space_before = Pt(14)
        lim_h.paragraph_format.space_after = Pt(4)
        lim_hr = lim_h.add_run("VI.  RESEARCH LIMITATIONS & UNCERTAINTIES")
        lim_hr.bold = True
        lim_hr.font.size = Pt(10.5)

        for lim in doc_model.limitations:
            lp = doc.add_paragraph()
            lp.paragraph_format.space_after = Pt(3)
            lr = lp.add_run(f"•  {lim}")
            lr.font.size = Pt(9.5)

        # ----------------------------------------------------
        # 7. CONCLUSION
        # ----------------------------------------------------
        con_h = doc.add_paragraph()
        con_h.paragraph_format.space_before = Pt(14)
        con_h.paragraph_format.space_after = Pt(4)
        con_hr = con_h.add_run("VII.  CONCLUSION")
        con_hr.bold = True
        con_hr.font.size = Pt(10.5)

        cp = doc.add_paragraph()
        cp.paragraph_format.space_after = Pt(12)
        cp.paragraph_format.line_spacing = 1.15
        cr = cp.add_run(doc_model.conclusion)
        cr.font.size = Pt(10)

        # ----------------------------------------------------
        # 8. REFERENCES (BIBLIOGRAPHY)
        # ----------------------------------------------------
        ref_h = doc.add_paragraph()
        ref_h.paragraph_format.space_before = Pt(16)
        ref_h.paragraph_format.space_after = Pt(6)
        ref_hr = ref_h.add_run("REFERENCES")
        ref_hr.bold = True
        ref_hr.font.size = Pt(10.5)

        if doc_model.references:
            for ref in doc_model.references:
                rp = doc.add_paragraph()
                rp.paragraph_format.space_after = Pt(3)
                rp.paragraph_format.left_indent = Inches(0.25)
                rp.paragraph_format.first_line_indent = Inches(-0.25)
                
                ref_text = ref.formatted_citation_text()
                rr = rp.add_run(ref_text)
                rr.font.size = Pt(9.0)
                rr.font.color.rgb = RGBColor(30, 41, 59)
        else:
            rp = doc.add_paragraph()
            rp.add_run("No references recorded in retrieved sources.").font.size = Pt(9.0)

        # Save DOCX
        doc.save(file_path)

        file_size = os.path.getsize(file_path)
        with open(file_path, "rb") as f:
            sha256 = hashlib.sha256(f.read()).hexdigest()

        logger.info(f"IEEEDocumentGenerator: Compiled '{filename}' ({file_size} bytes, SHA256: {sha256[:12]})")

        return {
            "status": "success",
            "task_id": task_id,
            "version": version,
            "file_name": filename,
            "file_path": file_path,
            "file_size": file_size,
            "sha256_hash": sha256,
            "doc_format": "docx",
            "reference_count": len(doc_model.references),
            "references_count": len(doc_model.references),
            "generation_status": "completed",
            "metadata_json": {
                "formal_title": doc_model.formal_title,
                "references_count": len(doc_model.references),
                "sections_count": len(doc_model.sections),
                "validation_report": val_report
            }
        }
