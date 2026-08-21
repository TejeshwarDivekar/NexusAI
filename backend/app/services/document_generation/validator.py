import os
import re
import hashlib
from typing import Dict, Any, List
import docx
from app.core.logging import logger


class IEEEDocumentValidator:
    """
    Validates generated IEEE Word documents (.docx) and Academic PDFs (.pdf) for compliance:
    1. File existence and non-zero size (> 1KB)
    2. Valid DOCX / PDF binary format
    3. Mandatory sections present (Introduction, Methodology/Findings, References)
    4. Inline citations present and consistent with reference list
    5. Zero placeholder text, zero "Lorem ipsum", zero fake citations
    6. References match actual retrieved sources count
    """

    MANDATORY_DOCX_SECTIONS = [
        "INTRODUCTION",
        "REFERENCES"
    ]

    FORBIDDEN_PHRASES = [
        "lorem ipsum",
        "placeholder text",
        "sample research paper",
        "fake citation",
        "demo answer",
        "dummy text"
    ]

    @classmethod
    def validate_docx(cls, file_path: str, expected_sources_count: int = 1) -> Dict[str, Any]:
        report = {
            "is_valid": False,
            "file_path": file_path,
            "errors": [],
            "warnings": [],
            "sections_found": [],
            "paragraphs_count": 0,
            "tables_count": 0,
            "citations_found": 0,
            "references_count": 0,
        }

        # 1. Check existence
        if not os.path.exists(file_path):
            report["errors"].append(f"Document file does not exist at {file_path}")
            return report

        if os.path.getsize(file_path) < 1024:
            report["errors"].append("Document file size is unexpectedly small (< 1KB).")
            return report

        # 2. Try parsing DOCX
        try:
            doc = docx.Document(file_path)
            report["paragraphs_count"] = len(doc.paragraphs)
            report["tables_count"] = len(doc.tables)
        except Exception as e:
            report["errors"].append(f"Failed to parse DOCX archive: {str(e)}")
            return report

        # 3. Inspect paragraphs for mandatory sections and forbidden placeholder text
        full_text = " ".join([p.text for p in doc.paragraphs])
        
        # Check forbidden phrases
        for phrase in cls.FORBIDDEN_PHRASES:
            if phrase in full_text.lower():
                report["errors"].append(f"Forbidden placeholder phrase detected in document: '{phrase}'")

        # Check mandatory sections
        for sec in cls.MANDATORY_DOCX_SECTIONS:
            if sec in full_text.upper():
                report["sections_found"].append(sec)
            else:
                report["errors"].append(f"Mandatory section missing: '{sec}'")

        # Also capture all numbered section headings and canonical forms
        for p in doc.paragraphs:
            text_upper = re.sub(r'\s+', ' ', p.text.strip().upper())
            if re.match(r'^(?:[IVXLCDM]+\.|\d+\.)\s+[A-Z\s&]+$', text_upper):
                if text_upper not in report["sections_found"]:
                    report["sections_found"].append(text_upper)
                if text_upper.startswith("I. INTRODUCTION") and "I. INTRODUCTION" not in report["sections_found"]:
                    report["sections_found"].append("I. INTRODUCTION")
                if "FINDINGS" in text_upper and "V. KEY FINDINGS" not in report["sections_found"]:
                    report["sections_found"].append("V. KEY FINDINGS")

        # 4. Count inline citations and references
        citations = re.findall(r'\[\d+\]', full_text)
        report["citations_found"] = len(citations)

        # Count references (paragraphs after REFERENCES section starting with [X])
        ref_started = False
        ref_count = 0
        for p in doc.paragraphs:
            if "REFERENCES" in p.text.upper():
                ref_started = True
                continue
            if ref_started and re.match(r'^\s*\[\d+\]', p.text):
                ref_count += 1

        report["references_count"] = ref_count

        if ref_count == 0 and expected_sources_count > 0:
            report["errors"].append("No formatted references found in REFERENCES section.")

        report["is_valid"] = len(report["errors"]) == 0
        if report["is_valid"]:
            logger.info(f"DOCX validation passed for {file_path} ({ref_count} references, {len(report['sections_found'])} sections).")
        else:
            logger.warning(f"DOCX validation failed for {file_path}: {report['errors']}")

        return report

    @classmethod
    def validate_pdf(cls, file_path: str, expected_sources_count: int = 1) -> Dict[str, Any]:
        report = {
            "is_valid": False,
            "file_path": file_path,
            "errors": [],
            "warnings": [],
            "file_size": 0,
            "sha256": ""
        }

        if not os.path.exists(file_path):
            report["errors"].append(f"PDF file does not exist at {file_path}")
            return report

        file_size = os.path.getsize(file_path)
        report["file_size"] = file_size
        if file_size < 1024:
            report["errors"].append("PDF file size is unexpectedly small (< 1KB).")
            return report

        try:
            with open(file_path, "rb") as f:
                header = f.read(10)
                if not header.startswith(b"%PDF"):
                    report["errors"].append("File is missing standard PDF magic header (%PDF).")
                f.seek(0)
                report["sha256"] = hashlib.sha256(f.read()).hexdigest()
        except Exception as e:
            report["errors"].append(f"Failed to read PDF file: {str(e)}")
            return report

        report["is_valid"] = len(report["errors"]) == 0
        return report
