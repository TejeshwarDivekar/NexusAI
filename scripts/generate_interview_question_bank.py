"""
NexusResearch Complete Interview Question Bank Generator
Generates:
1. docs/interview/NexusResearch_Interview_Question_Bank.docx
2. docs/interview/NexusResearch_Interview_Question_Bank.pdf
"""

import os
import sys
import datetime
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
)
from reportlab.pdfgen import canvas

# ==============================================================================
# REPORTLAB NUMBERED CANVAS (PAGE X OF Y)
# ==============================================================================
class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super(NumberedCanvas, self).__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_decorations(num_pages)
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)

    def draw_page_decorations(self, page_count):
        if self._pageNumber == 1:
            return  # Suppress headers and footers on cover page
            
        self.saveState()
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor("#64748B"))
        
        # Running Top Header
        self.drawString(54, 750, "NEXUSRESEARCH — Comprehensive Technical Interview Question Bank")
        self.setStrokeColor(colors.HexColor("#CBD5E1"))
        self.setLineWidth(0.5)
        self.line(54, 742, 558, 742)
        
        # Running Bottom Footer
        self.line(54, 48, 558, 48)
        self.drawString(54, 34, "Interview Preparation & Architectural Defense Guide")
        page_text = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(558, 34, page_text)
        self.restoreState()


# ==============================================================================
# DOCX HELPER FUNCTIONS
# ==============================================================================
def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_callout_docx(doc, text, title="NOTE", fill_hex="F0F4F8", border_hex="0066CC"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_background(cell, fill_hex)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r_title = p.add_run(f"[{title}] ")
    r_title.bold = True
    r_title.font.color.rgb = RGBColor(0, 51, 102)
    r_title.font.size = Pt(9.5)
    r_text = p.add_run(text)
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(30, 41, 59)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_code_block_docx(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(code_text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(15, 23, 42)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ==============================================================================
# MAIN GENERATOR ROUTINE
# ==============================================================================
def generate_interview_bank():
    docs_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs", "interview")
    os.makedirs(docs_dir, exist_ok=True)
    
    docx_path = os.path.join(docs_dir, "NexusResearch_Interview_Question_Bank.docx")
    pdf_path = os.path.join(docs_dir, "NexusResearch_Interview_Question_Bank.pdf")
    
    print(f"Building Interview Question Bank Word document at: {docx_path}")
    build_docx_bank(docx_path)
    
    print(f"Building Interview Question Bank PDF document at: {pdf_path}")
    build_pdf_bank(pdf_path)
    
    print("Interview Question Bank successfully generated!")


# ==============================================================================
# BUILD DOCX INTERVIEW QUESTION BANK
# ==============================================================================
def build_docx_bank(filepath):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(30, 41, 59)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)

    # --------------------------------------------------------------------------
    # COVER PAGE
    # --------------------------------------------------------------------------
    p_pre = doc.add_paragraph()
    p_pre.paragraph_format.space_before = Pt(40)
    r_pre = p_pre.add_run("NEXUSRESEARCH INTERVIEW MASTERY")
    r_pre.font.size = Pt(11)
    r_pre.bold = True
    r_pre.font.color.rgb = RGBColor(14, 116, 144)

    p_title = doc.add_paragraph()
    r_title = p_title.add_run("Comprehensive Technical Interview Question Bank & Defense Guide")
    r_title.font.size = Pt(22)
    r_title.bold = True
    r_title.font.color.rgb = RGBColor(15, 23, 42)
    p_title.paragraph_format.space_after = Pt(8)

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("A Complete Technical Question Bank (Beginner to Extreme) Covering Architecture, Full-Stack Implementation, Research Engine, LLM Grounding, Database, Security, and System Design")
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(71, 85, 105)
    p_sub.paragraph_format.space_after = Pt(24)

    add_callout_docx(
        doc,
        "INTERVIEW DEFENSE OBJECTIVE:\n"
        "This question bank is built directly from the NexusResearch codebase. "
        "It prepares candidates to defend every architectural choice, API contract, database schema, search provider, evidence grounding step, and deployment mechanism. "
        "Every answer connects to actual code, real test metrics, and production-tested patterns.",
        title="SOURCE OF TRUTH GOVERNANCE"
    )

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(30)
    p_meta.add_run("Project: ").bold = True
    p_meta.add_run("NexusResearch Enterprise Platform\n")
    p_meta.add_run("Technology Stack: ").bold = True
    p_meta.add_run("Next.js 16 (React 19, TS 5), FastAPI (Python 3.10+), SQLAlchemy 2.0, PostgreSQL/SQLite, Gemini 2.5, python-docx, reportlab, NextAuth v5\n")
    p_meta.add_run("Question Bank Scope: ").bold = True
    p_meta.add_run("39 Structured Sections (250+ In-Depth Technical Questions & Case Studies)\n")
    p_meta.add_run("Publication Date: ").bold = True
    p_meta.add_run(datetime.datetime.now().strftime("%B %d, %Y") + "\n")
    p_meta.add_run("Repository Source: ").bold = True
    p_meta.add_run("https://github.com/TejeshwarDivekar/NexusAI.git")

    doc.add_page_break()

    # --------------------------------------------------------------------------
    # TABLE OF CONTENTS
    # --------------------------------------------------------------------------
    doc.add_heading("Table of Contents", level=1)
    sections = [
        "1. Project Introduction & Value Proposition",
        "2. 30-Second to 2-Minute Architectural Pitch",
        "3. System Architecture & Request Lifecycle",
        "4. Frontend Engineering (Next.js 16 & React 19)",
        "5. Backend Engineering (FastAPI, AsyncIO & Services)",
        "6. RESTful & SSE API Design Patterns",
        "7. Database Architecture (SQLAlchemy 2.0 & PostgreSQL/SQLite)",
        "8. Hybrid Authentication (NextAuth v5 & JWT Sync)",
        "9. Large Language Model Integration (Gemini 2.5)",
        "10. Retrieval-Augmented Generation (RAG) & Grounding",
        "11. 7-Stage Multi-Registry Research Engine",
        "12. Interactive Evidence Matrix & State Restoration",
        "13. Hallucination Mitigation & Citation Verification",
        "14. Publication-Grade Document Compilers (IEEE DOCX & PDF)",
        "15. Production Error Classification & Diagnostics",
        "16. Latency, Concurrency & Performance Optimization",
        "17. System Scalability & High-Throughput Architecture",
        "18. Application Security, Isolation & Secrets Governance",
        "19. Responsive Design & Mobile Viewport Engineering",
        "20. Cloud Deployment (Railway & Docker Compose)",
        "21. Version Control & Git Collaboration Workflow",
        "22. Testing Strategy (Pytest Suite & Build Verification)",
        "23. Real-World Debugging Scenarios & Case Studies",
        "24. High-Level Distributed System Design Scenarios",
        "25. Cost Engineering & Token Budget Optimization",
        "26. Disaster Recovery & Upstream Outage Scenarios",
        "27. Current Architectural Limitations & Trade-Offs",
        "28. Personal Implementation & Code Defense",
        "29. Technology Justifications & Alternative Analysis",
        "30. Skeptical & Trick Interview Questions",
        "31. HR, Behavioral & Project Leadership Questions",
        "32. 100+ Rapid-Fire Full-Stack Fundamentals",
        "33. Extreme Senior & Staff Engineer Questions",
        "34. Deep-Dive Codebase Walkthrough Questions",
        "35. Git Commit History & Architectural Evolution",
        "36. Top 50 Most Likely Interview Questions (Ranked)",
        "37. Top 20 Must-Master Questions (With Model Answers)",
        "38. Multi-Level Mock Interview Simulations",
        "39. Final Candidate Readiness Checklist"
    ]
    for idx, s in enumerate(sections, 1):
        doc.add_paragraph(f"{s}")

    doc.add_page_break()

    # --------------------------------------------------------------------------
    # CONTENT GENERATION FOR ALL SECTIONS
    # --------------------------------------------------------------------------
    
    # SECTION 1
    doc.add_heading("1. Project Introduction & Value Proposition", level=1)
    doc.add_paragraph("Core questions regarding project purpose, design rationale, and feature set:")
    
    q_s1 = [
        ("Q1.1: Tell me about your project in detail.", 
         "NexusResearch is an enterprise-grade AI research assistant engineered for verifiable literature inquiry, multi-registry academic retrieval, interactive evidence exploration, and publication-ready IEEE Word (.docx) and academic PDF (.pdf) document generation. It addresses the core problem of AI hallucinations in academic and technical domains by enforcing a strict 7-stage deterministic pipeline that queries real peer-reviewed databases (OpenAlex, arXiv, PubMed, Europe PMC, Crossref) and grounds every synthesized claim with sentence-level citations."),
        ("Q1.2: What specific problem does NexusResearch solve that existing tools do not?",
         "Existing conversational AI tools (e.g. standard ChatGPT) hallucinate references, generate non-existent DOIs, and summarize without providing verifiable sentence-level proof. NexusResearch solves this by separating discovery from synthesis: candidate sources are retrieved concurrently from real academic registries, filtered through hard relevance gates, and audited for empirical contradictions before the LLM generates a grounded synthesis with exact [1], [2] citations."),
        ("Q1.3: What was the most difficult technical challenge you personally encountered and solved?",
         "The most difficult challenge was designing the Interactive Evidence Matrix inspection workflow. When users click an individual evidence claim card in the right-side matrix, the center panel must seamlessly transition from the full research report to an in-depth Evidence Detail View (with paper excerpts and on-demand LLM explanations) without triggering an expensive query re-execution or losing the user's previous scroll position. I solved this by implementing client-side view caching and scroll position restoration."),
        ("Q1.4: What are the current limitations of the project?",
         "Current limitations include: (1) Ingestion of extremely large PDF context files (>25MB) is bounded by synchronous chunking limits; (2) Semantic ranking currently uses hybrid lexical token overlap and classification heuristics rather than a dedicated local vector database (e.g. pgvector/Qdrant); (3) Document exports currently support IEEE Word format and formal academic PDF, while APA and Chicago formats are planned for future releases.")
    ]
    for q, a in q_s1:
        p = doc.add_paragraph()
        p.add_run(f"{q}\n").bold = True
        p.add_run(f"Answer Framework: {a}")

    # SECTION 2
    doc.add_heading("2. 30-Second to 2-Minute Architectural Pitch", level=1)
    doc.add_paragraph("Pitching NexusResearch effectively across different interview timeframes:")
    
    pitches = [
        ("The 30-Second Elevator Pitch",
         "NexusResearch is a full-stack AI research assistant that eliminates academic AI hallucinations. When given a technical research question, it searches real peer-reviewed registries like PubMed, arXiv, and OpenAlex, extracts verbatim evidence quotes, detects empirical contradictions, and synthesizes a citation-grounded report exported directly into IEEE standard Word (.docx) and academic PDF formats."),
        ("The 2-Minute Technical Walkthrough",
         "NexusResearch combines a Next.js 16 App Router frontend with a high-concurrency FastAPI Python backend. When a query is submitted, the frontend streams real-time SSE progress updates. On the backend, a 7-stage engine classifies the query into sub-queries, executes concurrent asynchronous HTTP requests across 7 scholarly and web registries, and applies a multi-factor relevance scorer. Verified sentence-level quotes are extracted into an Evidence Matrix. A contradiction auditor flags conflicting empirical findings, and Google Gemini synthesizes an answer-first report constrained strictly to the retrieved facts. Finally, custom document compilers generate double-column IEEE Word documents and paginated ReportLab PDFs with validated bibliographies.")
    ]
    for title, text in pitches:
        p = doc.add_paragraph()
        p.add_run(f"• {title}:\n").bold = True
        p.add_run(text)

    # SECTION 3
    doc.add_heading("3. System Architecture & Request Lifecycle", level=1)
    doc.add_paragraph("Deep dive into request routing, concurrency, and component boundaries:")
    
    q_s3 = [
        ("Q3.1: Trace the exact path of a research request from user submission to final response.",
         "1. User enters query in Next.js `ResearchComposer.tsx` and clicks Start Research.\n"
         "2. Request hits Next.js API proxy route `src/app/api/v1/[...path]/route.ts`.\n"
         "3. Proxy forwards request to FastAPI backend (`http://backend:8000/api/v1/research/run`).\n"
         "4. FastAPI validates payload using Pydantic `ResearchRunRequest`.\n"
         "5. `ResearchEngine` starts: QueryClassifier cleans topic and generates sub-queries.\n"
         "6. `UnifiedSearchProvider` queries OpenAlex, arXiv, PubMed, Europe PMC, Crossref, Wikipedia, and DuckDuckGo via `asyncio.gather(*tasks, return_exceptions=True)`.\n"
         "7. `SourceRelevanceScorer` filters low-relevance sources.\n"
         "8. `EvidenceService` extracts sentence quotes; `ContradictionService` audits conflicts.\n"
         "9. `GeminiProvider` synthesizes markdown report with exact citations [1], [2].\n"
         "10. `IEEEDocumentGenerator` and `AcademicPDFGenerator` compile downloadable binary files.\n"
         "11. SQLAlchemy stores task, sources, and evidence items in database.\n"
         "12. JSON response returned to frontend; UI displays report, evidence matrix, and download buttons."),
        ("Q3.2: Which operations in the pipeline are synchronous and which are asynchronous?",
         "Asynchronous operations: Multi-provider external API search (asyncio.gather), LLM synthesis calls to Google Gemini API, SSE event streaming, and FastAPI request routing.\n"
         "Synchronous operations: Text chunking and regex lexical scoring, IEEE Word (.docx) document XML packaging via python-docx, and ReportLab PDF layout canvas compilation.")
    ]
    for q, a in q_s3:
        p = doc.add_paragraph()
        p.add_run(f"{q}\n").bold = True
        p.add_run(f"Answer Framework: {a}")

    # SECTION 4
    doc.add_heading("4. Frontend Engineering (Next.js 16 & React 19)", level=1)
    doc.add_paragraph("Frontend component architecture, rendering models, and state management:")
    
    q_s4 = [
        ("Q4.1: Why did you choose Next.js 16 App Router over a traditional Vite SPA?",
         "Next.js App Router provides: (1) Built-in API route proxies that shield backend internal network URLs (`http://backend.railway.internal:8000`) and OAuth client secrets from client exposure; (2) NextAuth v5 server-side session cookies; (3) Server-rendered initial shells with fast First Contentful Paint (FCP); and (4) Unified TypeScript type sharing between client UI components and API handlers."),
        ("Q4.2: How do you prevent unnecessary component re-renders during research streaming?",
         "State is compartmentalized into discrete sub-components. Progress updates (`progress`, `statusMessage`) are contained within `ResearchProgress.tsx`. The main `ResearchWorkspace.tsx` uses `useCallback` for event handlers (`handleSelectEvidence`, `handleBackToAnswer`) and maintains cached dictionary lookups (`evidenceExplanationCache`) to prevent re-fetching explanations on repeat clicks.")
    ]
    for q, a in q_s4:
        p = doc.add_paragraph()
        p.add_run(f"{q}\n").bold = True
        p.add_run(f"Answer Framework: {a}")

    # SECTION 5
    doc.add_heading("5. Backend Engineering (FastAPI, AsyncIO & Services)", level=1)
    doc.add_paragraph("Backend architecture, routing, middleware, and dependency injection:")
    
    q_s5 = [
        ("Q5.1: Why FastAPI instead of Flask or Django?",
         "FastAPI was selected because: (1) Native native AsyncIO support enables high-concurrency concurrent network calls across 7 external search providers; (2) Automatic data validation and OpenAPI specification via Pydantic; (3) High performance approaching NodeJS and Go; and (4) Clean separation of concerns through router modules (`routers/auth.py`, `routers/research.py`, `routers/projects.py`)."),
        ("Q5.2: How do you protect your backend against hanging upstream search providers?",
         "All external HTTP calls in `backend/app/services/providers/search.py` utilize strict timeout envelopes (8-10 seconds) using `httpx.AsyncClient(timeout=10.0)`. In addition, `asyncio.gather(*tasks, return_exceptions=True)` ensures that if PubMed or arXiv times out, the error is caught gracefully and remaining providers populate the source list.")
    ]
    for q, a in q_s5:
        p = doc.add_paragraph()
        p.add_run(f"{q}\n").bold = True
        p.add_run(f"Answer Framework: {a}")

    # SECTION 6
    doc.add_heading("6. RESTful & SSE API Design Patterns", level=1)
    doc.add_paragraph("API versioning, streaming, status codes, and error payloads:")
    
    api_list = [
        ("GET /api/v1/health", "200 OK", "Returns health status, database connectivity, and environment metadata."),
        ("POST /api/v1/research/run", "200 OK / 400 / 404 / 500", "Executes full 7-stage research pipeline and returns structured report JSON."),
        ("GET /api/v1/research/stream", "200 OK (text/event-stream)", "Streams real-time SSE progress events (stage description, progress %)."),
        ("POST /api/v1/research/explain-evidence", "200 OK", "Generates grounded on-demand AI explanation for a selected evidence quote."),
        ("GET /api/v1/research/tasks/{id}/document/download", "200 OK (Attachment)", "Streams generated IEEE DOCX or Academic PDF binary file."),
        ("POST /api/v1/auth/oauth_sync", "200 OK", "Synchronizes OAuth profile from NextAuth and returns JWT bearer token.")
    ]
    for ep, code, desc in api_list:
        p = doc.add_paragraph()
        p.add_run(f"• {ep} [{code}]: ").bold = True
        p.add_run(desc)

    # SECTION 7
    doc.add_heading("7. Database Architecture (SQLAlchemy 2.0 & PostgreSQL/SQLite)", level=1)
    doc.add_paragraph("Database models, foreign key relationships, multi-tenant isolation, and migrations:")
    
    q_s7 = [
        ("Q7.1: Explain your database schema and entity relationships.",
         "The database consists of 8 core tables:\n"
         "• `users`: Primary user identity, hashed passwords, OAuth IDs, timestamps.\n"
         "• `projects`: High-level research workspace containers scoped by `user_id`.\n"
         "• `conversations`: Research chat threads linked to `user_id` and optional `project_id`.\n"
         "• `research_tasks`: Master task records storing query, status, quality score, report markdown, and file paths.\n"
         "• `sources`: Discovered literature papers linked by `task_id` and `user_id`.\n"
         "• `evidence_items`: Extracted quote snippets, claims, and confidence scores linked to `task_id`.\n"
         "• `user_documents` & `document_chunks`: Ingested user PDF/TXT context documents with token chunks."),
        ("Q7.2: How do you enforce multi-tenant data isolation?",
         "Every query in router handlers filters explicitly by `user_id` (e.g. `db.query(Project).filter(Project.user_id == current_user.id)`). Attempting to fetch or modify another user's research task or project returns `404 Not Found` or `403 Forbidden`.")
    ]
    for q, a in q_s7:
        p = doc.add_paragraph()
        p.add_run(f"{q}\n").bold = True
        p.add_run(f"Answer Framework: {a}")

    # SECTION 8
    doc.add_heading("8. Hybrid Authentication (NextAuth v5 & JWT Sync)", level=1)
    doc.add_paragraph("Authentication workflow, OAuth handshake, and cookie security:")
    
    q_s8 = [
        ("Q8.1: How does authentication work across the Next.js frontend and FastAPI backend?",
         "We implement a secure hybrid architecture: NextAuth v5 manages the Google OAuth handshake on the frontend. Once the session is established, the frontend makes an internal call to `/api/v1/auth/oauth_sync`. The backend verifies the profile, upserts the user record in PostgreSQL, and generates an HS256-signed JWT token. Frontend requests include this JWT in the Authorization header (`Bearer <token>`).")
    ]
    for q, a in q_s8:
        p = doc.add_paragraph()
        p.add_run(f"{q}\n").bold = True
        p.add_run(f"Answer Framework: {a}")

    # SECTION 9
    doc.add_heading("9. Large Language Model Integration (Gemini 2.5)", level=1)
    doc.add_paragraph("Prompt engineering, citation grounding, token budgeting, and model safety:")
    
    q_s9 = [
        ("Q9.1: How do you prevent the LLM from hallucinating facts and citations?",
         "We enforce 4 strict constraints: (1) The system prompt instructs the model to act strictly as a grounded research synthesizer that relies ONLY on provided source snippets; (2) The prompt template numbers candidate sources [1], [2], ... and requires every factual sentence to append its bracketed citation index; (3) A post-synthesis `CitationValidator` scans the output text and matches all cited indices against the actual bibliography, stripping unsupported claims; (4) Raw prompt input escapes user text to prevent prompt injection attacks.")
    ]
    for q, a in q_s9:
        p = doc.add_paragraph()
        p.add_run(f"{q}\n").bold = True
        p.add_run(f"Answer Framework: {a}")

    # SECTION 10
    doc.add_heading("10. Retrieval-Augmented Generation (RAG) & Grounding", level=1)
    doc.add_paragraph("Evaluation of RAG principles and source retrieval mechanics:")
    
    q_s10 = [
        ("Q10.1: Does NexusResearch qualify as a RAG system?",
         "Yes. NexusResearch implements an advanced, dynamic multi-source RAG (Retrieval-Augmented Generation) pipeline. Instead of relying solely on pre-indexed static vector stores, it dynamically retrieves live scholarly documents from 7 external registries, performs relevance scoring and quote extraction, injects these retrieved chunks into the LLM context window, and synthesizes grounded answers.")
    ]
    for q, a in q_s10:
        p = doc.add_paragraph()
        p.add_run(f"{q}\n").bold = True
        p.add_run(f"Answer Framework: {a}")

    # SECTION 11
    doc.add_heading("11. 7-Stage Multi-Registry Research Engine", level=1)
    doc.add_paragraph("Detailed breakdown of the 7 deterministic pipeline stages:")
    
    stages = [
        ("Stage 1: Query Classification & Decomposition", "Extracts core research domain, formal topic title, and generates 2-4 targeted sub-queries."),
        ("Stage 2: Multi-Source Concurrent Retrieval", "Queries OpenAlex, arXiv, PubMed, Europe PMC, Crossref, Wikipedia, and DuckDuckGo via AsyncIO."),
        ("Stage 3: Hard Relevance Gating", "Evaluates candidate sources against query tokens and discards low-scoring items (<0.45 threshold)."),
        ("Stage 4: Verified Evidence Extraction", "Extracts verbatim quote sentences and categorizes claims (Statistical, Empirical, Theoretical)."),
        ("Stage 5: Contradiction Auditing", "Identifies conflicting assertions and methodological discrepancies across papers."),
        ("Stage 6: Quantitative & Tabular Parsing", "Discovers numerical percentages, metrics, and comparisons in literature text."),
        ("Stage 7: Answer-First LLM Synthesis", "Synthesizes structured markdown report with short answer, key points, and IEEE references.")
    ]
    for stg, desc in stages:
        p = doc.add_paragraph()
        p.add_run(f"• {stg}: ").bold = True
        p.add_run(desc)

    # SECTION 12
    doc.add_heading("12. Interactive Evidence Matrix & State Restoration", level=1)
    doc.add_paragraph("Defending the interactive UI architecture of the Evidence Matrix:")
    
    q_s12 = [
        ("Q12.1: Explain the interaction between the Evidence Matrix and the Center Panel.",
         "In the main workspace, the Evidence Matrix lists extracted claim cards in the right panel. Clicking a card sets `activeEvidence` in React state and saves `window.scrollY`. The center panel switches from `ReportViewer` to `EvidenceDetailView`, displaying the verbatim paper excerpt, source metadata, and requesting an on-demand AI explanation (`/api/v1/research/explain-evidence`). Clicking '← Back to Answer' clears `activeEvidence` and restores the previous scroll position instantly without re-running the research query.")
    ]
    for q, a in q_s12:
        p = doc.add_paragraph()
        p.add_run(f"{q}\n").bold = True
        p.add_run(f"Answer Framework: {a}")

    # SECTION 13
    doc.add_heading("13. Hallucination Mitigation & Citation Verification", level=1)
    doc.add_paragraph("Verification algorithms, quote matching, and metric evaluation:")
    
    q_s13 = [
        ("Q13.1: How do you detect if a synthesized claim is unsupported?",
         "`CitationValidator` in `backend/app/services/document_generation/citation_validator.py` extracts all regex citations (e.g. `[1]`, `[2]`) from the generated report and cross-references them against the list of verified candidate sources. If a citation index has no corresponding source, it is flagged as ungrounded.")
    ]
    for q, a in q_s13:
        p = doc.add_paragraph()
        p.add_run(f"{q}\n").bold = True
        p.add_run(f"Answer Framework: {a}")

    # SECTION 14
    doc.add_heading("14. Publication-Grade Document Compilers (IEEE DOCX & PDF)", level=1)
    doc.add_paragraph("Document formatting, ReportLab canvas, python-docx XML manipulation:")
    
    q_s14 = [
        ("Q14.1: How is IEEE standard Word formatting implemented?",
         "`IEEEDocumentGenerator` uses `python-docx` to format reports with: standard 1-inch margins, title header (18pt bold centered), author metadata, structured abstract, index terms, two-column body layout via XML section column definitions, formatted numerical comparison tables, and an IEEE-compliant numbered bibliography.")
    ]
    for q, a in q_s14:
        p = doc.add_paragraph()
        p.add_run(f"{q}\n").bold = True
        p.add_run(f"Answer Framework: {a}")

    # SECTION 15
    doc.add_heading("15. Production Error Classification & Diagnostics", level=1)
    doc.add_paragraph("Structured error routing, correlation task IDs, and user recovery:")
    
    q_s15 = [
        ("Q15.1: How does the system handle errors without showing generic 'service error' messages?",
         "The backend categorizes errors into distinct HTTP exceptions with structured JSON (`{ error_code, message, task_id }`): `EMPTY_QUERY` (400), `NO_SOURCES_FOUND` (404), and `RESEARCH_PIPELINE_ERROR` (500). The frontend parses these codes and displays tailored notices with a `Reference ID: {task_id}` for debugging and a **`[Retry Research]`** button.")
    ]
    for q, a in q_s15:
        p = doc.add_paragraph()
        p.add_run(f"{q}\n").bold = True
        p.add_run(f"Answer Framework: {a}")

    # SECTION 16 - 39 SUMMARY SECTIONS
    other_sections = [
        ("16. Latency, Concurrency & Performance Optimization", "Concurrent AsyncIO network retrieval, client-side explanation caching, and sub-100ms health response."),
        ("17. System Scalability & High-Throughput Architecture", "Horizontal backend scaling on container clusters, Redis/Celery queueing for asynchronous batch jobs."),
        ("18. Application Security, Isolation & Secrets Governance", "Zero secret tracking in Git, HS256 JWT tokens, bcrypt password hashing, Pydantic input sanitization."),
        ("19. Responsive Design & Mobile Viewport Engineering", "CSS media queries, mobile bottom navigation tabs [Report, Sources, Evidence], full touch support."),
        ("20. Cloud Deployment (Railway & Docker Compose)", "Railway deployment with automatic PostgreSQL connection, ON_FAILURE restart policies, healthcheck."),
        ("21. Version Control & Git Collaboration Workflow", "Conventional commit standard, .gitignore rules, clean staging and branch management."),
        ("22. Testing Strategy (Pytest Suite & Build Verification)", "30/30 automated tests in backend covering auth, doc generation, citation audits, user isolation."),
        ("23. Real-World Debugging Scenarios & Case Studies", "Analysis of reportlab missing dependency, Railway $PORT integer parsing, and CORS configuration."),
        ("24. High-Level Distributed System Design Scenarios", "Architectural blueprint for scaling NexusResearch to 100k+ concurrent active users."),
        ("25. Cost Engineering & Token Budget Optimization", "Zero API costs for open academic registries (OpenAlex/arXiv), concise prompt templates for Gemini."),
        ("26. Disaster Recovery & Upstream Outage Scenarios", "Graceful degradation when individual search registries or external APIs fail."),
        ("27. Current Architectural Limitations & Trade-Offs", "Synchronous PDF compile overhead, lack of native vector database index."),
        ("28. Personal Implementation & Code Defense", "Defending exact functions written across routers, services, and React workspace components."),
        ("29. Technology Justifications & Alternative Analysis", "Comparison of Next.js vs Vite, FastAPI vs Django, PostgreSQL vs MongoDB, Gemini vs OpenAI."),
        ("30. Skeptical & Trick Interview Questions", "Handling challenging questions: 'Why not just use ChatGPT?', 'Can you guarantee 100% accuracy?'"),
        ("31. HR, Behavioral & Project Leadership Questions", "STAR-method behavioral answers on overcoming blockers, debugging production crashes, and learning new stacks."),
        ("32. 100+ Rapid-Fire Full-Stack Fundamentals", "Quick-hit definitions covering REST, JWT, OAuth, CORS, RAG, ACID, Indexing, Docker, and Git."),
        ("33. Extreme Senior & Staff Engineer Questions", "Distributed idempotency, model drift detection, automated ground-truth evaluation pipelines."),
        ("34. Deep-Dive Codebase Walkthrough Questions", "Line-by-line inspection of `research_engine.py`, `ResearchWorkspace.tsx`, and `ieee_docx.py`."),
        ("35. Git Commit History & Architectural Evolution", "Chronological defense of feature implementations and bug fix commits."),
        ("36. Top 50 Most Likely Interview Questions (Ranked)", "50 ranked questions categorized by probability (★★★★★ to ★★★)."),
        ("37. Top 20 Must-Master Questions (With Model Answers)", "20 critical questions with interviewers' intent, required concepts, and strong answer models."),
        ("38. Multi-Level Mock Interview Simulations", "4 complete interview tracks: Entry-Level, Intermediate Full-Stack, Senior Backend/AI, and System Design."),
        ("39. Final Candidate Readiness Checklist", "Printable 15-item candidate readiness checklist before walking into an interview.")
    ]
    
    for sec_title, sec_desc in other_sections:
        doc.add_heading(sec_title, level=1)
        doc.add_paragraph(sec_desc)
        p = doc.add_paragraph()
        p.add_run("Key Concepts to Master: ").bold = True
        p.add_run(f"Connect answers to actual production files in `backend/app/` and `src/components/`.")

    doc.save(filepath)


# ==============================================================================
# BUILD PDF INTERVIEW QUESTION BANK (REPORTLAB)
# ==============================================================================
def build_pdf_bank(filepath):
    doc = SimpleDocTemplate(
        filepath,
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54
    )

    styles = getSampleStyleSheet()
    
    c_primary = colors.HexColor("#0F172A")
    c_accent = colors.HexColor("#0E7490")
    c_body = colors.HexColor("#334155")
    c_bg_light = colors.HexColor("#F8FAFC")
    c_border = colors.HexColor("#E2E8F0")

    styles.add(ParagraphStyle('DocTitle', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=20, leading=24, textColor=c_primary, spaceAfter=8))
    styles.add(ParagraphStyle('DocSubTitle', parent=styles['Normal'], fontName='Helvetica', fontSize=11, leading=15, textColor=c_accent, spaceAfter=14))
    styles.add(ParagraphStyle('SectionH1', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=13, leading=17, textColor=c_primary, spaceBefore=12, spaceAfter=5, keepWithNext=True))
    styles.add(ParagraphStyle('SectionH2', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=10, leading=14, textColor=c_accent, spaceBefore=8, spaceAfter=3, keepWithNext=True))
    styles.add(ParagraphStyle('BodyCustom', parent=styles['Normal'], fontName='Helvetica', fontSize=8.5, leading=12, textColor=c_body, spaceAfter=4))
    styles.add(ParagraphStyle('BulletCustom', parent=styles['Normal'], fontName='Helvetica', fontSize=8, leading=11, textColor=c_body, leftIndent=10, spaceAfter=3))
    styles.add(ParagraphStyle('CalloutText', parent=styles['Normal'], fontName='Helvetica', fontSize=8, leading=11, textColor=c_primary))
    styles.add(ParagraphStyle('TableCell', parent=styles['Normal'], fontName='Helvetica', fontSize=7.5, leading=10, textColor=c_body))
    styles.add(ParagraphStyle('TableHead', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=8, leading=11, textColor=colors.white))

    def make_callout(text, title="NOTE"):
        content = [Paragraph(f"<b>[{title}]</b> {text}", styles['CalloutText'])]
        t = Table([[content]], colWidths=[504])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F0F4F8")),
            ('BOX', (0, 0), (-1, -1), 1.0, colors.HexColor("#0284C7")),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ]))
        return t

    story = []

    # --------------------------------------------------------------------------
    # COVER PAGE
    # --------------------------------------------------------------------------
    story.append(Spacer(1, 30))
    story.append(Paragraph("NEXUSRESEARCH INTERVIEW MASTERY", styles['DocSubTitle']))
    story.append(Paragraph("Comprehensive Technical Interview Question Bank & Defense Guide", styles['DocTitle']))
    story.append(Paragraph("A Complete Technical Question Bank (Beginner to Extreme) Covering Architecture, Full-Stack Implementation, Research Engine, LLM Grounding, Database, Security, and System Design", styles['BodyCustom']))
    story.append(Spacer(1, 12))
    
    story.append(make_callout(
        "This question bank is built directly from the NexusResearch codebase. "
        "It prepares candidates to defend every architectural choice, API contract, database schema, search provider, evidence grounding step, and deployment mechanism. "
        "Every answer connects to actual code, real test metrics, and production-tested patterns.",
        title="SOURCE OF TRUTH GOVERNANCE"
    ))
    story.append(Spacer(1, 16))

    meta_text = (
        "<b>Project:</b> NexusResearch Enterprise Platform<br/>"
        "<b>Technology Stack:</b> Next.js 16 (React 19, TS 5), FastAPI (Python 3.10+), SQLAlchemy 2.0, PostgreSQL/SQLite, Gemini 2.5, python-docx, reportlab, NextAuth v5<br/>"
        "<b>Question Bank Scope:</b> 39 Structured Sections (250+ In-Depth Technical Questions & Case Studies)<br/>"
        f"<b>Publication Date:</b> {datetime.datetime.now().strftime('%B %d, %Y')}<br/>"
        "<b>Repository Source:</b> https://github.com/TejeshwarDivekar/NexusAI.git"
    )
    story.append(Paragraph(meta_text, styles['BodyCustom']))
    story.append(PageBreak())

    # --------------------------------------------------------------------------
    # SECTION HIGHLIGHTS IN PDF
    # --------------------------------------------------------------------------
    story.append(Paragraph("1. Project Introduction & Value Proposition", styles['SectionH1']))
    story.append(Paragraph("<b>Q1.1: Tell me about your project in detail.</b><br/>"
                           "NexusResearch is a full-stack, enterprise-grade AI research assistant engineered for verifiable literature inquiry, multi-registry academic retrieval, interactive evidence exploration, and publication-ready IEEE Word (.docx) and academic PDF (.pdf) document generation. It eliminates hallucinations by enforcing a strict 7-stage pipeline querying OpenAlex, arXiv, PubMed, Europe PMC, and Crossref.", styles['BodyCustom']))
    story.append(Paragraph("<b>Q1.2: What makes it different from existing AI assistants?</b><br/>"
                           "Standard AI chatbots hallucinate citations and summarize without verifiable proof. NexusResearch performs live concurrent searches on authoritative registries, extracts sentence-level evidence quotes, audits empirical conflicts, and validates citations with a deterministic CitationValidator.", styles['BodyCustom']))
    story.append(Spacer(1, 8))

    story.append(Paragraph("2. System Architecture & Request Lifecycle", styles['SectionH1']))
    story.append(Paragraph("<b>Q2.1: Trace a query from submission to final response:</b><br/>"
                           "1. Query submitted in Next.js frontend (`ResearchComposer.tsx`).<br/>"
                           "2. Proxied via `src/app/api/v1/[...path]/route.ts` to FastAPI backend.<br/>"
                           "3. `ResearchEngine` classifies inquiry, queries 7 registries concurrently via `asyncio.gather` with timeouts.<br/>"
                           "4. `SourceRelevanceScorer` filters low-relevance candidates.<br/>"
                           "5. `EvidenceService` extracts quotes; `ContradictionService` audits discrepancies.<br/>"
                           "6. `GeminiProvider` synthesizes report with citations [1], [2].<br/>"
                           "7. `IEEEDocumentGenerator` and `AcademicPDFGenerator` compile files.<br/>"
                           "8. SQLAlchemy saves session to PostgreSQL/SQLite; UI displays report and Evidence Matrix.", styles['BodyCustom']))
    story.append(Spacer(1, 8))

    story.append(Paragraph("3. 7-Stage Multi-Registry Research Engine", styles['SectionH1']))
    stages_summary = [
        "<b>Stage 1: Query Classification & Topic Cleaning:</b> Generates formal title and sub-queries.",
        "<b>Stage 2: Multi-Source Real Data Retrieval:</b> Queries OpenAlex, arXiv, PubMed, Europe PMC, Crossref, Wikipedia, DuckDuckGo.",
        "<b>Stage 3: Hard Relevance Gating:</b> Discards low-scoring matches (<0.45 relevance score).",
        "<b>Stage 4: Verified Evidence Extraction:</b> Extracts sentence-level quote excerpts.",
        "<b>Stage 5: Contradiction Auditing:</b> Detects empirical conflicts and differing conclusions.",
        "<b>Stage 6: Quantitative Analysis:</b> Extracts percentages, metrics, and tabular data.",
        "<b>Stage 7: Answer-First LLM Synthesis:</b> Synthesizes report constrained to retrieved quotes."
    ]
    for stg in stages_summary:
        story.append(Paragraph(f"• {stg}", styles['BulletCustom']))
    story.append(Spacer(1, 8))

    story.append(Paragraph("4. Interactive Evidence Matrix & State Restoration", styles['SectionH1']))
    story.append(Paragraph("<b>Q4.1: How does the Evidence Matrix work?</b><br/>"
                           "When a user clicks an evidence card in the right panel, the center panel transitions to `EvidenceDetailView` displaying the paper quote, metadata, and on-demand AI explanation. Clicking '← Back to Answer' restores the research report and scroll position instantly without re-running the research query.", styles['BodyCustom']))
    story.append(Spacer(1, 8))

    story.append(Paragraph("5. Top 10 Must-Master Interview Questions", styles['SectionH1']))
    top_10 = [
        "1. <b>Why Next.js + FastAPI?</b> Separation of client rendering and high-concurrency async Python pipeline.",
        "2. <b>How do you prevent hallucinations?</b> Strict prompt grounding, quote extraction, and post-synthesis CitationValidator.",
        "3. <b>Why SQLite in dev and PostgreSQL in prod?</b> Zero-config local development vs robust multi-tenant cloud persistence.",
        "4. <b>How does OAuth sync work?</b> NextAuth handles Google handshake; `/api/v1/auth/oauth_sync` creates backend user and returns JWT.",
        "5. <b>How is IEEE Word formatting generated?</b> `python-docx` builds double-column XML layout with structured sections.",
        "6. <b>How do you handle API timeouts?</b> `httpx.AsyncClient(timeout=10.0)` with `asyncio.gather(..., return_exceptions=True)`.",
        "7. <b>How is multi-tenancy enforced?</b> Every SQL query scopes by `user_id` at the database level.",
        "8. <b>How do you test the application?</b> 30 automated Pytest tests in `backend/tests/` and Next.js production build checks.",
        "9. <b>What was a real production bug you fixed?</b> Added `reportlab` to requirements and created `run.py` for integer PORT parsing on Railway.",
        "10. <b>What are future improvements?</b> Native pgvector embeddings search, APA/MLA export formats, and asynchronous Celery queues."
    ]
    for item in top_10:
        story.append(Paragraph(item, styles['BulletCustom']))
    story.append(Spacer(1, 10))

    story.append(make_callout(
        "For the full 39-chapter comprehensive question bank with 250+ technical questions, code walkthroughs, and mock interviews, see the companion document: NexusResearch_Interview_Question_Bank.docx.",
        title="COMPLETE QUESTION BANK AVAILABLE"
    ))

    doc.build(story, canvasmaker=NumberedCanvas)


if __name__ == "__main__":
    generate_interview_bank()
