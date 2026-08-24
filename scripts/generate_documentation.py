"""
NexusResearch Complete Technical Architecture & Reconstruction Guide Generator
Generates both:
1. NexusResearch_Complete_Technical_Documentation.docx
2. NexusResearch_Complete_Technical_Documentation.pdf
"""

import os
import sys
import datetime
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether, HRFlowable
)
from reportlab.pdfgen import canvas

def create_documents():
    print("Starting NexusResearch technical documentation generation...")
    docx_path = "NexusResearch_Complete_Technical_Documentation.docx"
    pdf_path = "NexusResearch_Complete_Technical_Documentation.pdf"
    
    # 1. Build DOCX
    build_docx(docx_path)
    print(f"Generated DOCX: {docx_path}")
    
    # 2. Build PDF
    build_pdf(pdf_path)
    print(f"Generated PDF: {pdf_path}")

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_callout_docx(doc, text, title="NOTE", fill_hex="F0F4F8", border_hex="0066CC"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_background(cell, fill_hex)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r_title = p.add_run(f"[{title}] ")
    r_title.bold = True
    r_title.font.color.rgb = RGBColor(0, 51, 102)
    r_title.font.size = Pt(9.5)
    r_text = p.add_run(text)
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(30, 41, 59)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_code_block_docx(doc, code_text, language=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_background(cell, "1E293B")
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(code_text.strip())
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(226, 232, 240)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def build_docx(docx_path):
    doc = Document()
    
    # Page setup - Standard Letter, 0.8 inch margins
    sections = doc.sections
    for section in sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
        # Header / Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hr = hp.add_run("NexusResearch — Complete Technical Architecture & Reconstruction Guide")
        hr.font.size = Pt(8.5)
        hr.font.color.rgb = RGBColor(148, 163, 184)
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run("Confidential & Proprietary — Ground Truth Engineering Manual — 2026")
        fr.font.size = Pt(8.0)
        fr.font.color.rgb = RGBColor(148, 163, 184)

    # Base Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(30, 41, 59)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)

    # ----------------------------------------------------
    # TITLE PAGE
    # ----------------------------------------------------
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(36)
    
    p_badge = doc.add_paragraph()
    p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_badge = p_badge.add_run("PRODUCTION SYSTEM ARCHITECTURE & ENGINEERING MANUAL")
    r_badge.bold = True
    r_badge.font.size = Pt(10)
    r_badge.font.color.rgb = RGBColor(0, 102, 204)
    
    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_main_title = p_main_title.add_run("NexusResearch — Complete Technical Architecture,\nDevelopment & Reconstruction Guide")
    r_main_title.bold = True
    r_main_title.font.size = Pt(22)
    r_main_title.font.color.rgb = RGBColor(15, 23, 42)
    p_main_title.paragraph_format.space_after = Pt(12)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("How to Understand, Run, Rebuild, Maintain and Deploy the Complete Real-Time AI Research Platform from Scratch Without AI")
    r_sub.font.size = Pt(12)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(71, 85, 105)
    p_sub.paragraph_format.space_after = Pt(28)
    
    # Metadata Table
    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.columns[0].width = Inches(2.2)
    meta_table.columns[1].width = Inches(4.3)
    
    meta_data = [
        ("Author & System Architect:", "Tejeshwar Divekar"),
        ("Project Repository:", "https://github.com/TejeshwarDivekar/NexusAI.git"),
        ("Live Production Frontend:", "https://frontend-production-2df5.up.railway.app"),
        ("Live Production API:", "https://backend-production-873b.up.railway.app/api/v1/health"),
        ("Architecture Version:", "1.0.0-Production (Next.js 16.3.1 + FastAPI 1.0.0 + Gemini 2.5)"),
        ("Release Date:", datetime.datetime.now().strftime("%B %d, %Y"))
    ]
    for idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[idx]
        c0, c1 = row.cells[0], row.cells[1]
        set_cell_background(c0, "F8FAFC")
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c0, 60, 60, 100, 100)
        set_cell_margins(c1, 60, 60, 100, 100)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = RGBColor(51, 65, 85)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = RGBColor(15, 23, 42)
        if "http" in val:
            r1.font.color.rgb = RGBColor(0, 102, 204)
            r1.underline = True

    doc.add_page_break()

    # ----------------------------------------------------
    # TABLE OF CONTENTS
    # ----------------------------------------------------
    h_toc = doc.add_heading("Table of Contents", level=1)
    h_toc.paragraph_format.space_before = Pt(12)
    h_toc.paragraph_format.space_after = Pt(12)
    
    toc_items = [
        ("1. Project Overview & Executive Summary", "3"),
        ("2. What NexusResearch Does: Core Capabilities", "4"),
        ("3. Core Features & Functional Architecture", "5"),
        ("4. Complete Technology Stack & Version Inventory", "7"),
        ("5. Why Each Technology Is Used (In-Depth Rationale)", "9"),
        ("6. Complete System Architecture & Communication Flow", "11"),
        ("7. Full Project Directory & File-by-File Reference", "14"),
        ("8. Frontend Architecture: Next.js 16, React 19 & Mobile UI", "18"),
        ("9. Backend Architecture: FastAPI, Middleware & Endpoints", "22"),
        ("10. Database Architecture: PostgreSQL, SQLAlchemy 2.0 & Schema", "26"),
        ("11. Authentication Architecture: NextAuth v5, Google OAuth & User Scoping", "29"),
        ("12. AI & LLM Engine: Google Gemini 2.5, Prompts & Grounding", "31"),
        ("13. Deep Research Pipeline: Multi-Stage Orchestrator", "34"),
        ("14. External Research & Academic APIs: Endpoints, Payloads & Fallbacks", "37"),
        ("15. Source Retrieval, Multi-Factor Relevance Scoring & Hard Gate", "40"),
        ("16. Evidence Matrix & Grounding Verification System", "43"),
        ("17. Answer Generation & Multi-Archetype Synthesis Engine", "45"),
        ("18. Document Generation Architecture: IEEE Two-Column DOCX", "47"),
        ("19. File Uploads & Local Document Processing Pipeline", "49"),
        ("20. User History, Workspaces & Project Isolation", "51"),
        ("21. Complete Environment Variables & Secrets Reference", "53"),
        ("22. Local Development Setup from Scratch", "55"),
        ("23. Running in VS Code & Windows Environment Guide", "58"),
        ("24. First-Run Verification Checklist", "60"),
        ("25. Testing Strategy & Automated Benchmark Suites", "62"),
        ("26. Logging, Observability & Debugging Guide", "64"),
        ("27. Cloud Deployment Architecture on Railway", "66"),
        ("28. Production Configuration & Next.js API Proxying", "68"),
        ("29. Security, Isolation & Compliance Measures", "70"),
        ("30. Manual Step-by-Step Rebuild Guide (15 Phases)", "72"),
        ("31. Exact Recommended Rebuild Order", "76"),
        ("32. Beginner Guide to Core Technical Concepts", "78"),
        ("33. Comprehensive Troubleshooting Guide", "81"),
        ("34. Maintenance, Schema Upgrades & Model Operations", "84"),
        ("35. Current Known Limitations & Trade-offs", "86"),
        ("36. Recommended Future Improvements", "88"),
        ("37. Technical Glossary of Terms", "90"),
        ("38. Command Reference & Final Reproduction Checklist", "92"),
    ]
    
    t_toc = doc.add_table(rows=len(toc_items), cols=2)
    t_toc.columns[0].width = Inches(5.8)
    t_toc.columns[1].width = Inches(0.7)
    for idx, (title, page) in enumerate(toc_items):
        row = t_toc.rows[idx]
        set_cell_background(row.cells[0], "FFFFFF" if idx % 2 == 0 else "F8FAFC")
        set_cell_background(row.cells[1], "FFFFFF" if idx % 2 == 0 else "F8FAFC")
        set_cell_margins(row.cells[0], 20, 20, 40, 40)
        set_cell_margins(row.cells[1], 20, 20, 40, 40)
        
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(title)
        r0.font.size = Pt(9.5)
        if title.startswith("1.") or title.startswith("6.") or title.startswith("8.") or title.startswith("9.") or title.startswith("13.") or title.startswith("30."):
            r0.bold = True
            r0.font.color.rgb = RGBColor(15, 23, 42)
        else:
            r0.font.color.rgb = RGBColor(51, 65, 85)
            
        p1 = row.cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1 = p1.add_run(page)
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_page_break()

    # ----------------------------------------------------
    # CHAPTERS CONTENT BUILDER
    # ----------------------------------------------------
    
    # CHAPTER 1
    doc.add_heading("1. Project Overview & Executive Summary", level=1)
    doc.add_paragraph(
        "NexusResearch is an enterprise-grade, autonomous real-time research assistant and academic intelligence workspace. "
        "Unlike conventional LLM wrappers that merely generate plausible-sounding text, NexusResearch implements an evidence-grounded, "
        "multi-stage verification pipeline. It retrieves authentic peer-reviewed literature, live web disclosures, and uploaded user documents, "
        "subjects all candidate sources to a strict hard relevance gate, maps exact quote-level evidence into a claim matrix, "
        "and synthesizes structured, publication-grade intelligence with interactive citations and professional IEEE Word (.docx) document generation."
    )
    add_callout_docx(
        doc,
        "Ground Truth Core Principle: A source is never assumed relevant simply because an API returned it. "
        "The system explicitly separates Source Validity (registry authentication) from Query Relevance (topical alignment), "
        "filtering out ungrounded or distantly related papers before synthesis begins.",
        title="PRIMARY ARCHITECTURAL MANDATE"
    )
    
    p = doc.add_paragraph()
    p.add_run("Key Operational Problems Solved:\n").bold = True
    p.add_run("1. LLM Hallucinations: Eliminate fabricated citations, authors, DOIs, and statistics through strict evidence grounding.\n")
    p.add_run("2. Noise & False Matches: Stop irrelevant search results (e.g. veterinary papers in AI farming queries) via multi-signal semantic filtering.\n")
    p.add_run("3. Cognitive Overload: Transform dense academic literature into an answer-first experience with key takeaways, deep dives, and evidence tracing.\n")
    p.add_run("4. Document Publishing Bottleneck: Automatically generate IEEE formatted two-column Word manuscripts ready for formal peer submission.")

    # CHAPTER 2
    doc.add_heading("2. What NexusResearch Does: Core Capabilities", level=1)
    doc.add_paragraph(
        "NexusResearch handles diverse research archetypes dynamically based on user intent and domain categorization:"
    )
    p = doc.add_paragraph()
    p.add_run("• Academic & Scientific Deep Dives: ").bold = True
    p.add_run("Performs multi-query expansion across OpenAlex, arXiv, PubMed, Europe PMC, and Crossref. Synthesizes formal state-of-the-art reports with mathematical formulations, methodologies, and benchmarking data.\n")
    p.add_run("• Real-Time & Live Market Intelligence: ").bold = True
    p.add_run("Fetches real-time web data via DuckDuckGo and Wikipedia with live timestamp verification for breaking news, stock metrics, and technology releases.\n")
    p.add_run("• Step-by-Step Roadmaps & Curricula: ").bold = True
    p.add_run("Generates structured milestone tracks (Foundational -> Intermediate -> Advanced) with concrete prerequisites, toolsets, and project milestones.\n")
    p.add_run("• Comparative Analyses: ").bold = True
    p.add_run("Constructs side-by-side dimensional matrices evaluating trade-offs, performance benchmarks, architectural bottlenecks, and cost profiles.\n")
    p.add_run("• Numerical & Statistical Calculations: ").bold = True
    p.add_run("Executes deterministic Python calculations (mean, standard deviation, variance, compound annual growth rate) alongside scientific analysis.")

    # CHAPTER 3
    doc.add_heading("3. Core Features & Functional Architecture", level=1)
    doc.add_paragraph("The platform is architected around 7 core subsystems:")
    p = doc.add_paragraph()
    p.add_run("1. Answer-First Research Workspace: ").bold = True
    p.add_run("Desktop 3-panel workspace (Sources | Main Answer | Evidence Matrix) and Purpose-Built Mobile Experience with segmented navigation and 44px+ touch targets.\n")
    p.add_run("2. Multi-Signal Source Relevance Scorer: ").bold = True
    p.add_run("Multi-factor heuristic scoring engine with negative domain disqualification and hard gating (threshold >= 0.48).\n")
    p.add_run("3. Evidence Matrix & Claim Grounding: ").bold = True
    p.add_run("Direct mapping from high-level assertions to exact source excerpts, why-relevant rationales, and qualitative confidence scores.\n")
    p.add_run("4. Contradiction & Consensus Engine: ").bold = True
    p.add_run("Identifies conflicting claims, methodological divergences, and disputes across retrieved literature.\n")
    p.add_run("5. IEEE Two-Column DOCX Document Builder: ").bold = True
    p.add_run("Generates fully compliant IEEE Word documents with professional styling, title blocks, abstract, and citation bibliographies.\n")
    p.add_run("6. NextAuth v5 & User Isolation: ").bold = True
    p.add_run("Secure Google and GitHub OAuth authentication with PostgreSQL user-level project isolation and conversation persistence.\n")
    p.add_run("7. Dynamic Next.js API Proxying: ").bold = True
    p.add_run("Edge route proxy at /api/v1/[...path] forwarding requests, handling large timeouts (120s), and passing binary document streams seamlessly.")

    # CHAPTER 4
    doc.add_heading("4. Complete Technology Stack & Version Inventory", level=1)
    doc.add_paragraph("The following table lists every production dependency extracted directly from package.json and requirements.txt:")
    
    tech_data = [
        ("Next.js", "16.3.1", "Frontend Framework", "App Router, SSR, Route Proxying", "package.json"),
        ("React & React-DOM", "19.2.8", "UI Library", "Component tree, hooks, hydration", "package.json"),
        ("Tailwind CSS", "4.0.0", "Styling Engine", "CSS-first token architecture, dark mode", "package.json"),
        ("NextAuth", "5.0.0-beta.32", "Authentication", "Google/GitHub OAuth, JWT sessions", "package.json"),
        ("Lucide React", "1.31.0", "Iconography", "Modern iconography across UI", "package.json"),
        ("React Markdown", "10.1.0", "Markdown Parser", "Renders rich synthesis & LaTeX tables", "package.json"),
        ("Zod", "4.4.3", "Validation", "Frontend schema and type validation", "package.json"),
        ("FastAPI", ">=0.110.0", "Backend Framework", "High-performance async REST API", "requirements.txt"),
        ("Uvicorn", ">=0.28.0", "ASGI Server", "Production async HTTP worker", "requirements.txt"),
        ("SQLAlchemy", ">=2.0.28", "ORM & Database", "Database modeling & transaction queries", "requirements.txt"),
        ("Psycopg2-Binary", ">=2.9.9", "PostgreSQL Driver", "High-throughput PostgreSQL connection", "requirements.txt"),
        ("Pydantic", ">=2.6.0", "Data Validation", "Request/Response schema enforcement", "requirements.txt"),
        ("Python-JOSE & Passlib", "3.3.0 / 1.7.4", "Security & Auth", "JWT token encoding & bcrypt hashing", "requirements.txt"),
        ("HTTPX", ">=0.27.0", "HTTP Client", "Async API requests to search providers", "requirements.txt"),
        ("PyPDF", ">=4.1.0", "PDF Ingestion", "Extracts text from uploaded user papers", "requirements.txt"),
        ("DuckDuckGo-Search", ">=5.0.0", "Web Search", "Live web search & news discovery", "requirements.txt"),
        ("Python-Docx", ">=1.2.0", "Document Builder", "Programmatic IEEE DOCX compilation", "requirements.txt"),
        ("Google Gemini 2.5", "v1beta REST", "LLM Intelligence", "Query classification, synthesis & claims", "Backend / Env")
    ]
    
    t_tech = doc.add_table(rows=len(tech_data) + 1, cols=5)
    t_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, col in enumerate(["Technology", "Version", "Category", "Purpose in NexusResearch", "Source File"]):
        cell = t_tech.rows[0].cells[i]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, 60, 60, 80, 80)
        p = cell.paragraphs[0]
        r = p.add_run(col)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    for r_idx, row_values in enumerate(tech_data):
        row = t_tech.rows[r_idx + 1]
        for c_idx, val in enumerate(row_values):
            cell = row.cells[c_idx]
            set_cell_background(cell, "FFFFFF" if r_idx % 2 == 0 else "F8FAFC")
            set_cell_margins(cell, 40, 40, 60, 60)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8.0)
            if c_idx == 0:
                r.bold = True

    # CHAPTER 5
    doc.add_heading("5. Why Each Technology Is Used (In-Depth Rationale)", level=1)
    doc.add_paragraph("Engineering justifications for core architectural selections:")
    
    why_points = [
        ("Next.js 16 + React 19", 
         "Provides cutting-edge App Router server rendering, dynamic streaming capabilities, and built-in API proxy routing. "
         "React 19 hooks and optimistic UI updates enable instant tab switching and seamless state transitions between Sources, Answers, and Evidence without UI flicker."),
        ("FastAPI + Python 3.10+",
         "Python is the undisputed standard for scientific research algorithms, academic API parsing (XML/JSON/BibTeX), numerical analysis, and LLM orchestration. "
         "FastAPI provides native asynchronous I/O (asyncio) allowing simultaneous concurrent querying of 5+ academic search providers in parallel."),
        ("PostgreSQL + SQLAlchemy 2.0",
         "Relational integrity is critical for mapping multi-tiered academic data: Users -> Projects -> Conversations -> Research Tasks -> Claims -> Evidence Items -> Sources. "
         "SQLAlchemy 2.0 provides type-safe ORM querying, JSON column support for rich metadata, and zero data loss automated migrations."),
        ("Google Gemini 2.5 Flash / Pro",
         "Delivers an industry-leading 1M+ token context window, sub-second latency for real-time streaming synthesis, and superior JSON structured output compliance "
         "for deterministic claim-evidence mapping."),
        ("Python-Docx for IEEE Document Generation",
         "Allows complete programmatic control over binary Word XML packaging. Compiles precise IEEE standard two-column layouts, title headers, abstracts, and reference sections directly on the server without third-party office dependencies.")
    ]
    for tech, rationale in why_points:
        p = doc.add_paragraph()
        p.add_run(f"• {tech}: ").bold = True
        p.add_run(rationale)

    # CHAPTER 6
    doc.add_heading("6. Complete System Architecture & Communication Flow", level=1)
    doc.add_paragraph("The complete end-to-end data communication lifecycle operates as follows:")
    
    add_code_block_docx(doc, """
+---------------------------------------------------------------------------------------------------+
|                                     NEXUSRESEARCH ARCHITECTURE                                    |
+---------------------------------------------------------------------------------------------------+

[ Client Browser (Desktop / Mobile) ]
             |
             | (HTTPS / JSON / Auth Bearer)
             v
[ Next.js 16 Edge / App Router (Railway Frontend) ]
   |-- /login, /chat, /workspace UI (React 19)
   |-- NextAuth v5 (Google / GitHub OAuth)
   `-- Dynamic Proxy Route: /api/v1/[...path]/route.ts (120s timeout, Binary Buffering)
             |
             | (Internal Railway Network / Public HTTPS)
             v
[ FastAPI 1.0.0 Async Engine (Railway Backend) ]
   |-- Middleware: CORS, Request Timing, Correlation ID (X-Request-ID), Exception Handlers
   |-- Database Layer: SQLAlchemy 2.0 -> PostgreSQL Managed Instance
   |
   |-- [ RESEARCH ORCHESTRATOR (ResearchEngine) ]
   |     |
   |     +--> Stage 1: Query Understanding & Cleaning
   |     |      `-- QueryClassifier (Domain, Archetype, Real-Time Check)
   |     |
   |     +--> Stage 2: Multi-Source Retrieval (Concurrent Async HTTPX)
   |     |      |-- OpenAlex (Broad Academic Works API)
   |     |      |-- arXiv (Computer Science, Physics, Mathematics XML)
   |     |      |-- PubMed & Europe PMC (Life Sciences, Medical)
   |     |      |-- Crossref (DOI / Metadata Registry)
   |     |      |-- DuckDuckGo & Wikipedia (Live Web & Encyclopedic)
   |     |      `-- DocumentFile Store (User Uploaded PDFs via PyPDF)
   |     |
   |     +--> Stage 3: Multi-Signal Scoring & HARD RELEVANCE GATE
   |     |      `-- SourceRelevanceScorer (Negative Penalty, Concept Intersection, Threshold >= 0.48)
   |     |
   |     +--> Stage 4: Evidence Extraction & Grounding
   |     |      `-- EvidenceService (Quote Extraction, Why-Relevant Rationale, Confidence)
   |     |
   |     +--> Stage 5: Contradiction & Consensus Analysis
   |     |      `-- ContradictionService (Direct Conflicts, Methodological Divergences)
   |     |
   |     +--> Stage 6: Multi-Archetype LLM Synthesis (Gemini 2.5)
   |     |      `-- Structured Markdown, Interactive Citations [X], Key Takeaways
   |     |
   |     `--> Stage 7: IEEE DOCX Manuscript Generation
   |            `-- IEEEDocumentGenerator (Title Block, 2-Column Body, References)
   |
   `-- [ Persistence & Download Delivery ]
         |-- Store Task, Claims, Evidence, Sources in PostgreSQL
         `-- Stream ResearchResult JSON & Binary .docx File to User
    """, language="text")

    # CHAPTER 7
    doc.add_heading("7. Full Project Directory & File-by-File Reference", level=1)
    doc.add_paragraph("Actual repository layout and source code manifest:")
    
    add_code_block_docx(doc, """
AI Research assistant/
├── backend/                              # FastAPI Python Backend Service
│   ├── app/
│   │   ├── core/                         # Core logging, security & exception handling
│   │   │   ├── exceptions.py             # Custom AppException, NotFound, Validation errors
│   │   │   ├── logging.py                # Structured stdout/file logging
│   │   │   └── security.py               # JWT encoding/decoding, bcrypt password hashing
│   │   ├── db/                           # Database access & SQLAlchemy models
│   │   │   ├── database.py               # Database engine, sessionmaker, get_db dependency
│   │   │   ├── init_db.py                # Table creation & non-destructive schema migrations
│   │   │   └── models.py                 # User, Project, Conversation, Task, Claim, Evidence, Doc
│   │   ├── routers/                      # REST API Endpoints
│   │   │   ├── auth.py                   # /api/v1/auth (login, register, oauth_sync, me)
│   │   │   ├── conversations.py          # /api/v1/conversations (CRUD & messages)
│   │   │   ├── documents.py              # /api/v1/documents (PDF upload & chunking)
│   │   │   ├── health.py                 # /api/v1/health (Health check & DB ping)
│   │   │   ├── projects.py               # /api/v1/projects (Project workspace management)
│   │   │   ├── research.py               # /api/v1/research (run, tasks, evidence, doc download)
│   │   │   └── sources.py                # /api/v1/sources (Search & source inspection)
│   │   ├── schemas/                      # Pydantic v2 Request/Response validation schemas
│   │   │   ├── auth.py, conversation.py, document.py, project.py, research.py
│   │   ├── services/                     # Core Business Logic & AI Engines
│   │   │   ├── chunking_service.py       # Sliding window document text chunker
│   │   │   ├── contradiction_service.py  # Conflict & consensus detection
│   │   │   ├── data_analysis_service.py  # Deterministic mathematical & statistical analysis
│   │   │   ├── evidence_service.py       # Quote extraction & why-relevant mapping
│   │   │   ├── query_classifier.py       # Domain identification & formal title generator
│   │   │   ├── relevance_service.py      # SourceRelevanceScorer multi-factor hard gate
│   │   │   ├── research_engine.py        # Master 7-stage research pipeline orchestrator
│   │   │   ├── document_generation/
│   │   │   │   ├── ieee_docx.py          # IEEE two-column Word manuscript generator
│   │   │   │   └── validator.py          # DOCX integrity and compliance validator
│   │   │   └── providers/
│   │   │       ├── base.py               # Abstract search provider interface
│   │   │       ├── gemini_llm.py         # Google Gemini 2.5 REST/SDK client
│   │   │       └── search.py             # MultiSearchAggregator (OpenAlex, arXiv, PubMed, DDG)
│   │   ├── config.py                     # pydantic-settings environment configuration
│   │   └── main.py                       # FastAPI application entry point & CORS
│   ├── tests/                            # Pytest Automated Test Suites
│   │   ├── test_docx_generation.py       # Word document compliance tests
│   │   ├── test_production_research_features.py # Real-time, roadmap & math tests
│   │   ├── test_research.py              # End-to-end research API tests
│   │   └── test_research_evaluation.py   # Relevance gate & 20-query evaluation benchmarks
│   ├── Dockerfile                        # Backend production container configuration
│   ├── Procfile                          # Process definition (uvicorn app.main:app)
│   └── requirements.txt                  # Python dependencies manifest
│
├── src/                                  # Next.js 16 Frontend Application
│   ├── app/                              # Next.js App Router
│   │   ├── api/
│   │   │   ├── auth/[...nextauth]/       # NextAuth API route handlers
│   │   │   └── v1/[...path]/route.ts     # Dynamic Reverse Proxy to FastAPI Backend
│   │   ├── chat/page.tsx                 # Dedicated chat & inquiry view
│   │   ├── login/page.tsx                # Authentication & OAuth sign-in page
│   │   ├── layout.tsx                    # Root layout, ThemeProvider, SessionProvider
│   │   ├── page.tsx                      # Root redirect & home route
│   │   └── globals.css                   # Tailwind v4 styles, custom scrollbars & glassmorphism
│   ├── auth.ts                           # NextAuth v5 configuration (Google, GitHub)
│   ├── components/                       # React 19 UI Components
│   │   ├── ResearchWorkspace.tsx         # Master 3-panel research workspace orchestrator
│   │   ├── workspace/
│   │   │   ├── ReportViewer.tsx          # Answer-first center panel with interactive citations
│   │   │   ├── SourcesPanel.tsx          # Sources list with Query Relevance badges & modal
│   │   │   ├── EvidencePanel.tsx         # Evidence Matrix (Claims, Excerpts, Confidence)
│   │   │   ├── ResearchComposer.tsx      # Multi-mode research prompt input
│   │   │   ├── ResearchProgress.tsx      # Real-time multi-stage pipeline status progress
│   │   │   ├── Sidebar.tsx               # Workspace navigation & history drawer
│   │   │   └── TopBar.tsx                # Action header, document export, theme toggle
│   │   ├── chat/MarkdownRenderer.tsx     # Enhanced Markdown & interactive [X] citation parser
│   │   └── ui/                           # Atoms & Primitives (Button, Modal, Tabs, Badge)
│   └── lib/                              # Frontend client utilities & API fetchers
├── package.json                          # Node dependencies & npm scripts
├── next.config.ts                        # Next.js configuration
├── tsconfig.json                         # TypeScript strict compiler options
└── railway.json                          # Railway cloud deployment blueprint
    """, language="text")

    # CHAPTER 8
    doc.add_heading("8. Frontend Architecture: Next.js 16, React 19 & Mobile UI", level=1)
    doc.add_paragraph(
        "The frontend is built on Next.js 16.3.1 (App Router) and React 19.2.8. "
        "The workspace interface provides a seamless desktop 3-panel layout and a custom mobile-first tabbed interaction model:"
    )
    p = doc.add_paragraph()
    p.add_run("1. Master Workspace (ResearchWorkspace.tsx): ").bold = True
    p.add_run("Maintains the global research state (currentTask, sources, evidence, claims, contradictions, activeTab). Handles real-time polling or streaming updates.\n")
    p.add_run("2. Center Panel (ReportViewer.tsx): ").bold = True
    p.add_run("Presents the Answer-First experience: Direct Answer Summary, Key Takeaways Bullet Points, and Deep-Dive Sections. Parses interactive citations [X] that immediately highlight the corresponding evidence card in the Evidence Matrix.\n")
    p.add_run("3. Left Panel (SourcesPanel.tsx): ").bold = True
    p.add_run("Lists verified academic and web sources. Displays Source Validity (Verified) alongside Query Relevance (High / Moderate / Limited) and provides a full detail modal with abstract and DOI.\n")
    p.add_run("4. Right Panel (EvidencePanel.tsx): ").bold = True
    p.add_run("Renders the formal Evidence Matrix. Displays atomic claims, verbatim quotes, why-relevant explanations, and qualitative confidence levels.\n")
    p.add_run("5. Mobile-First Optimization: ").bold = True
    p.add_run("On mobile viewports (<1024px), the layout switches to a segmented bottom tab navigator (Sources | Answer | Evidence) with 44px+ touch targets and zero horizontal scroll.")

    # CHAPTER 9
    doc.add_heading("9. Backend Architecture: FastAPI, Middleware & Endpoints", level=1)
    doc.add_paragraph("FastAPI application architecture, lifecycle, and API endpoint reference:")
    
    api_endpoints = [
        ("POST", "/api/v1/auth/register", "Public", "Registers new user with email and password"),
        ("POST", "/api/v1/auth/login", "Public", "Authenticates user and returns JWT bearer token"),
        ("POST", "/api/v1/auth/oauth_sync", "Public", "Synchronizes OAuth profile from NextAuth to database"),
        ("GET", "/api/v1/auth/me", "Bearer Auth", "Retrieves current authenticated user profile"),
        ("GET", "/api/v1/health", "Public", "System health check and database connectivity ping"),
        ("POST", "/api/v1/research/run", "Optional / Auth", "Executes full 7-stage research inquiry pipeline"),
        ("GET", "/api/v1/research/tasks/{id}", "Optional / Auth", "Fetches complete research task status and artifacts"),
        ("GET", "/api/v1/research/tasks/{id}/evidence", "Optional / Auth", "Returns structured evidence matrix for task"),
        ("GET", "/api/v1/research/tasks/{id}/contradictions", "Optional / Auth", "Returns detected claim contradictions"),
        ("GET", "/api/v1/research/tasks/{id}/document/download", "Optional / Auth", "Streams binary IEEE Word (.docx) manuscript"),
        ("POST", "/api/v1/research/tasks/{id}/document/regenerate", "Optional / Auth", "Regenerates DOCX with updated styling/title"),
        ("GET", "/api/v1/research/history", "Optional / Auth", "Returns historical research tasks for current user"),
        ("GET", "/api/v1/conversations", "Bearer Auth", "Lists all research conversations for user"),
        ("POST", "/api/v1/conversations", "Bearer Auth", "Creates new research conversation thread"),
        ("GET", "/api/v1/conversations/{id}", "Bearer Auth", "Retrieves conversation thread messages and history"),
        ("POST", "/api/v1/conversations/{id}/messages", "Bearer Auth", "Appends message to conversation thread"),
        ("GET", "/api/v1/projects", "Bearer Auth", "Lists all research project workspaces for user"),
        ("POST", "/api/v1/projects", "Bearer Auth", "Creates new research project workspace"),
        ("POST", "/api/v1/documents/upload", "Optional / Auth", "Uploads and parses PDF research papers")
    ]
    
    t_api = doc.add_table(rows=len(api_endpoints) + 1, cols=4)
    t_api.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, col in enumerate(["Method", "Endpoint Route", "Auth Type", "Operational Description"]):
        cell = t_api.rows[0].cells[i]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, 50, 50, 70, 70)
        p = cell.paragraphs[0]
        r = p.add_run(col)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    for r_idx, row_values in enumerate(api_endpoints):
        row = t_api.rows[r_idx + 1]
        for c_idx, val in enumerate(row_values):
            cell = row.cells[c_idx]
            set_cell_background(cell, "FFFFFF" if r_idx % 2 == 0 else "F8FAFC")
            set_cell_margins(cell, 35, 35, 50, 50)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8.0)
            if c_idx == 0:
                r.bold = True
                if val == "POST":
                    r.font.color.rgb = RGBColor(16, 185, 129)
                elif val == "GET":
                    r.font.color.rgb = RGBColor(59, 130, 246)

    # CHAPTER 10
    doc.add_heading("10. Database Architecture: PostgreSQL, SQLAlchemy 2.0 & Schema", level=1)
    doc.add_paragraph("Database entity relationships and table schema specifications:")
    
    db_tables = [
        ("users", "User accounts, hashed passwords, OAuth IDs, profile pictures"),
        ("conversations", "Research chat threads linked to user, tracking message sessions"),
        ("messages", "Individual chat messages (user/assistant/system)"),
        ("projects", "Top-level research project workspaces"),
        ("research_questions", "Target research hypotheses and questions within a project"),
        ("document_files", "Uploaded PDF papers with extracted text and storage paths"),
        ("document_chunks", "Sliding window text chunks with token counts and embeddings"),
        ("sources", "Discovered academic and web sources with reliability scores and metadata"),
        ("research_tasks", "Master research execution records storing reports, scores, and tokens"),
        ("claims", "Extracted atomic factual statements linked to research task"),
        ("evidence_items", "Verbatim quote excerpts, context, and relevance scores supporting claims"),
        ("contradictions", "Detected conflicting assertions and methodological divergences"),
        ("generated_documents", "Compiled IEEE DOCX files with SHA-256 integrity hashes")
    ]
    for tbl, desc in db_tables:
        p = doc.add_paragraph()
        p.add_run(f"• {tbl}: ").bold = True
        p.add_run(desc)

    # CHAPTER 11
    doc.add_heading("11. Authentication Architecture: NextAuth v5, Google OAuth & User Scoping", level=1)
    doc.add_paragraph(
        "Authentication uses a hybrid architecture combining NextAuth v5 beta on the frontend with JWT token verification and OAuth synchronization on the backend:"
    )
    add_callout_docx(
        doc,
        "Google & GitHub OAuth Workflow:\n"
        "1. User clicks 'Sign in with Google' on /login.\n"
        "2. NextAuth redirects to Google OAuth consent screen.\n"
        "3. Google returns OAuth token to NextAuth callback handler.\n"
        "4. Frontend automatically calls /api/v1/auth/oauth_sync to create or link the user in the PostgreSQL database.\n"
        "5. Backend returns a secure JWT bearer token stored in session cookies.",
        title="OAUTH & USER SYNCHRONIZATION FLOW"
    )

    # CHAPTER 12
    doc.add_heading("12. AI & LLM Engine: Google Gemini 2.5, Prompts & Grounding", level=1)
    doc.add_paragraph(
        "NexusResearch utilizes Google Gemini 2.5 Flash / Pro via direct HTTP REST API calls and SDK clients. "
        "The LLM is strictly constrained to synthesize evidence from retrieved candidate sources:"
    )
    p = doc.add_paragraph()
    p.add_run("• Query Classification Prompt: ").bold = True
    p.add_run("Extracts primary domain, intent archetype, real-time necessity, and domain sub-queries.\n")
    p.add_run("• Evidence Grounding Prompt: ").bold = True
    p.add_run("Extracts atomic claims and maps each claim to exact source indices [X] with verbatim quote excerpts.\n")
    p.add_run("• Synthesis Prompt: ").bold = True
    p.add_run("Generates Answer Summary, Key Takeaways, Deep-Dive Academic Analysis, and Qualitative Confidence metrics. Never produces fabricated numerical percentages.")

    # CHAPTER 13
    doc.add_heading("13. Deep Research Pipeline: Multi-Stage Orchestrator", level=1)
    doc.add_paragraph("The 7-stage execution pipeline implemented in ResearchEngine (backend/app/services/research_engine.py):")
    
    stages = [
        ("Stage 1: Intent & Domain Classification", "Classifies query into domain (e.g. agriculture_farming) and intent (e.g. academic_scientific). Generates high-precision sub-queries."),
        ("Stage 2: Concurrent Multi-Source Retrieval", "Executes parallel async queries across OpenAlex, arXiv, PubMed, Crossref, DuckDuckGo, and Wikipedia."),
        ("Stage 3: Multi-Signal Scoring & Hard Gate", "Applies SourceRelevanceScorer. Rejects unrelated papers (score < 0.48) and eliminates negative domain mismatches."),
        ("Stage 4: Evidence Extraction & Mapping", "Extracts verbatim quote evidence, attaches 'why_relevant' rationales, and assigns qualitative confidence tiers."),
        ("Stage 5: Contradiction Detection", "Analyzes claims across diverse papers to detect conflicting findings or methodological variations."),
        ("Stage 6: Multi-Archetype Synthesis", "Synthesizes final answer in structured markdown with interactive [X] citations grounded in the evidence matrix."),
        ("Stage 7: IEEE DOCX Document Generation", "Compiles publication-grade IEEE manuscript and stores generated document record.")
    ]
    for stg, desc in stages:
        p = doc.add_paragraph()
        p.add_run(f"• {stg}: ").bold = True
        p.add_run(desc)

    # CHAPTER 14
    doc.add_heading("14. External Research & Academic APIs", level=1)
    doc.add_paragraph("External search providers utilized and integrated in MultiSearchAggregator:")
    
    apis = [
        ("OpenAlex API", "https://api.openalex.org/works", "Broad academic metadata discovery, citation counts, open access links"),
        ("arXiv API", "http://export.arxiv.org/api/query", "Pre-print literature in Computer Science, Machine Learning, Physics, Mathematics"),
        ("PubMed E-Utilities", "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/", "Biomedical, clinical, genomic, and life science peer-reviewed papers"),
        ("Europe PMC API", "https://www.ebi.ac.uk/europepmc/webservices/rest/", "European life science and biomedical research repository"),
        ("Crossref API", "https://api.crossref.org/works", "Global DOI registry, publisher metadata, author records"),
        ("DuckDuckGo HTML", "https://html.duckduckgo.com/html/", "Live web indexing, breaking news, technical documentation"),
        ("Wikipedia API", "https://en.wikipedia.org/w/api.php", "Foundational encyclopedic definitions and historical background")
    ]
    for name, base, desc in apis:
        p = doc.add_paragraph()
        p.add_run(f"• {name} ({base}): ").bold = True
        p.add_run(desc)

    # CHAPTER 15
    doc.add_heading("15. Source Retrieval, Multi-Factor Relevance Scoring & Hard Gate", level=1)
    doc.add_paragraph(
        "Source relevance is evaluated by SourceRelevanceScorer (backend/app/services/relevance_service.py) using a multi-signal composite algorithm:"
    )
    add_code_block_docx(doc, """
Scoring Formula:
Composite Score = (Title_Score * 0.35) + (Abstract_Score * 0.25) + (Concept_Intersection * 0.30) + (Exact_Phrase * 0.10) - (Negative_Domain_Penalty)

Hard Negative Disqualification:
If query domain is Computing / AI / Agriculture and candidate paper contains medical / veterinary terms 
(e.g., 'doxycycline', 'pharmacokinetics', 'catfish', 'periodontitis', 'swine'):
Score is capped at <= 0.10 and classified as 'IRRELEVANT'.

Hard Relevance Gate Threshold:
If Composite Score >= 0.70  --> Tier: HIGH (Passed Gate)
If Composite Score >= 0.48  --> Tier: MODERATE (Passed Gate)
If Composite Score < 0.48   --> Tier: IRRELEVANT (Hard Rejected, Discarded from Evidence & Synthesis)
    """, language="text")

    # CHAPTER 16
    doc.add_heading("16. Evidence Matrix & Grounding Verification System", level=1)
    doc.add_paragraph(
        "EvidenceService extracts atomic factual claims from sources passing the hard relevance gate. "
        "Every evidence item contains:"
    )
    p = doc.add_paragraph()
    p.add_run("1. Citation Index [X]: ").bold = True
    p.add_run("Direct pointer linking the claim to the specific verified source in the Sources list.\n")
    p.add_run("2. Claim Text: ").bold = True
    p.add_run("Clear, atomic statement asserting a factual research finding.\n")
    p.add_run("3. Exact Verbatim Quote: ").bold = True
    p.add_run("Unmodified text excerpt extracted directly from the paper's title or abstract.\n")
    p.add_run("4. Why Relevant: ").bold = True
    p.add_run("Concrete rationale explaining how this finding answers the user's specific inquiry.\n")
    p.add_run("5. Qualitative Confidence: ").bold = True
    p.add_run("Assigned as 'High', 'Moderate', or 'Limited' based on evidence support (never fake percentages).")

    # CHAPTER 17
    doc.add_heading("17. Answer Generation & Multi-Archetype Synthesis Engine", level=1)
    doc.add_paragraph(
        "The synthesis engine formats research results into an answer-first, multi-tiered report structure:"
    )
    p = doc.add_paragraph()
    p.add_run("• Tier 1: Direct Answer Summary (1-2 clear paragraphs answering the core question immediately).\n")
    p.add_run("• Tier 2: Key Takeaways (3-5 structured bullet points with inline [X] citations).\n")
    p.add_run("• Tier 3: In-Depth Analysis (Detailed thematic sections, technical trade-offs, and empirical findings).\n")
    p.add_run("• Tier 4: Methodological & Contradiction Notes (Divergences or conflicting results across literature).")

    # CHAPTER 18
    doc.add_heading("18. Document Generation Architecture: IEEE Two-Column DOCX", level=1)
    doc.add_paragraph(
        "IEEEDocumentGenerator (backend/app/services/document_generation/ieee_docx.py) compiles publication-ready manuscripts adhering strictly to IEEE specifications:"
    )
    p = doc.add_paragraph()
    p.add_run("• Document Geometry: Letter paper (8.5 x 11.0 in), 0.75 in outer margins, 0.5 in gutter.\n")
    p.add_run("• Typography: Times New Roman, Title at 24pt Bold Centered, Section Headings at 10pt Small Caps Bold, Body Text at 10pt Regular.\n")
    p.add_run("• Two-Column Section: Body text flows through a continuous two-column section layout.\n")
    p.add_run("• Formal Research Title: Automatically converts casual user prompts into formal academic titles.\n")
    p.add_run("• References Section: Formatted IEEE bibliographic citations numbered [1], [2], [3].")

    # CHAPTER 19
    doc.add_heading("19. File Uploads & Local Document Processing Pipeline", level=1)
    doc.add_paragraph(
        "Users can upload local research papers (.pdf) via /api/v1/documents/upload. "
        "The backend uses PyPDF to extract text streams, performs sliding-window chunking (chunking_service.py), "
        "stores chunks in document_chunks, and integrates user document chunks directly into the multi-source retrieval pool."
    )

    # CHAPTER 20
    doc.add_heading("20. User History, Workspaces & Project Isolation", level=1)
    doc.add_paragraph(
        "All database records (conversations, projects, research tasks, generated documents) are strictly foreign-keyed to user_id. "
        "SQLAlchemy queries in all router endpoints enforce 'where user_id == current_user.id', ensuring complete data isolation between research accounts."
    )

    # CHAPTER 21
    doc.add_heading("21. Complete Environment Variables & Secrets Reference", level=1)
    doc.add_paragraph("Required environment variables across frontend and backend services:")
    
    env_vars = [
        ("DATABASE_URL", "Backend", "Yes", "postgresql://user:pass@host:5432/dbname", "PostgreSQL database connection string"),
        ("GEMINI_API_KEY", "Backend", "Yes", "<YOUR_GOOGLE_GEMINI_API_KEY>", "Google Gemini 2.5 API key for LLM intelligence"),
        ("SECRET_KEY", "Backend", "Yes", "<YOUR_32_CHAR_JWT_SECRET>", "Secret key for JWT token encryption & signing"),
        ("CORS_ORIGINS", "Backend", "No", "http://localhost:3000,https://frontend.up.railway.app", "Allowed CORS frontend origins"),
        ("NEXTAUTH_SECRET", "Frontend", "Yes", "<YOUR_NEXTAUTH_SECRET>", "NextAuth session cookie encryption secret"),
        ("AUTH_GOOGLE_ID", "Frontend", "No", "<YOUR_GOOGLE_CLIENT_ID>", "Google OAuth client ID for user login"),
        ("AUTH_GOOGLE_SECRET", "Frontend", "No", "<YOUR_GOOGLE_CLIENT_SECRET>", "Google OAuth client secret"),
        ("AUTH_GITHUB_ID", "Frontend", "No", "<YOUR_GITHUB_CLIENT_ID>", "GitHub OAuth client ID"),
        ("AUTH_GITHUB_SECRET", "Frontend", "No", "<YOUR_GITHUB_CLIENT_SECRET>", "GitHub OAuth client secret"),
        ("BACKEND_INTERNAL_URL", "Frontend", "No", "https://backend-production-873b.up.railway.app", "Target URL for Next.js API route proxy")
    ]
    
    t_env = doc.add_table(rows=len(env_vars) + 1, cols=5)
    t_env.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, col in enumerate(["Variable Name", "Service", "Required?", "Example Format", "Purpose"]):
        cell = t_env.rows[0].cells[i]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, 50, 50, 70, 70)
        p = cell.paragraphs[0]
        r = p.add_run(col)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    for r_idx, row_values in enumerate(env_vars):
        row = t_env.rows[r_idx + 1]
        for c_idx, val in enumerate(row_values):
            cell = row.cells[c_idx]
            set_cell_background(cell, "FFFFFF" if r_idx % 2 == 0 else "F8FAFC")
            set_cell_margins(cell, 35, 35, 50, 50)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8.0)
            if c_idx == 0:
                r.bold = True

    # CHAPTER 22
    doc.add_heading("22. Local Development Setup from Scratch", level=1)
    doc.add_paragraph("Step-by-step instructions to run the entire application locally on a clean computer:")
    
    add_code_block_docx(doc, """
# 1. Clone Repository
git clone https://github.com/TejeshwarDivekar/NexusAI.git
cd NexusAI

# 2. Setup Backend Environment (Terminal 1)
cd backend
python -m venv .venv
# Windows PowerShell:
.venv\\Scripts\\Activate.ps1
# Mac/Linux:
source .venv/bin/activate

pip install -r requirements.txt
cp .env.example .env
# Edit .env and configure GEMINI_API_KEY and DATABASE_URL

# Start Backend Server
uvicorn app.main:app --host 0.0.0.0 --port 8000 --reload

# 3. Setup Frontend Environment (Terminal 2)
# In repository root:
npm install
cp .env.example .env.local
# Edit .env.local and configure NEXTAUTH_SECRET and BACKEND_INTERNAL_URL=http://localhost:8000

# Start Frontend Server
npm run dev
    """, language="bash")

    # CHAPTER 23
    doc.add_heading("23. Running in VS Code & Windows Environment Guide", level=1)
    doc.add_paragraph(
        "Windows-specific guidelines for VS Code developers:\n"
        "• Terminal Execution: Always use PowerShell inside VS Code (Terminal -> New Terminal).\n"
        "• Execution Policy: If script activation fails, run 'Set-ExecutionPolicy -Scope Process -ExecutionPolicy Bypass'.\n"
        "• Line Endings: Git autocrlf should be set to 'git config core.autocrlf true' on Windows.\n"
        "• Port Conflicts: If port 8000 or 3000 is occupied, use 'Get-Process -Id (Get-NetTCPConnection -LocalPort 8000).OwningProcess | Stop-Process'."
    )

    # CHAPTER 24
    doc.add_heading("24. First-Run Verification Checklist", level=1)
    doc.add_paragraph("Checklist to verify full operational status upon first execution:")
    p = doc.add_paragraph()
    p.add_run("[ ] Repository cloned and dependencies installed (npm install, pip install -r requirements.txt)\n")
    p.add_run("[ ] Database tables created automatically on startup by init_and_upgrade_db()\n")
    p.add_run("[ ] Health check ping returning status 200 at http://localhost:8000/api/v1/health\n")
    p.add_run("[ ] Frontend loaded at http://localhost:3000\n")
    p.add_run("[ ] Test research query submitted (e.g. 'Applications of AI in Precision Agriculture')\n")
    p.add_run("[ ] Sources appearing with 'Query Relevance: High / Moderate' badges\n")
    p.add_run("[ ] Evidence Matrix populating claims and exact quotes\n")
    p.add_run("[ ] Word document download streaming valid IEEE formatted .docx file")

    # CHAPTER 25
    doc.add_heading("25. Testing Strategy & Automated Benchmark Suites", level=1)
    doc.add_paragraph("Automated test suites located in backend/tests/:")
    p = doc.add_paragraph()
    p.add_run("• test_docx_generation.py: Verifies IEEE Word document styling, title block, and tables.\n")
    p.add_run("• test_production_research_features.py: Verifies query classification, roadmaps, and statistical math analysis.\n")
    p.add_run("• test_research_evaluation.py: Validates the Hard Relevance Gate and runs a 20-query evaluation benchmark suite.\n")
    p.add_run("Command to run tests: ").bold = True
    p.add_run("python -m pytest backend/tests\n")
    p.add_run("Frontend build verification: ").bold = True
    p.add_run("npm run build")

    # CHAPTER 26
    doc.add_heading("26. Logging, Observability & Debugging Guide", level=1)
    doc.add_paragraph(
        "Structured logging is implemented via app.core.logging. Every incoming request receives a unique X-Request-ID header. "
        "Process execution duration is recorded in milliseconds (X-Process-Time-MS)."
    )

    # CHAPTER 27
    doc.add_heading("27. Cloud Deployment Architecture on Railway", level=1)
    doc.add_paragraph(
        "Production deployment on Railway comprises 3 interconnected services:\n"
        "1. Backend Service: Python FastAPI container running uvicorn on port 8000 with managed secrets.\n"
        "2. Database Service: Managed PostgreSQL 16 instance with automated SSL and persistent volume storage.\n"
        "3. Frontend Service: Next.js 16 container proxying /api/v1 calls to the backend service."
    )

    # CHAPTER 28
    doc.add_heading("28. Production Configuration & Next.js API Proxying", level=1)
    doc.add_paragraph(
        "The edge proxy at src/app/api/v1/[...path]/route.ts forwards client requests to BACKEND_INTERNAL_URL. "
        "It supports streaming responses, custom binary buffers (.docx downloads), and sets maxDuration = 120s to accommodate deep academic research syntheses."
    )

    # CHAPTER 29
    doc.add_heading("29. Security, Isolation & Compliance Measures", level=1)
    doc.add_paragraph(
        "Security safeguards implemented in production:\n"
        "• Secret Masking: All API keys and JWT secrets stored in environment variables, never committed.\n"
        "• SQL Injection Prevention: 100% parameterized queries via SQLAlchemy ORM.\n"
        "• XSS Sanitization: React JSX auto-escaping and sanitized Markdown rendering in MarkdownRenderer.tsx.\n"
        "• User Isolation: Strict foreign-key constraints on user_id across all data tables."
    )

    # CHAPTER 30
    doc.add_heading("30. Manual Step-by-Step Rebuild Guide (15 Phases)", level=1)
    doc.add_paragraph("A comprehensive 15-phase manual implementation plan to rebuild NexusResearch from scratch without AI assistance:")
    
    phases = [
        ("Phase 1: Environment & Toolchain Initialization", "Install Node.js 20, Python 3.10, Git, PostgreSQL. Initialize Next.js 16 project and FastAPI backend skeleton."),
        ("Phase 2: Database Schema & ORM Setup", "Configure SQLAlchemy database.py and create models.py (User, Conversation, ResearchTask, Claim, EvidenceItem, Source)."),
        ("Phase 3: Core Backend Middleware & Security", "Implement JWT security, bcrypt password hashing, CORS middleware, and custom AppException handler."),
        ("Phase 4: Authentication Endpoints & NextAuth", "Implement /api/v1/auth routes and configure NextAuth v5 in src/auth.ts with Google/GitHub OAuth providers."),
        ("Phase 5: Search Providers & External APIs", "Create MultiSearchAggregator in backend/app/services/providers/search.py integrating OpenAlex, arXiv, PubMed, DDG."),
        ("Phase 6: SourceRelevanceScorer & Hard Gating", "Build relevance_service.py with multi-factor scoring, negative domain penalties, and hard rejection threshold (0.48)."),
        ("Phase 7: QueryClassifier & Sub-Query Expander", "Build query_classifier.py to classify intent archetypes, domains, and generate domain-aware sub-queries."),
        ("Phase 8: Evidence & Contradiction Engines", "Implement evidence_service.py for verbatim quote extraction and contradiction_service.py for conflict detection."),
        ("Phase 9: Gemini 2.5 LLM Integration", "Implement gemini_llm.py to execute structured query classification and answer-first synthesis."),
        ("Phase 10: Master ResearchEngine Orchestrator", "Connect all 7 stages in research_engine.py into a single deterministic async pipeline."),
        ("Phase 11: IEEE DOCX Manuscript Builder", "Build ieee_docx.py using python-docx to generate formal two-column research manuscripts."),
        ("Phase 12: Next.js Dynamic Reverse Proxy", "Create src/app/api/v1/[...path]/route.ts to seamlessly forward frontend requests to FastAPI backend."),
        ("Phase 13: 3-Panel Workspace UI", "Build ResearchWorkspace.tsx, ReportViewer.tsx, SourcesPanel.tsx, and EvidencePanel.tsx."),
        ("Phase 14: Mobile-First Responsive Optimization", "Implement segmented mobile bottom navigation, touch targets >= 44px, and modal inspection views."),
        ("Phase 15: Cloud Deployment & CI/CD", "Configure railway.json, Dockerfile, environment variables, and deploy to Railway.")
    ]
    for ph, desc in phases:
        p = doc.add_paragraph()
        p.add_run(f"• {ph}: ").bold = True
        p.add_run(desc)

    # CHAPTER 31
    doc.add_heading("31. Exact Recommended Rebuild Order", level=1)
    doc.add_paragraph(
        "Sequential build order: 1. Setup git & repo -> 2. Backend database & models -> 3. Backend auth -> "
        "4. Search providers -> 5. Relevance scorer -> 6. Research engine -> 7. Document builder -> "
        "8. Next.js frontend proxy -> 9. UI Workspace & Panels -> 10. Mobile responsive layout -> 11. Cloud deployment."
    )

    # CHAPTER 32
    doc.add_heading("32. Beginner Guide to Core Technical Concepts", level=1)
    doc.add_paragraph("Essential architectural concepts explained simply:")
    p = doc.add_paragraph()
    p.add_run("• REST API: An architectural standard for exchanging JSON data over HTTP methods (GET, POST, PATCH, DELETE).\n")
    p.add_run("• OAuth 2.0: A secure delegated authorization framework allowing users to log in with Google without sharing passwords.\n")
    p.add_run("• ORM (Object-Relational Mapping): Translates database tables into Python classes (SQLAlchemy).\n")
    p.add_run("• Evidence Grounding: Ensuring every sentence in an AI report is directly anchored to an authentic quote from a peer-reviewed source.\n")
    p.add_run("• Hard Gate: A programmatic filter that automatically discards sources failing quality or relevance criteria.")

    # CHAPTER 33
    doc.add_heading("33. Comprehensive Troubleshooting Guide", level=1)
    
    trouble_data = [
        ("CORS error in browser", "Backend CORS_ORIGINS missing frontend URL", "Add frontend URL to CORS_ORIGINS in backend .env"),
        ("Database connection refused", "PostgreSQL service stopped or bad URI", "Verify DATABASE_URL and ensure PostgreSQL is listening on port 5432"),
        ("Search returns 0 sources", "API rate limits or network timeout", "Check internet connection; search aggregator will fallback to arXiv/DDG"),
        ("Word download fails", "Missing python-docx or invalid path", "Verify python-docx>=1.2.0 installed and generated_docs/ directory is writable"),
        ("Next.js Proxy 504 Timeout", "Backend research taking > 60s", "Ensure maxDuration = 120 is set in src/app/api/v1/[...path]/route.ts"),
        ("Mobile horizontal scroll", "Unconstrained table or code block width", "Apply 'overflow-x-auto' and 'max-w-full' to markdown containers")
    ]
    t_tr = doc.add_table(rows=len(trouble_data) + 1, cols=3)
    t_tr.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, col in enumerate(["Observed Symptom", "Probable Root Cause", "Step-by-Step Resolution"]):
        cell = t_tr.rows[0].cells[i]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, 50, 50, 70, 70)
        p = cell.paragraphs[0]
        r = p.add_run(col)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    for r_idx, row_values in enumerate(trouble_data):
        row = t_tr.rows[r_idx + 1]
        for c_idx, val in enumerate(row_values):
            cell = row.cells[c_idx]
            set_cell_background(cell, "FFFFFF" if r_idx % 2 == 0 else "F8FAFC")
            set_cell_margins(cell, 35, 35, 50, 50)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8.0)
            if c_idx == 0:
                r.bold = True

    # CHAPTER 34 - 38
    doc.add_heading("34. Maintenance, Schema Upgrades & Model Operations", level=1)
    doc.add_paragraph(
        "Database migrations are handled non-destructively by init_and_upgrade_db() in backend/app/db/init_db.py. "
        "To upgrade dependencies, run 'npm update' in frontend and 'pip install --upgrade -r requirements.txt' in backend."
    )
    
    doc.add_heading("35. Current Known Limitations & Trade-offs", level=1)
    doc.add_paragraph(
        "1. Academic API Rate Limits: Unauthenticated OpenAlex / Crossref queries are throttled under high load.\n"
        "2. LLM Latency: Multi-stage synthesis takes 8-15 seconds for deep academic inquiries.\n"
        "3. Synchronous Execution: Very long queries rely on HTTP polling; background Celery/Redis queue is recommended for enterprise scale."
    )
    
    doc.add_heading("36. Recommended Future Improvements", level=1)
    doc.add_paragraph(
        "• Implement Redis / Celery for distributed background task execution.\n"
        "• Add vector semantic caching for repeated academic queries.\n"
        "• Integrate real-time WebSocket communication for token-by-token streaming synthesis.\n"
        "• Expand search providers to include Semantic Scholar and IEEE Xplore APIs."
    )
    
    doc.add_heading("37. Technical Glossary of Terms", level=1)
    doc.add_paragraph(
        "• Grounding: Verifying that synthetic text is mathematically supported by source citations.\n"
        "• Hard Gate: Programmatic rejection threshold ensuring zero low-quality data enters the evidence pool.\n"
        "• IEEE Format: Two-column standard layout published by the Institute of Electrical and Electronics Engineers."
    )
    
    doc.add_heading("38. Command Reference & Final Reproduction Checklist", level=1)
    add_code_block_docx(doc, """
Quick Command Reference:
Frontend Dev:           npm run dev
Frontend Build:         npm run build
Frontend Lint:          npm run lint
Backend Dev:            uvicorn app.main:app --reload --port 8000
Backend Test:           python -m pytest backend/tests
Git Status:             git status
Git Push:               git push -u origin main
    """, language="bash")
    
    p_final = doc.add_paragraph()
    p_final.add_run("Final Reproduction Verification: ").bold = True
    p_final.add_run("Following Chapters 22 through 30 allows any software engineer to manually rebuild, configure, execute, and deploy NexusResearch from scratch with 100% fidelity.")

    # Save DOCX
    doc.save(docx_path)

def build_pdf(pdf_path):
    # Professional ReportLab PDF generation matching the complete technical manual
    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        leftMargin=0.75*inch,
        rightMargin=0.75*inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    style_title = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#0F172A'),
        alignment=1, # Center
        spaceAfter=10
    )
    
    style_subtitle = ParagraphStyle(
        'DocSubTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=11,
        leading=15,
        textColor=colors.HexColor('#475569'),
        alignment=1,
        spaceAfter=20
    )
    
    style_badge = ParagraphStyle(
        'DocBadge',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#0066CC'),
        alignment=1,
        spaceAfter=12
    )
    
    style_h1 = ParagraphStyle(
        'H1',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=17,
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=14,
        spaceAfter=6,
        keepWithNext=True
    )
    
    style_body = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=colors.HexColor('#1E293B'),
        spaceAfter=6
    )
    
    style_callout = ParagraphStyle(
        'Callout',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=13,
        textColor=colors.HexColor('#0F172A')
    )
    
    style_code = ParagraphStyle(
        'CodeStyle',
        parent=styles['Code'],
        fontName='Courier',
        fontSize=7.5,
        leading=10.5,
        textColor=colors.HexColor('#E2E8F0')
    )

    story = []
    
    # Title Page
    story.append(Spacer(1, 30))
    story.append(Paragraph("PRODUCTION SYSTEM ARCHITECTURE & ENGINEERING MANUAL", style_badge))
    story.append(Paragraph("NexusResearch — Complete Technical Architecture,<br/>Development & Reconstruction Guide", style_title))
    story.append(Paragraph("How to Understand, Run, Rebuild, Maintain and Deploy the Complete Real-Time AI Research Platform from Scratch Without AI", style_subtitle))
    story.append(Spacer(1, 15))
    
    meta_table_data = [
        [Paragraph("<b>Author & System Architect:</b>", style_body), Paragraph("Tejeshwar Divekar", style_body)],
        [Paragraph("<b>Project Repository:</b>", style_body), Paragraph("https://github.com/TejeshwarDivekar/NexusAI.git", style_body)],
        [Paragraph("<b>Live Production Frontend:</b>", style_body), Paragraph("https://frontend-production-2df5.up.railway.app", style_body)],
        [Paragraph("<b>Live Production API:</b>", style_body), Paragraph("https://backend-production-873b.up.railway.app/api/v1/health", style_body)],
        [Paragraph("<b>Architecture Version:</b>", style_body), Paragraph("1.0.0-Production (Next.js 16.3.1 + FastAPI 1.0.0 + Gemini 2.5)", style_body)],
        [Paragraph("<b>Release Date:</b>", style_body), Paragraph(datetime.datetime.now().strftime("%B %d, %Y"), style_body)]
    ]
    t_meta = Table(meta_table_data, colWidths=[2.2*inch, 4.5*inch])
    t_meta.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#F8FAFC')),
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#CBD5E1')),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('LEFTPADDING', (0,0), (-1,-1), 8),
        ('RIGHTPADDING', (0,0), (-1,-1), 8),
    ]))
    story.append(t_meta)
    story.append(PageBreak())
    
    # Table of Contents
    story.append(Paragraph("Table of Contents", style_h1))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#CBD5E1'), spaceAfter=10))
    
    toc_data = [
        [Paragraph("<b>1. Project Overview & Executive Summary</b>", style_body), Paragraph("Page 3", style_body)],
        [Paragraph("<b>2. What NexusResearch Does: Core Capabilities</b>", style_body), Paragraph("Page 4", style_body)],
        [Paragraph("<b>3. Core Features & Functional Architecture</b>", style_body), Paragraph("Page 5", style_body)],
        [Paragraph("<b>4. Complete Technology Stack & Version Inventory</b>", style_body), Paragraph("Page 6", style_body)],
        [Paragraph("<b>5. Why Each Technology Is Used (In-Depth Rationale)</b>", style_body), Paragraph("Page 7", style_body)],
        [Paragraph("<b>6. Complete System Architecture & Communication Flow</b>", style_body), Paragraph("Page 8", style_body)],
        [Paragraph("<b>7. Full Project Directory & File-by-File Reference</b>", style_body), Paragraph("Page 9", style_body)],
        [Paragraph("<b>8. Frontend Architecture: Next.js 16, React 19 & Mobile UI</b>", style_body), Paragraph("Page 10", style_body)],
        [Paragraph("<b>9. Backend Architecture: FastAPI, Middleware & Endpoints</b>", style_body), Paragraph("Page 11", style_body)],
        [Paragraph("<b>10. Database Architecture: PostgreSQL, SQLAlchemy 2.0 & Schema</b>", style_body), Paragraph("Page 12", style_body)],
        [Paragraph("<b>11. Authentication Architecture: NextAuth v5 & Google OAuth</b>", style_body), Paragraph("Page 13", style_body)],
        [Paragraph("<b>12. AI & LLM Engine: Google Gemini 2.5, Prompts & Grounding</b>", style_body), Paragraph("Page 14", style_body)],
        [Paragraph("<b>13. Deep Research Pipeline: Multi-Stage Orchestrator</b>", style_body), Paragraph("Page 15", style_body)],
        [Paragraph("<b>14. External Research & Academic APIs</b>", style_body), Paragraph("Page 16", style_body)],
        [Paragraph("<b>15. Source Retrieval, Multi-Factor Relevance Scoring & Hard Gate</b>", style_body), Paragraph("Page 17", style_body)],
        [Paragraph("<b>16. Evidence Matrix & Grounding Verification System</b>", style_body), Paragraph("Page 18", style_body)],
        [Paragraph("<b>17. Answer Generation & Multi-Archetype Synthesis Engine</b>", style_body), Paragraph("Page 19", style_body)],
        [Paragraph("<b>18. Document Generation Architecture: IEEE Two-Column DOCX</b>", style_body), Paragraph("Page 20", style_body)],
        [Paragraph("<b>19. File Uploads & Local Document Processing Pipeline</b>", style_body), Paragraph("Page 21", style_body)],
        [Paragraph("<b>20. User History, Workspaces & Project Isolation</b>", style_body), Paragraph("Page 22", style_body)],
        [Paragraph("<b>21. Complete Environment Variables & Secrets Reference</b>", style_body), Paragraph("Page 23", style_body)],
        [Paragraph("<b>22. Local Development Setup from Scratch</b>", style_body), Paragraph("Page 24", style_body)],
        [Paragraph("<b>23. Running in VS Code & Windows Environment Guide</b>", style_body), Paragraph("Page 25", style_body)],
        [Paragraph("<b>24. First-Run Verification Checklist</b>", style_body), Paragraph("Page 26", style_body)],
        [Paragraph("<b>25. Testing Strategy & Automated Benchmark Suites</b>", style_body), Paragraph("Page 27", style_body)],
        [Paragraph("<b>26. Logging, Observability & Debugging Guide</b>", style_body), Paragraph("Page 28", style_body)],
        [Paragraph("<b>27. Cloud Deployment Architecture on Railway</b>", style_body), Paragraph("Page 29", style_body)],
        [Paragraph("<b>28. Production Configuration & Next.js API Proxying</b>", style_body), Paragraph("Page 30", style_body)],
        [Paragraph("<b>29. Security, Isolation & Compliance Measures</b>", style_body), Paragraph("Page 31", style_body)],
        [Paragraph("<b>30. Manual Step-by-Step Rebuild Guide (15 Phases)</b>", style_body), Paragraph("Page 32", style_body)],
        [Paragraph("<b>31. Exact Recommended Rebuild Order</b>", style_body), Paragraph("Page 34", style_body)],
        [Paragraph("<b>32. Beginner Guide to Core Technical Concepts</b>", style_body), Paragraph("Page 35", style_body)],
        [Paragraph("<b>33. Comprehensive Troubleshooting Guide</b>", style_body), Paragraph("Page 36", style_body)],
        [Paragraph("<b>34. Maintenance, Schema Upgrades & Model Operations</b>", style_body), Paragraph("Page 37", style_body)],
        [Paragraph("<b>35. Current Known Limitations & Trade-offs</b>", style_body), Paragraph("Page 38", style_body)],
        [Paragraph("<b>36. Recommended Future Improvements</b>", style_body), Paragraph("Page 39", style_body)],
        [Paragraph("<b>37. Technical Glossary of Terms</b>", style_body), Paragraph("Page 40", style_body)],
        [Paragraph("<b>38. Command Reference & Final Reproduction Checklist</b>", style_body), Paragraph("Page 41", style_body)],
    ]
    t_toc_pdf = Table(toc_data, colWidths=[5.8*inch, 0.9*inch])
    t_toc_pdf.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('TOPPADDING', (0,0), (-1,-1), 2),
        ('BOTTOMPADDING', (0,0), (-1,-1), 2),
        ('LINEBELOW', (0,0), (-1,-1), 0.5, colors.HexColor('#F1F5F9'))
    ]))
    story.append(t_toc_pdf)
    story.append(PageBreak())
    
    # Chapters Content
    chapters = [
        ("1. Project Overview & Executive Summary", 
         "NexusResearch is an enterprise-grade, autonomous real-time research assistant and academic intelligence workspace. "
         "Unlike conventional LLM wrappers that merely generate plausible-sounding text, NexusResearch implements an evidence-grounded, "
         "multi-stage verification pipeline. It retrieves authentic peer-reviewed literature, live web disclosures, and uploaded user documents, "
         "subjects all candidate sources to a strict hard relevance gate, maps exact quote-level evidence into a claim matrix, "
         "and synthesizes structured, publication-grade intelligence with interactive citations and professional IEEE Word (.docx) document generation.\n\n"
         "Core Mandate: Source Validity (registry authentication) is strictly separated from Query Relevance (topical alignment). "
         "Candidate papers are evaluated against a multi-factor relevance scorer with hard negative disqualification before synthesis begins."),
         
        ("2. What NexusResearch Does: Core Capabilities",
         "• Academic & Scientific Deep Dives: Executes multi-query expansion across OpenAlex, arXiv, PubMed, Europe PMC, and Crossref. Synthesizes formal state-of-the-art reports with mathematical formulations, methodologies, and benchmarking data.\n"
         "• Real-Time & Live Market Intelligence: Fetches real-time web data via DuckDuckGo and Wikipedia with live timestamp verification for breaking news, stock metrics, and technology releases.\n"
         "• Step-by-Step Roadmaps: Generates structured milestone tracks (Foundational -> Intermediate -> Advanced) with concrete prerequisites, toolsets, and project milestones.\n"
         "• Comparative Analyses: Constructs side-by-side dimensional matrices evaluating trade-offs, performance benchmarks, and cost profiles.\n"
         "• Numerical Calculations: Executes deterministic Python calculations (mean, standard deviation, growth rates) alongside scientific analysis."),
         
        ("3. Core Features & Functional Architecture",
         "1. Answer-First Research Workspace: Desktop 3-panel workspace (Sources | Main Answer | Evidence Matrix) and Purpose-Built Mobile Experience with segmented navigation and 44px+ touch targets.\n"
         "2. Multi-Signal Source Relevance Scorer: Multi-factor heuristic scoring engine with negative domain disqualification and hard gating (threshold >= 0.48).\n"
         "3. Evidence Matrix & Claim Grounding: Direct mapping from assertions to exact source excerpts, why-relevant rationales, and qualitative confidence scores.\n"
         "4. Contradiction & Consensus Engine: Identifies conflicting claims and methodological divergences across literature.\n"
         "5. IEEE Two-Column DOCX Document Builder: Generates fully compliant IEEE Word documents ready for publication.\n"
         "6. NextAuth v5 & User Isolation: Secure Google/GitHub OAuth authentication with PostgreSQL user-level project isolation.\n"
         "7. Dynamic Next.js API Proxying: Edge route proxy at /api/v1/[...path] forwarding requests with 120s timeout and binary streaming support."),
         
        ("4. Complete Technology Stack & Version Inventory",
         "• Next.js: v16.3.1 (App Router, Server Actions, Route Proxying)\n"
         "• React & React-DOM: v19.2.8 (Component tree, hooks, hydration)\n"
         "• Tailwind CSS: v4.0.0 (CSS-first token architecture, dark mode)\n"
         "• NextAuth: v5.0.0-beta.32 (Google/GitHub OAuth, JWT sessions)\n"
         "• FastAPI: >=0.110.0 (High-performance async REST API)\n"
         "• Uvicorn: >=0.28.0 (Production async ASGI server)\n"
         "• SQLAlchemy: >=2.0.28 (Database ORM & automated migration)\n"
         "• Psycopg2-Binary: >=2.9.9 (PostgreSQL high-performance driver)\n"
         "• Pydantic: >=2.6.0 (Request/Response data validation schemas)\n"
         "• HTTPX: >=0.27.0 (Async HTTP search requests)\n"
         "• Python-Docx: >=1.2.0 (Programmatic IEEE DOCX compilation)\n"
         "• Google Gemini 2.5: REST / SDK (Query classification, synthesis & grounding)"),
         
        ("5. Why Each Technology Is Used",
         "• Next.js 16 + React 19: Provides high-speed App Router server rendering, dynamic streaming capabilities, and built-in API proxy routing without UI flicker.\n"
         "• FastAPI + Python 3.10+: Native asynchronous I/O (asyncio) allows concurrent multi-source querying across 5+ academic search providers in parallel.\n"
         "• PostgreSQL + SQLAlchemy 2.0: Relational integrity ensures rigorous mapping between Users, Projects, Research Tasks, Claims, Evidence, and Sources.\n"
         "• Google Gemini 2.5: 1M+ token context window and superior JSON structured output compliance for deterministic claim-evidence mapping.\n"
         "• Python-Docx: Allows complete programmatic control over binary Word XML packaging to build IEEE standard manuscripts directly on the server."),
         
        ("6. Complete System Architecture & Communication Flow",
         "Client Browser -> Next.js 16 Edge / App Router (Frontend) -> Reverse Proxy (/api/v1/[...path]) -> FastAPI Backend -> ResearchEngine Orchestrator -> Search Providers (OpenAlex, arXiv, PubMed, Crossref, DDG) -> SourceRelevanceScorer (Hard Gate >= 0.48) -> EvidenceService -> ContradictionService -> Gemini 2.5 LLM -> IEEEDocumentGenerator (.docx) -> PostgreSQL Persistence -> Client Streaming Response."),
         
        ("7. Full Project Directory & File Manifest",
         "backend/app/main.py (FastAPI Bootstrap), backend/app/routers/ (auth.py, conversations.py, documents.py, health.py, projects.py, research.py, sources.py), backend/app/services/ (research_engine.py, relevance_service.py, query_classifier.py, evidence_service.py, contradiction_service.py, ieee_docx.py, gemini_llm.py, search.py), src/app/api/v1/[...path]/route.ts (Proxy), src/components/ResearchWorkspace.tsx (Master UI), src/components/workspace/ (ReportViewer.tsx, SourcesPanel.tsx, EvidencePanel.tsx)."),
         
        ("8. Frontend Architecture: Next.js 16 & React 19",
         "Desktop 3-Panel Workspace (SourcesPanel | ReportViewer | EvidencePanel). Interactive [X] citations highlight linked evidence cards in real-time. Mobile viewports (<1024px) utilize a segmented bottom tab navigator with 44px+ touch targets and zero horizontal scroll."),
         
        ("9. Backend Architecture: FastAPI & API Endpoints",
         "FastAPI router modules provide 24+ REST endpoints across /api/v1/auth, /api/v1/research, /api/v1/conversations, /api/v1/projects, /api/v1/documents, and /api/v1/health. All endpoints support standard HTTP methods, bearer authentication, correlation IDs, and structured exception handlers."),
         
        ("10. Database Architecture & Schema Reference",
         "PostgreSQL database tables: users, conversations, messages, projects, research_questions, document_files, document_chunks, sources, research_tasks, claims, evidence_items, contradictions, generated_documents. All tables strictly enforce cascade deletions and foreign keys."),
         
        ("11. Authentication: NextAuth v5 & Google OAuth",
         "NextAuth v5 beta manages frontend OAuth sessions. Upon login, the client calls /api/v1/auth/oauth_sync to create/link the PostgreSQL user record and issue an application JWT token."),
         
        ("12. AI & LLM Engine: Gemini 2.5 Grounding",
         "The LLM is strictly constrained via prompt engineering and JSON schema enforcement. It synthesizes findings strictly from retrieved candidate sources passing the relevance gate, ensuring zero ungrounded hallucinations."),
         
        ("13. Deep Research Pipeline: 7-Stage Orchestrator",
         "Stage 1: Intent & Domain Classification -> Stage 2: Concurrent Multi-Source Retrieval -> Stage 3: Multi-Signal Scoring & Hard Gate -> Stage 4: Evidence Extraction & Grounding -> Stage 5: Contradiction Detection -> Stage 6: Multi-Archetype Synthesis -> Stage 7: IEEE DOCX Document Generation."),
         
        ("14. External Research & Academic APIs",
         "Integrated search providers: OpenAlex API (works & citations), arXiv API (preprints in CS/Physics/Math), PubMed & Europe PMC (biomedical & life sciences), Crossref API (DOI registry), DuckDuckGo (live web news), Wikipedia API (encyclopedic summaries)."),
         
        ("15. Source Relevance Scoring & Hard Gate",
         "Composite Score = (Title_Score * 0.35) + (Abstract_Score * 0.25) + (Concept_Intersection * 0.30) + (Exact_Phrase * 0.10) - (Negative_Domain_Penalty). Sources with score >= 0.70 are tagged HIGH; >= 0.48 tagged MODERATE; < 0.48 are hard-rejected and discarded."),
         
        ("16. Evidence Matrix & Grounding System",
         "Evidence items store verbatim quotes, context, page numbers, 'why_relevant' rationales, and qualitative confidence ratings (High / Moderate / Limited). Linked directly to [X] citations in the synthesized report."),
         
        ("17. Answer Generation & Synthesis Structure",
         "Synthesizes structured multi-tiered output: 1. Direct Answer Summary, 2. Key Takeaways, 3. In-Depth Analysis, 4. Methodological Notes and Contradictions."),
         
        ("18. Document Generation: IEEE Two-Column DOCX",
         "IEEEDocumentGenerator formats formal IEEE manuscripts: 24pt bold title, author metadata block, italic abstract, continuous two-column body text, and numbered references."),
         
        ("19. File Uploads & Local Document Extraction",
         "PDF papers uploaded via /api/v1/documents/upload are parsed by PyPDF, segmented by sliding-window chunker, and added to the candidate source pool."),
         
        ("20. User History, Workspaces & Isolation",
         "Users can organize inquiries into Project workspaces and Conversation threads. All queries enforce 'where user_id == current_user.id'."),
         
        ("21. Environment Variables & Secrets Reference",
         "DATABASE_URL, GEMINI_API_KEY, SECRET_KEY, CORS_ORIGINS, NEXTAUTH_SECRET, AUTH_GOOGLE_ID, AUTH_GOOGLE_SECRET, BACKEND_INTERNAL_URL."),
         
        ("22. Local Development Setup from Scratch",
         "1. Clone repo -> 2. Backend: venv, pip install -r requirements.txt, uvicorn app.main:app --port 8000 -> 3. Frontend: npm install, npm run dev -> 4. Open http://localhost:3000."),
         
        ("23. Running in VS Code & Windows Guide",
         "Use PowerShell in VS Code. Configure execution policy if script activation is restricted. Verify ports 3000 and 8000 are open."),
         
        ("24. First-Run Verification Checklist",
         "Verify health check at /api/v1/health, test research query submission, verify source relevance badges, check evidence matrix, test Word document download."),
         
        ("25. Testing Strategy & Automated Benchmark Suites",
         "Pytest test suites in backend/tests: test_docx_generation.py, test_production_research_features.py, test_research_evaluation.py (20-query evaluation benchmark suite)."),
         
        ("26. Logging & Observability Guide",
         "Structured logging records request method, path, status, and duration in ms. Request correlation tracked via X-Request-ID."),
         
        ("27. Cloud Deployment Architecture on Railway",
         "Three production services: Frontend (Next.js), Backend (FastAPI Python), and Database (Managed PostgreSQL 16)."),
         
        ("28. Production Configuration & Next.js Proxying",
         "Dynamic reverse proxy at src/app/api/v1/[...path]/route.ts forwards requests to backend with 120s timeout and binary buffer streaming."),
         
        ("29. Security, Isolation & Compliance",
         "Parameterized SQL queries, XSS sanitization, secret redaction, and user data isolation across all database operations."),
         
        ("30. Manual Step-by-Step Rebuild Guide (15 Phases)",
         "Comprehensive 15-phase blueprint: Environment -> DB Models -> Middleware -> Auth -> Search Providers -> Relevance Gate -> Classifier -> Evidence Engine -> Gemini LLM -> Pipeline Orchestrator -> IEEE DOCX Builder -> Proxy -> UI Workspace -> Mobile Layout -> Deployment."),
         
        ("31. Recommended Rebuild Order",
         "Strict chronological build order from database foundations to API providers, relevance scorers, LLM synthesis, UI panels, and deployment."),
         
        ("32. Beginner Guide to Core Technical Concepts",
         "Explanations of REST APIs, OAuth 2.0, ORM, Evidence Grounding, and Hard Gating in simple, accessible language."),
         
        ("33. Comprehensive Troubleshooting Guide",
         "Diagnostic tables covering CORS errors, database connection failures, search provider timeouts, and Word document streaming issues."),
         
        ("34. Maintenance & Operations",
         "Procedures for non-destructive database migrations, dependency updates, and AI model version upgrades."),
         
        ("35. Current Known Limitations & Trade-offs",
         "Academic search provider rate limits, LLM synthesis latency (8-15s), and memory considerations for large PDF uploads."),
         
        ("36. Recommended Future Improvements",
         "Redis caching, Celery background worker queues, WebSocket real-time streaming, and Semantic Scholar API integration."),
         
        ("37. Technical Glossary of Terms",
         "Key terminology definitions: Grounding, Hard Gate, IEEE Format, DOI, OpenAlex, PubMed, Relevance Tier."),
         
        ("38. Command Reference & Final Checklist",
         "CLI reference for dev, build, test, and deployment. Final verification checklist confirms complete project reproducibility from zero.")
    ]
    
    for title, content in chapters:
        story.append(Paragraph(title, style_h1))
        story.append(HRFlowable(width="100%", thickness=0.75, color=colors.HexColor('#E2E8F0'), spaceAfter=8))
        
        # Split into paragraphs if contains double newlines
        paras = content.split("\n\n")
        for p_text in paras:
            story.append(Paragraph(p_text.replace("\n", "<br/>"), style_body))
        story.append(Spacer(1, 10))

    # Page number canvas
    def add_page_number(canvas_obj, doc_obj):
        page_num = canvas_obj.getPageNumber()
        canvas_obj.saveState()
        canvas_obj.setFont('Helvetica', 8)
        canvas_obj.setFillColor(colors.HexColor('#94A3B8'))
        canvas_obj.drawString(0.75*inch, 0.4*inch, "NexusResearch — Complete Technical Architecture & Reconstruction Guide")
        canvas_obj.drawRightString(8.5*inch - 0.75*inch, 0.4*inch, f"Page {page_num}")
        canvas_obj.restoreState()

    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)

if __name__ == "__main__":
    create_documents()
