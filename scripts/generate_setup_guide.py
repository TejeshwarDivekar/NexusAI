"""
NexusResearch Complete Project Setup & Reconstruction Guide Generator
Generates:
1. docs/NexusResearch_Complete_Project_Setup_Guide.docx
2. docs/NexusResearch_Complete_Project_Setup_Guide.pdf
"""

import os
import sys
import datetime
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether, HRFlowable
)
from reportlab.pdfgen import canvas

# ==============================================================================
# REPORTLAB NUMBERED CANVAS (PAGE X OF Y + HEADERS / FOOTERS)
# ==============================================================================
class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super(NumberedCanvas, self).__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_decorations(num_pages)
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)

    def draw_page_decorations(self, page_count):
        if self._pageNumber == 1:
            return  # Suppress headers and footers on cover page
            
        self.saveState()
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor("#64748B"))
        
        # Running Top Header
        self.drawString(54, 750, "NEXUSRESEARCH — Complete Project Setup, Execution & Deployment Manual")
        self.setStrokeColor(colors.HexColor("#CBD5E1"))
        self.setLineWidth(0.5)
        self.line(54, 742, 558, 742)
        
        # Running Bottom Footer
        self.line(54, 48, 558, 48)
        self.drawString(54, 34, "Confidential & Proprietary — Developed for Single-Source Reconstruction")
        page_text = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(558, 34, page_text)
        self.restoreState()


# ==============================================================================
# DOCX HELPER FUNCTIONS
# ==============================================================================
def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_callout_docx(doc, text, title="NOTE", fill_hex="F0F4F8", border_hex="0066CC"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_background(cell, fill_hex)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r_title = p.add_run(f"[{title}] ")
    r_title.bold = True
    r_title.font.color.rgb = RGBColor(0, 51, 102)
    r_title.font.size = Pt(9.5)
    r_text = p.add_run(text)
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(30, 41, 59)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_code_block_docx(doc, code_text, language=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(code_text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(15, 23, 42)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ==============================================================================
# MAIN BUILD SCRIPT
# ==============================================================================
def generate_all_guides():
    docs_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs")
    os.makedirs(docs_dir, exist_ok=True)
    
    docx_path = os.path.join(docs_dir, "NexusResearch_Complete_Project_Setup_Guide.docx")
    pdf_path = os.path.join(docs_dir, "NexusResearch_Complete_Project_Setup_Guide.pdf")
    
    print(f"Building Word (.docx) manual at: {docx_path}")
    build_docx_guide(docx_path)
    
    print(f"Building PDF (.pdf) manual at: {pdf_path}")
    build_pdf_guide(pdf_path)
    
    print("All setup guides successfully compiled!")


# ==============================================================================
# BUILD DOCX GUIDE
# ==============================================================================
def build_docx_guide(filepath):
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Style defaults
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(30, 41, 59)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)

    # --------------------------------------------------------------------------
    # COVER PAGE
    # --------------------------------------------------------------------------
    p_pre = doc.add_paragraph()
    p_pre.paragraph_format.space_before = Pt(40)
    r_pre = p_pre.add_run("NEXUSRESEARCH ENTERPRISE PLATFORM")
    r_pre.font.size = Pt(11)
    r_pre.bold = True
    r_pre.font.color.rgb = RGBColor(14, 116, 144)

    p_title = doc.add_paragraph()
    r_title = p_title.add_run("Complete Project Setup, Local Reconstruction & Deployment Guide")
    r_title.font.size = Pt(22)
    r_title.bold = True
    r_title.font.color.rgb = RGBColor(15, 23, 42)
    p_title.paragraph_format.space_after = Pt(8)

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("A Complete, Standalone Manual to Set Up, Run, Test, and Deploy NexusResearch on Any Windows Computer from Scratch Without AI Assistance")
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(71, 85, 105)
    p_sub.paragraph_format.space_after = Pt(24)

    add_callout_docx(
        doc,
        "SECURITY & INTEGRITY NOTICE:\n"
        "This document is authored directly from the NexusResearch production repository. "
        "Every command, port, environment variable, database table, API endpoint, and deployment step in this guide has been verified against active code. "
        "No real API keys, passwords, or OAuth secrets are published in this guide; placeholder formats are provided throughout.",
        title="AUTHENTIC SOURCE OF TRUTH"
    )

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(30)
    p_meta.add_run("Author: ").bold = True
    p_meta.add_run("NexusResearch Engineering & Architecture Group\n")
    p_meta.add_run("Version: ").bold = True
    p_meta.add_run("1.0.0 Production Release\n")
    p_meta.add_run("Target Environment: ").bold = True
    p_meta.add_run("Windows 10 / Windows 11 (64-bit)\n")
    p_meta.add_run("Date of Publication: ").bold = True
    p_meta.add_run(datetime.datetime.now().strftime("%B %d, %Y") + "\n")
    p_meta.add_run("Repository Source: ").bold = True
    p_meta.add_run("https://github.com/TejeshwarDivekar/NexusAI.git")

    doc.add_page_break()

    # --------------------------------------------------------------------------
    # TABLE OF CONTENTS
    # --------------------------------------------------------------------------
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        ("1. Project Overview & Capabilities", "3"),
        ("2. System Requirements & Prerequisites Table", "4"),
        ("3. Visual Studio Code Installation & Terminal Setup", "5"),
        ("4. Git Version Control Installation", "6"),
        ("5. Node.js & Python Runtime Environments", "7"),
        ("6. Obtaining the Repository (Git Clone & ZIP)", "8"),
        ("7. Verified Repository Directory Structure", "9"),
        ("8. Dependency Installation (Frontend npm & Backend pip)", "10"),
        ("9. Comprehensive Environment Variables Master Table", "11"),
        ("10. Creating .env Configuration Files", "13"),
        ("11. Database Setup: SQLite (Dev) & PostgreSQL (Prod)", "14"),
        ("12. Google Cloud OAuth & Login Configuration", "15"),
        ("13. AI LLM Engine: Google Gemini Configuration", "17"),
        ("14. Scholarly Research Registries & Search Providers", "18"),
        ("15. Starting the FastAPI Backend Server", "19"),
        ("16. Starting the Next.js Frontend Server", "20"),
        ("17. Dual-Terminal Local Execution Workflow", "21"),
        ("18. Performing First Login & Account Sync", "22"),
        ("19. Executing Real Research Inquiries", "23"),
        ("20. Testing Interactive Evidence Matrix & Restoration", "24"),
        ("21. Validating IEEE DOCX & Academic PDF Generation", "25"),
        ("22. Multi-Tenant Project Isolation & History", "26"),
        ("23. Mobile View & Responsive Layout Testing", "27"),
        ("24. Comprehensive Troubleshooting Matrix", "28"),
        ("25. How to Read & Diagnose Errors Across Layers", "30"),
        ("26. Development vs. Production Configuration", "31"),
        ("27. Cloud Deployment: Railway & Docker Compose", "32"),
        ("28. Database Backup & Disaster Recovery", "34"),
        ("29. Standard Developer Contribution Workflow", "35"),
        ("30. Git Branching & Commit Conventions", "36"),
        ("31. Printable Fresh-Laptop Setup Checklist", "37"),
        ("32. Unified Command Cheat Sheet", "38"),
        ("33. System Architecture Topology Diagram", "39"),
        ("34. The 18-Step 'Lost Laptop' Rapid Reconstruction", "40"),
        ("35. What is (and is NOT) Stored in GitHub", "41"),
        ("36. Critical Security & Secrets Governance Warning", "42"),
    ]
    
    t_toc = doc.add_table(rows=len(toc_items) + 1, cols=2)
    t_toc.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_toc.columns[0].width = Inches(5.5)
    t_toc.columns[1].width = Inches(1.0)
    
    hdr_cells = t_toc.rows[0].cells
    hdr_cells[0].text = "Section / Chapter Title"
    hdr_cells[1].text = "Page"
    set_cell_background(hdr_cells[0], "0F172A")
    set_cell_background(hdr_cells[1], "0F172A")
    hdr_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hdr_cells[0].paragraphs[0].runs[0].bold = True
    hdr_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hdr_cells[1].paragraphs[0].runs[0].bold = True
    
    for idx, (item, page) in enumerate(toc_items):
        row = t_toc.rows[idx + 1]
        row.cells[0].text = item
        row.cells[1].text = page
        set_cell_background(row.cells[0], "FFFFFF" if idx % 2 == 0 else "F8FAFC")
        set_cell_background(row.cells[1], "FFFFFF" if idx % 2 == 0 else "F8FAFC")
        set_cell_margins(row.cells[0], 40, 40, 60, 60)
        set_cell_margins(row.cells[1], 40, 40, 60, 60)
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_page_break()

    # --------------------------------------------------------------------------
    # CHAPTER 1: PROJECT OVERVIEW
    # --------------------------------------------------------------------------
    doc.add_heading("1. Project Overview & Capabilities", level=1)
    doc.add_paragraph(
        "NexusResearch is a full-stack, enterprise-grade AI research assistant engineered to automate grounded scholarly research, "
        "evidence synthesis, conflict auditing, and publication-ready academic document compilation. "
        "Unlike generic LLM wrappers that hallucinate facts or generate fabricated statistics, NexusResearch enforces a strict 7-stage "
        "deterministic pipeline that connects directly to real peer-reviewed registries and authoritative web databases."
    )
    
    doc.add_heading("Key Problems Solved", level=2)
    p = doc.add_paragraph()
    p.add_run("• Academic Hallucinations: ").bold = True
    p.add_run("Every factual claim synthesized by the system is linked to a verified sentence-level quote with an exact bracketed citation ([1], [2]).\n")
    p.add_run("• Research Fragmentation: ").bold = True
    p.add_run("Aggregates findings concurrently across OpenAlex, PubMed, Europe PMC, arXiv, Crossref, and Wikipedia into a single unified workspace.\n")
    p.add_run("• Manual Paper Formatting: ").bold = True
    p.add_run("Automatically exports validated research findings into double-column IEEE standard Word documents (.docx) and formal academic PDFs (.pdf).\n")
    p.add_run("• Data Governance & Multi-Tenancy: ").bold = True
    p.add_run("Ensures strict database-level isolation of user projects, research histories, and uploaded context documents.")

    doc.add_heading("Actual Verified Features", level=2)
    features = [
        ("User Authentication & OAuth Sync", "Hybrid NextAuth v5 + FastAPI JWT system supporting Google OAuth, GitHub OAuth, and local credentials."),
        ("Multi-Registry Search Engine", "Concurrent queries across OpenAlex, arXiv, PubMed, Europe PMC, Crossref, Wikipedia, and DuckDuckGo."),
        ("Hard Relevance Gating", "Automated scoring of candidate sources against research questions with rejection of irrelevant matches."),
        ("Interactive Evidence Matrix", "Selectable evidence items in the right-side panel that load grounded AI explanations into the center view with instant state restoration."),
        ("Contradiction & Conflict Detection", "Automated discovery of empirical conflicts, differing conclusions, and methodological divergences across literature."),
        ("Publication-Grade Document Compilers", "Deterministic generation of IEEE standard Word documents and ReportLab academic PDFs with citation audits."),
        ("Multi-Project Research Workspace", "Isolated project workspaces with conversation history, uploaded document chunking, and search scope toggles.")
    ]
    for feat, desc in features:
        p = doc.add_paragraph()
        p.add_run(f"• {feat}: ").bold = True
        p.add_run(desc)

    # --------------------------------------------------------------------------
    # CHAPTER 2: SYSTEM REQUIREMENTS
    # --------------------------------------------------------------------------
    doc.add_heading("2. System Requirements & Prerequisites Table", level=1)
    doc.add_paragraph("The following software prerequisites are required to build, test, and execute NexusResearch on a new computer:")

    req_table = [
        ("Software Component", "Required Version", "Purpose in NexusResearch", "Download URL", "Verification Command"),
        ("Windows OS", "10 / 11 (64-bit)", "Host operating system", "microsoft.com/windows", "systeminfo"),
        ("Visual Studio Code", "Latest Stable", "Integrated development environment", "code.visualstudio.com", "code --version"),
        ("Git for Windows", "v2.40.0+", "Version control & repository cloning", "git-scm.com/download/win", "git --version"),
        ("Node.js", "v18.17+ or v20+", "Next.js frontend runtime & tooling", "nodejs.org/en/download", "node --version"),
        ("npm", "v9.0+ or v10+", "Frontend package manager", "Bundled with Node.js", "npm --version"),
        ("Python", "v3.10.x - v3.12.x", "FastAPI backend & document engine", "python.org/downloads", "python --version"),
        ("pip", "Latest", "Python package manager", "Bundled with Python", "pip --version"),
        ("Google Gemini API", "API Access Key", "Synthesizing research answers", "aistudio.google.com", "N/A (Web Portal)"),
        ("PostgreSQL (Optional)", "v14+ / v16+", "Production database (SQLite dev)", "postgresql.org/download", "psql --version")
    ]
    
    t_req = doc.add_table(rows=len(req_table), cols=5)
    t_req.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row_vals in enumerate(req_table):
        row = t_req.rows[r_idx]
        for c_idx, val in enumerate(row_vals):
            cell = row.cells[c_idx]
            set_cell_background(cell, "0F172A" if r_idx == 0 else ("FFFFFF" if r_idx % 2 == 1 else "F8FAFC"))
            set_cell_margins(cell, 35, 35, 40, 40)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8)
            if r_idx == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    # --------------------------------------------------------------------------
    # CHAPTER 3: INSTALL VS CODE
    # --------------------------------------------------------------------------
    doc.add_heading("3. Visual Studio Code Installation & Terminal Setup", level=1)
    doc.add_paragraph("Visual Studio Code (VS Code) is the recommended development environment for editing, debugging, and running both the Next.js frontend and FastAPI backend.")
    
    steps_vscode = [
        "1. Navigate to https://code.visualstudio.com and download the 'User Installer (64-bit)' for Windows.",
        "2. Run the downloaded installer (.exe). Check the options: 'Add to PATH', 'Add Open with Code action to Windows Explorer file context menu'.",
        "3. Launch Visual Studio Code.",
        "4. Open the integrated terminal using keyboard shortcut: Ctrl + ` (backtick) or go to Terminal > New Terminal from the top menu.",
        "5. In the top right of the terminal window, verify that the shell dropdown says 'powershell' (or 'pwsh').",
        "6. Type `echo 'VS Code Terminal Ready'` and press Enter to verify execution."
    ]
    for s in steps_vscode:
        doc.add_paragraph(s)

    # --------------------------------------------------------------------------
    # CHAPTER 4: INSTALL GIT
    # --------------------------------------------------------------------------
    doc.add_heading("4. Git Version Control Installation", level=1)
    doc.add_paragraph(
        "Git is required to clone the repository, track revisions, and synchronize code with GitHub."
    )
    doc.add_paragraph("1. Visit https://git-scm.com/download/win and download the 64-bit Git for Windows Setup.")
    doc.add_paragraph("2. During installation, select 'Git from the command line and also from 3rd-party software' and 'Use Windows' default console window'.")
    doc.add_paragraph("3. Complete installation, open a new PowerShell terminal in VS Code, and verify:")
    add_code_block_docx(doc, "git --version\n# Expected output: git version 2.4x.x.windows.1")

    # --------------------------------------------------------------------------
    # CHAPTER 5: INSTALL RUNTIME ENVIRONMENTS
    # --------------------------------------------------------------------------
    doc.add_heading("5. Node.js & Python Runtime Environments", level=1)
    doc.add_paragraph("NexusResearch is a polyglot application requiring both Node.js (for Next.js) and Python (for FastAPI).")
    
    doc.add_heading("Installing Node.js & npm", level=2)
    doc.add_paragraph("1. Visit https://nodejs.org and download the 'LTS' (Long Term Support) installer (v20.x recommended).")
    doc.add_paragraph("2. Run the installer and check 'Automatically install the necessary tools'.")
    doc.add_paragraph("3. Verify installation in PowerShell:")
    add_code_block_docx(doc, "node --version   # Expected: v20.x.x\nnpm --version    # Expected: 10.x.x")

    doc.add_heading("Installing Python & pip", level=2)
    doc.add_paragraph("1. Visit https://www.python.org/downloads and download Python 3.10.x or 3.11.x.")
    doc.add_paragraph("2. CRITICAL: In the installer dialog, check the box: 'Add python.exe to PATH' before clicking Install Now.")
    doc.add_paragraph("3. Verify Python and pip in PowerShell:")
    add_code_block_docx(doc, "python --version # Expected: Python 3.10.x (or 3.11.x)\npip --version    # Expected: pip 24.x from ...")

    # --------------------------------------------------------------------------
    # CHAPTER 6: GETTING THE PROJECT
    # --------------------------------------------------------------------------
    doc.add_heading("6. Obtaining the Repository (Git Clone & ZIP)", level=1)
    doc.add_paragraph("You can download the project using either Git Clone (Recommended) or ZIP Extraction:")

    doc.add_heading("Method A: Git Clone (Recommended)", level=2)
    add_code_block_docx(
        doc,
        "# 1. Open PowerShell and navigate to your desired workspace\n"
        "cd C:\\Users\\YourUsername\\Projects\n\n"
        "# 2. Clone the repository\n"
        "git clone https://github.com/TejeshwarDivekar/NexusAI.git\n\n"
        "# 3. Enter project directory and open in VS Code\n"
        "cd NexusAI\n"
        "code ."
    )

    doc.add_heading("Method B: Download ZIP Archive", level=2)
    doc.add_paragraph("1. Open https://github.com/TejeshwarDivekar/NexusAI in your browser.")
    doc.add_paragraph("2. Click the green '<> Code' button and select 'Download ZIP'.")
    doc.add_paragraph("3. Extract the ZIP file into a permanent folder (e.g. `C:\\Projects\\NexusAI`).")
    doc.add_paragraph("4. Right-click the extracted folder and choose 'Open with Code'.")

    # --------------------------------------------------------------------------
    # CHAPTER 7: PROJECT STRUCTURE
    # --------------------------------------------------------------------------
    doc.add_heading("7. Verified Repository Directory Structure", level=1)
    doc.add_paragraph("The actual repository is structured as follows:")
    
    add_code_block_docx(
        doc,
        "NexusResearch/\n"
        "├── backend/                     # FastAPI Python backend application\n"
        "│   ├── app/                     # Backend core source code\n"
        "│   │   ├── core/                # Logging, security, exception handlers\n"
        "│   │   ├── db/                  # Database session, models, init_db\n"
        "│   │   ├── routers/             # API routes (auth, research, documents, projects)\n"
        "│   │   ├── schemas/             # Pydantic validation models\n"
        "│   │   └── services/            # Research engine, search providers, doc gen\n"
        "│   ├── tests/                   # Pytest automated test suite (30 tests)\n"
        "│   ├── Dockerfile               # Backend container definition\n"
        "│   ├── requirements.txt         # Python backend dependencies\n"
        "│   ├── railway.toml             # Railway deployment configuration\n"
        "│   └── run.py                   # Programmatic Uvicorn entrypoint\n"
        "├── docs/                        # Complete technical documentation\n"
        "│   ├── assets/                  # Architecture diagrams & PDF/DOCX assets\n"
        "│   ├── ARCHITECTURE.md          # Full system architecture guide\n"
        "│   ├── DEVELOPMENT.md           # Local developer onboarding\n"
        "│   ├── DATABASE.md              # Relational schema specification\n"
        "│   ├── API.md                   # REST & SSE API reference\n"
        "│   ├── DEPLOYMENT.md            # Railway & Docker Compose guide\n"
        "│   ├── TROUBLESHOOTING.md       # Operational troubleshooting guide\n"
        "│   └── FINAL_CLEANUP_REPORT.md  # Repository audit report\n"
        "├── public/                      # Static assets and icons\n"
        "├── scripts/                     # Documentation and validation scripts\n"
        "├── src/                         # Next.js 16 App Router frontend\n"
        "│   ├── app/                     # App router pages (/, /chat, /login, /api/...)\n"
        "│   ├── components/              # UI & Research Workspace components\n"
        "│   ├── lib/                     # Frontend client utilities\n"
        "│   └── auth.ts                  # NextAuth v5 authentication configuration\n"
        "├── .env.example                 # Root environment template (placeholders only)\n"
        "├── .gitignore                   # Git exclusion rules\n"
        "├── docker-compose.yml           # Multi-service local container orchestrator\n"
        "├── package.json                 # Frontend dependencies & Next.js scripts\n"
        "└── README.md                    # Project overview\n"
    )

    # --------------------------------------------------------------------------
    # CHAPTER 8: INSTALL DEPENDENCIES
    # --------------------------------------------------------------------------
    doc.add_heading("8. Dependency Installation (Frontend npm & Backend pip)", level=1)
    doc.add_paragraph("Dependencies must be installed for both the Node.js frontend and Python backend:")

    doc.add_heading("Step 1: Frontend Installation (npm)", level=2)
    doc.add_paragraph("In the root folder of the project (`NexusResearch/`), run:")
    add_code_block_docx(doc, "npm install")
    doc.add_paragraph("This installs Next.js 16.3.1, React 19, Lucide Icons, React Markdown, and NextAuth v5.")

    doc.add_heading("Step 2: Backend Installation (pip)", level=2)
    doc.add_paragraph("Open a terminal or navigate to the `backend/` directory:")
    add_code_block_docx(
        doc,
        "# Navigate to backend folder\n"
        "cd backend\n\n"
        "# Install Python dependencies\n"
        "pip install -r requirements.txt\n\n"
        "# Return to root\n"
        "cd .."
    )
    doc.add_paragraph("This installs FastAPI, Uvicorn, SQLAlchemy, Pydantic, python-docx, reportlab, and pytest.")

    # --------------------------------------------------------------------------
    # CHAPTER 9: ENVIRONMENT VARIABLES
    # --------------------------------------------------------------------------
    doc.add_heading("9. Comprehensive Environment Variables Master Table", level=1)
    doc.add_paragraph("The following environment variables govern system runtime, security, database access, and external APIs:")

    env_vars = [
        ("Variable Name", "Req?", "Used Where", "Purpose & Description", "Development Example"),
        ("AUTH_SECRET", "Yes", "Frontend (NextAuth)", "Cryptographic secret for signing session cookies", "Generate with: npx auth secret"),
        ("AUTH_TRUST_HOST", "Yes", "Frontend (NextAuth)", "Allows auth behind reverse proxy / localhost", "true"),
        ("NEXTAUTH_URL", "Yes", "Frontend (NextAuth)", "Base public URL for OAuth redirects", "http://localhost:3000"),
        ("BACKEND_INTERNAL_URL", "Yes", "Frontend Proxy", "Internal target for API proxy route", "http://127.0.0.1:8000"),
        ("GOOGLE_GENERATIVE_AI_API_KEY", "Yes", "Frontend & Backend", "API Key for Google Gemini LLM synthesis", "AIzaSy... (from aistudio.google.com)"),
        ("GOOGLE_API_KEY", "Yes", "Backend Engine", "API Key alias for backend Gemini provider", "Same as GOOGLE_GENERATIVE_AI_API_KEY"),
        ("SECRET_KEY", "Yes", "Backend Core", "Secret key for signing backend JWT tokens", "Generate 32+ random characters"),
        ("DATABASE_URL", "Yes", "Backend DB", "SQLAlchemy connection string", "sqlite:///./nexusai_research.db"),
        ("AUTH_GOOGLE_ID", "No", "Frontend OAuth", "Google Cloud OAuth Client ID", "your-client-id.apps.googleusercontent.com"),
        ("AUTH_GOOGLE_SECRET", "No", "Frontend OAuth", "Google Cloud OAuth Client Secret", "GOCSPX-your-secret-string"),
        ("AUTH_GITHUB_ID", "No", "Frontend OAuth", "GitHub OAuth Client ID", "Optional GitHub OAuth ID"),
        ("AUTH_GITHUB_SECRET", "No", "Frontend OAuth", "GitHub OAuth Client Secret", "Optional GitHub OAuth Secret"),
        ("TAVILY_API_KEY", "No", "Backend Search", "Optional Tavily web search fallback", "Optional (Public providers used by default)"),
        ("ENVIRONMENT", "No", "Backend Runtime", "Runtime mode: development or production", "development"),
        ("LOG_LEVEL", "No", "Backend Logging", "Logging verbosity: INFO or DEBUG", "INFO")
    ]
    
    t_env = doc.add_table(rows=len(env_vars), cols=5)
    t_env.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row_vals in enumerate(env_vars):
        row = t_env.rows[r_idx]
        for c_idx, val in enumerate(row_vals):
            cell = row.cells[c_idx]
            set_cell_background(cell, "0F172A" if r_idx == 0 else ("FFFFFF" if r_idx % 2 == 1 else "F8FAFC"))
            set_cell_margins(cell, 35, 35, 35, 35)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(7.5)
            if r_idx == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    # --------------------------------------------------------------------------
    # CHAPTER 10: CREATE .ENV FILES
    # --------------------------------------------------------------------------
    doc.add_heading("10. Creating .env Configuration Files", level=1)
    doc.add_paragraph("Two environment configuration files must be created from their templates:")

    doc.add_heading("1. Root Frontend File: `.env.local`", level=2)
    doc.add_paragraph("In the root directory, create a file named `.env.local`:")
    add_code_block_docx(
        doc,
        "# ==============================================================================\n"
        "# NexusResearch — Frontend Environment (.env.local)\n"
        "# ==============================================================================\n"
        "BACKEND_INTERNAL_URL=http://127.0.0.1:8000\n"
        "AUTH_SECRET=YOUR_RANDOM_32_CHARACTER_AUTH_SECRET_HERE\n"
        "AUTH_TRUST_HOST=true\n"
        "NEXTAUTH_URL=http://localhost:3000\n\n"
        "# Google Gemini API Key (Required for LLM research synthesis)\n"
        "GOOGLE_GENERATIVE_AI_API_KEY=YOUR_GEMINI_API_KEY_HERE\n\n"
        "# Google OAuth (Optional for local development; required for Google Login)\n"
        "AUTH_GOOGLE_ID=YOUR_GOOGLE_CLIENT_ID_HERE\n"
        "AUTH_GOOGLE_SECRET=YOUR_GOOGLE_CLIENT_SECRET_HERE\n"
    )

    doc.add_heading("2. Backend File: `backend/.env`", level=2)
    doc.add_paragraph("In the `backend/` directory, create a file named `backend/.env`:")
    add_code_block_docx(
        doc,
        "# ==============================================================================\n"
        "# NexusResearch — Backend Environment (backend/.env)\n"
        "# ==============================================================================\n"
        "PROJECT_NAME=\"Enterprise AI Research Assistant API\"\n"
        "VERSION=\"1.0.0\"\n"
        "ENVIRONMENT=development\n"
        "LOG_LEVEL=INFO\n\n"
        "# Backend Security\n"
        "SECRET_KEY=YOUR_RANDOM_32_CHARACTER_BACKEND_SECRET_KEY\n"
        "ALGORITHM=HS256\n"
        "ACCESS_TOKEN_EXPIRE_MINUTES=10080\n\n"
        "# Database (Zero-config SQLite for local development)\n"
        "DATABASE_URL=sqlite:///./nexusai_research.db\n\n"
        "# AI Model Providers\n"
        "GOOGLE_API_KEY=YOUR_GEMINI_API_KEY_HERE\n"
        "GOOGLE_GENERATIVE_AI_API_KEY=YOUR_GEMINI_API_KEY_HERE\n\n"
        "MAX_UPLOAD_SIZE_MB=25\n"
        "CORS_ORIGINS=[\"http://localhost:3000\", \"http://127.0.0.1:3000\"]\n"
    )

    # --------------------------------------------------------------------------
    # CHAPTER 11: DATABASE SETUP
    # --------------------------------------------------------------------------
    doc.add_heading("11. Database Setup: SQLite (Dev) & PostgreSQL (Prod)", level=1)
    doc.add_paragraph(
        "NexusResearch features zero-configuration database setup for local development. "
        "Local database installation (such as installing standalone PostgreSQL) is NOT required."
    )
    
    add_callout_docx(
        doc,
        "ZERO-CONFIG SQLITE FOR LOCAL DEVELOPMENT:\n"
        "When DATABASE_URL is set to `sqlite:///./nexusai_research.db`, the FastAPI backend automatically creates the database file and initializes all tables and indexes upon application startup via `init_db()`. "
        "No database server installation, user creation, or manual migration commands are required.",
        title="ZERO-CONFIGURATION LOCAL DATABASE"
    )
    
    doc.add_heading("Production PostgreSQL Setup (Cloud / Railway)", level=2)
    doc.add_paragraph(
        "In production environments (e.g. Railway), the database URL is provided automatically as a connection string:\n"
        "`postgresql://postgres:password@postgres.railway.internal:5432/railway`\n"
        "SQLAlchemy automatically adapts to PostgreSQL without any code changes."
    )

    # --------------------------------------------------------------------------
    # CHAPTER 12: GOOGLE LOGIN SETUP
    # --------------------------------------------------------------------------
    doc.add_heading("12. Google Cloud OAuth & Login Configuration", level=1)
    doc.add_paragraph("To enable 'Sign in with Google' on local and production environments:")
    
    google_steps = [
        "1. Open the Google Cloud Console at https://console.cloud.google.com and sign in.",
        "2. Create a new project named 'NexusResearch' (or select an existing project).",
        "3. In the left navigation menu, go to APIs & Services > OAuth consent screen.",
        "4. Choose User Type: 'External' and click Create.",
        "5. Fill in App Name ('NexusResearch'), User support email, and Developer contact information.",
        "6. In the left menu, click Credentials > Create Credentials > OAuth client ID.",
        "7. Select Application type: 'Web application'.",
        "8. Under 'Authorized JavaScript origins', add: `http://localhost:3000` (and your production domain if deployed).",
        "9. Under 'Authorized redirect URIs', add: `http://localhost:3000/api/auth/callback/google`.",
        "10. Click Create. Copy the Client ID and Client Secret.",
        "11. Paste the values into your `.env.local` file as `AUTH_GOOGLE_ID` and `AUTH_GOOGLE_SECRET`."
    ]
    for s in google_steps:
        doc.add_paragraph(s)

    # --------------------------------------------------------------------------
    # CHAPTER 13: LLM SETUP
    # --------------------------------------------------------------------------
    doc.add_heading("13. AI LLM Engine: Google Gemini Configuration", level=1)
    doc.add_paragraph(
        "NexusResearch uses Google Gemini (`gemini-2.5-flash` / `gemini-1.5-pro`) for research synthesis, query decomposition, and evidence explanations."
    )
    doc.add_paragraph("1. Visit Google AI Studio at https://aistudio.google.com.")
    doc.add_paragraph("2. Sign in with your Google Account and click 'Get API key'.")
    doc.add_paragraph("3. Click 'Create API key in new project' and copy the generated key string.")
    doc.add_paragraph("4. Add the key to `.env.local` and `backend/.env` as `GOOGLE_GENERATIVE_AI_API_KEY` and `GOOGLE_API_KEY`.")
    doc.add_paragraph("5. Fallback behavior: If the Gemini API experiences network interruption, the search provider returns raw grounded citations with an explicit notice.")

    # --------------------------------------------------------------------------
    # CHAPTER 14: RESEARCH PROVIDERS
    # --------------------------------------------------------------------------
    doc.add_heading("14. Scholarly Research Registries & Search Providers", level=1)
    doc.add_paragraph(
        "NexusResearch queries real peer-reviewed scientific registries and authoritative web databases concurrently. "
        "Zero API keys are required for the primary scholarly providers:"
    )

    prov_table = [
        ("Provider Name", "Target Data", "API Key Req?", "Protocol / Endpoint", "What It Retrieves"),
        ("OpenAlex", "Global Scholarly Works", "No (Open)", "REST API (`api.openalex.org`)", "Peer-reviewed papers, DOIs, abstracts, authors"),
        ("arXiv", "Preprints & CS / Physics", "No (Open)", "XML Query (`export.arxiv.org`)", "Recent preprint papers, abstracts, categories"),
        ("PubMed", "Biomedical & Life Sciences", "No (Open)", "E-utilities (`eutils.ncbi.nlm.nih.gov`)", "Medical journal abstracts, PMID citations"),
        ("Europe PMC", "Open Access Life Sciences", "No (Open)", "REST API (`europepmc.org`)", "Open access journal articles, PMC articles"),
        ("Crossref", "Scholarly DOI Metadata", "No (Open)", "REST API (`api.crossref.org`)", "Publication metadata, journal titles, dates"),
        ("Wikipedia", "Encyclopedic Context", "No (Open)", "REST API (`en.wikipedia.org`)", "Definitional summaries and historical context"),
        ("DuckDuckGo", "Authoritative Web Data", "No (Open)", "HTML Scraper / Instant API", "Contemporary web articles and industry news")
    ]
    
    t_prov = doc.add_table(rows=len(prov_table), cols=5)
    t_prov.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row_vals in enumerate(prov_table):
        row = t_prov.rows[r_idx]
        for c_idx, val in enumerate(row_vals):
            cell = row.cells[c_idx]
            set_cell_background(cell, "0F172A" if r_idx == 0 else ("FFFFFF" if r_idx % 2 == 1 else "F8FAFC"))
            set_cell_margins(cell, 35, 35, 35, 35)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(7.5)
            if r_idx == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    # --------------------------------------------------------------------------
    # CHAPTER 15: RUN BACKEND
    # --------------------------------------------------------------------------
    doc.add_heading("15. Starting the FastAPI Backend Server", level=1)
    doc.add_paragraph("Open a dedicated PowerShell terminal in VS Code to run the backend:")
    add_code_block_docx(
        doc,
        "# 1. Ensure you are in the project root\n"
        "cd C:\\Projects\\NexusAI\n\n"
        "# 2. Run backend via entrypoint script\n"
        "python backend/run.py\n\n"
        "# Expected Output:\n"
        "# Starting Uvicorn server on 0.0.0.0:8000...\n"
        "# INFO:     Application startup complete.\n"
        "# INFO:     Uvicorn running on http://0.0.0.0:8000"
    )
    doc.add_paragraph("Verify by visiting `http://127.0.0.1:8000/api/v1/health` in your browser. Expected response: `{\"status\":\"healthy\"}`.")

    # --------------------------------------------------------------------------
    # CHAPTER 16: RUN FRONTEND
    # --------------------------------------------------------------------------
    doc.add_heading("16. Starting the Next.js Frontend Server", level=1)
    doc.add_paragraph("Open a second PowerShell terminal in VS Code to run the frontend:")
    add_code_block_docx(
        doc,
        "# 1. Ensure you are in the project root\n"
        "cd C:\\Projects\\NexusAI\n\n"
        "# 2. Start Next.js development server\n"
        "npm run dev\n\n"
        "# Expected Output:\n"
        "# ▲ Next.js 16.3.1 (webpack)\n"
        "# - Local:        http://localhost:3000\n"
        "# ✓ Ready in 1.5s"
    )
    doc.add_paragraph("Open [http://localhost:3000](http://localhost:3000) in Google Chrome or Microsoft Edge.")

    # --------------------------------------------------------------------------
    # CHAPTER 17: DUAL TERMINAL WORKFLOW
    # --------------------------------------------------------------------------
    doc.add_heading("17. Dual-Terminal Local Execution Workflow", level=1)
    doc.add_paragraph("NexusResearch runs using two coordinated terminal processes in VS Code:")
    
    add_code_block_docx(
        doc,
        "Visual Studio Code Workspace\n"
        "├── Terminal 1 (Backend Server)  → python backend/run.py   (Port 8000)\n"
        "└── Terminal 2 (Frontend Client) → npm run dev            (Port 3000)\n\n"
        "User Browser (http://localhost:3000)\n"
        "      │\n"
        "      ▼\n"
        "Next.js Frontend (Port 3000)\n"
        "      │ (Internal Reverse Proxy: /api/v1/...)\n"
        "      ▼\n"
        "FastAPI Backend (Port 8000)"
    )

    # --------------------------------------------------------------------------
    # CHAPTER 18: FIRST LOGIN
    # --------------------------------------------------------------------------
    doc.add_heading("18. Performing First Login & Account Sync", level=1)
    doc.add_paragraph("1. Open `http://localhost:3000` in your web browser.")
    doc.add_paragraph("2. If not authenticated, you will be presented with the sign-in options: 'Continue with Google', 'Continue with GitHub', or demo credentials.")
    doc.add_paragraph("3. Click 'Continue with Google' and complete the Google consent screen.")
    doc.add_paragraph("4. Upon redirect, the frontend automatically synchronizes the OAuth user profile with the backend database via `/api/v1/auth/oauth_sync`.")

    # --------------------------------------------------------------------------
    # CHAPTER 19: FIRST RESEARCH TEST
    # --------------------------------------------------------------------------
    doc.add_heading("19. Executing Real Research Inquiries", level=1)
    doc.add_paragraph("In the main Launchpad, enter a test research query:")
    doc.add_paragraph("• 'What is artificial intelligence?'")
    doc.add_paragraph("• 'How does machine learning work?'")
    doc.add_paragraph("• 'What are the applications of AI in healthcare?'")
    doc.add_paragraph("• 'What are recent developments in renewable energy?'")
    doc.add_paragraph("Click 'Start Research'. You will observe real-time progress indicators as the system executes Stage 1 through Stage 7, displaying real academic papers and evidence items.")

    # --------------------------------------------------------------------------
    # CHAPTER 20: TEST EVIDENCE MATRIX
    # --------------------------------------------------------------------------
    doc.add_heading("20. Testing Interactive Evidence Matrix & Restoration", level=1)
    doc.add_paragraph("1. In the completed research workspace, locate the Evidence Matrix in the right-side panel.")
    doc.add_paragraph("2. Click on any specific evidence claim card.")
    doc.add_paragraph("3. The center panel transitions to the Evidence Detail View, displaying the verbatim paper quote, source metadata, and an on-demand AI explanation.")
    doc.add_paragraph("4. Click the '← Back to Answer' button in the center top bar.")
    doc.add_paragraph("5. Verify that the center panel restores the full research report with exact scroll position preserved without re-running the research query.")

    # --------------------------------------------------------------------------
    # CHAPTER 21: TEST DOCUMENT GENERATION
    # --------------------------------------------------------------------------
    doc.add_heading("21. Validating IEEE DOCX & Academic PDF Generation", level=1)
    doc.add_paragraph("1. In the Report Viewer header, locate the 'Export IEEE (.docx)' and 'Export PDF (.pdf)' download buttons.")
    doc.add_paragraph("2. Click 'Export IEEE (.docx)' and open the downloaded Word document in Microsoft Word.")
    doc.add_paragraph("3. Verify double-column layout, formal abstract, numbered headings, and bracketed IEEE references [1], [2].")
    doc.add_paragraph("4. Click 'Export PDF (.pdf)' and verify the formal ReportLab PDF with page headers, footers, and page numbers.")

    # --------------------------------------------------------------------------
    # CHAPTER 22: TEST RESEARCH HISTORY
    # --------------------------------------------------------------------------
    doc.add_heading("22. Multi-Tenant Project Isolation & History", level=1)
    doc.add_paragraph("1. In the left sidebar, locate the 'Recent Investigations' history list.")
    doc.add_paragraph("2. Click on a previous conversation to reload its full report, sources, and evidence matrix instantly.")
    doc.add_paragraph("3. Verify that user research sessions remain permanently stored and isolated by `user_id` in the database.")

    # --------------------------------------------------------------------------
    # CHAPTER 23: TEST MOBILE
    # --------------------------------------------------------------------------
    doc.add_heading("23. Mobile View & Responsive Layout Testing", level=1)
    doc.add_paragraph("1. Press F12 in Google Chrome to open Developer Tools.")
    doc.add_paragraph("2. Click the 'Toggle Device Toolbar' icon (Ctrl + Shift + M) and select 'iPhone 14' or 'Pixel 7'.")
    doc.add_paragraph("3. Verify the mobile bottom navigation bar with three dedicated tabs: 'Report', 'Sources', and 'Evidence'.")
    doc.add_paragraph("4. Verify evidence selection and document downloads function smoothly on mobile viewports.")

    # --------------------------------------------------------------------------
    # CHAPTER 24: COMMON ERRORS & TROUBLESHOOTING
    # --------------------------------------------------------------------------
    doc.add_heading("24. Comprehensive Troubleshooting Matrix", level=1)
    doc.add_paragraph("Quick solutions for common operational and configuration errors:")

    trouble_table = [
        ("Observed Problem", "Possible Root Cause", "How to Verify", "Permanent Solution"),
        ("Research Service Error Notice", "Backend container crashed or port parsing failed", "Check terminal: ValueError $PORT or missing reportlab", "Ensure reportlab is in requirements.txt and use python run.py"),
        ("'node' or 'npm' not recognized", "Node.js not added to system PATH", "Type `node -v` in fresh terminal", "Reinstall Node.js and check 'Add to PATH'"),
        ("'python' opens Microsoft Store", "Windows app execution alias intercepting", "Type `python` in cmd", "Disable 'App Execution Aliases' in Windows Settings"),
        ("Port 8000 or 3000 in use", "Another process holding the port", "Run `netstat -ano | findstr :8000`", "Kill conflicting process: `taskkill /PID <PID> /F`"),
        ("OAuth redirect_uri_mismatch", "Google Cloud redirect URI mismatch", "Check Google Console Credentials", "Set authorized redirect to `http://localhost:3000/api/auth/callback/google`"),
        ("Zero sources found for query", "Query terms too obscure or narrow", "Check backend search logs", "Click [Retry Research] or broaden inquiry terms")
    ]
    
    t_tr = doc.add_table(rows=len(trouble_table), cols=4)
    t_tr.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row_vals in enumerate(trouble_table):
        row = t_tr.rows[r_idx]
        for c_idx, val in enumerate(row_vals):
            cell = row.cells[c_idx]
            set_cell_background(cell, "0F172A" if r_idx == 0 else ("FFFFFF" if r_idx % 2 == 1 else "F8FAFC"))
            set_cell_margins(cell, 35, 35, 35, 35)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(7.5)
            if r_idx == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    # --------------------------------------------------------------------------
    # CHAPTER 25: HOW TO READ ERRORS
    # --------------------------------------------------------------------------
    doc.add_heading("25. How to Read & Diagnose Errors Across Layers", level=1)
    doc.add_paragraph("• Frontend (Browser Console): Press F12 > Console. Inspect `[Research Service Error]` payloads containing `{ status, message, taskId }`.")
    doc.add_paragraph("• Frontend (Network Tab): Press F12 > Network > Fetch/XHR. Click failed requests to inspect exact JSON error details.")
    doc.add_paragraph("• Backend Terminal Logs: Inspect structured logs prefixed with `[Research {task_id}]` showing exact pipeline stage failure.")

    # --------------------------------------------------------------------------
    # CHAPTER 26: PROD VS DEV
    # --------------------------------------------------------------------------
    doc.add_heading("26. Development vs. Production Configuration", level=1)
    doc.add_paragraph("• Database: SQLite (`sqlite:///./nexusai_research.db`) in Dev vs PostgreSQL in Production.")
    doc.add_paragraph("• Host & Ports: `localhost:3000` & `127.0.0.1:8000` in Dev vs Railway private networking `http://backend.railway.internal:8000` in Prod.")
    doc.add_paragraph("• HTTPS & Cookies: `AUTH_TRUST_HOST=true` required behind cloud reverse proxies.")

    # --------------------------------------------------------------------------
    # CHAPTER 27: DEPLOYMENT
    # --------------------------------------------------------------------------
    doc.add_heading("27. Cloud Deployment: Railway & Docker Compose", level=1)
    doc.add_paragraph("The repository is configured for cloud deployment on Railway:")
    doc.add_paragraph("1. Create a project on Railway and link the GitHub repository.")
    doc.add_paragraph("2. Add a PostgreSQL database service.")
    doc.add_paragraph("3. Add the Backend service with `rootDirectory = 'backend'`, `dockerfilePath = 'Dockerfile'`, and `startCommand = 'python run.py'`.")
    doc.add_paragraph("4. Add the Frontend service with `rootDirectory = '/'` and `BACKEND_INTERNAL_URL = 'http://backend.railway.internal:8000'`.")

    # --------------------------------------------------------------------------
    # CHAPTER 28: BACKUP & RESTORE
    # --------------------------------------------------------------------------
    doc.add_heading("28. Database Backup & Disaster Recovery", level=1)
    doc.add_paragraph("• SQLite Backup: Simply copy `nexusai_research.db` to a backup location.")
    doc.add_paragraph("• PostgreSQL Backup: `pg_dump -U postgres -d railway > backup.sql`")
    doc.add_paragraph("• PostgreSQL Restore: `psql -U postgres -d railway < backup.sql`")

    # --------------------------------------------------------------------------
    # CHAPTER 29: DEV WORKFLOW
    # --------------------------------------------------------------------------
    doc.add_heading("29. Standard Developer Contribution Workflow", level=1)
    doc.add_paragraph("1. `git pull origin main`")
    doc.add_paragraph("2. Start backend (`python backend/run.py`) and frontend (`npm run dev`).")
    doc.add_paragraph("3. Execute automated tests (`python -m pytest backend/tests -v`).")
    doc.add_paragraph("4. Build production bundle (`npm run build`).")
    doc.add_paragraph("5. Commit and push (`git commit -m 'feat: ...'` and `git push origin main`).")

    # --------------------------------------------------------------------------
    # CHAPTER 30: GIT WORKFLOW
    # --------------------------------------------------------------------------
    doc.add_heading("30. Git Branching & Commit Conventions", level=1)
    doc.add_paragraph("Use conventional commit prefixes: `feat:`, `fix:`, `chore:`, `docs:`, `test:`. Always verify `git status` before staging.")

    # --------------------------------------------------------------------------
    # CHAPTER 31: FRESH LAPTOP CHECKLIST
    # --------------------------------------------------------------------------
    doc.add_heading("31. Printable Fresh-Laptop Setup Checklist", level=1)
    checklist_items = [
        "[ ] 1. Windows updated to latest stable version",
        "[ ] 2. Visual Studio Code installed with PowerShell terminal configured",
        "[ ] 3. Git for Windows installed (`git --version` verified)",
        "[ ] 4. Node.js v20.x and npm installed (`node --version` verified)",
        "[ ] 5. Python 3.10+ installed with 'Add to PATH' checked (`python --version` verified)",
        "[ ] 6. Repository cloned into local workspace directory",
        "[ ] 7. Frontend dependencies installed via `npm install`",
        "[ ] 8. Backend dependencies installed via `cd backend && pip install -r requirements.txt`",
        "[ ] 9. Root `.env.local` created with valid `AUTH_SECRET` and `GOOGLE_GENERATIVE_AI_API_KEY`",
        "[ ] 10. Backend `backend/.env` created with `SECRET_KEY` and `GOOGLE_API_KEY`",
        "[ ] 11. Backend server starts cleanly on port 8000 (`python backend/run.py`)",
        "[ ] 12. Frontend server starts cleanly on port 3000 (`npm run dev`)",
        "[ ] 13. Health endpoint `http://127.0.0.1:8000/api/v1/health` returns status healthy",
        "[ ] 14. Application opens in browser at `http://localhost:3000`",
        "[ ] 15. Research query 'What is artificial intelligence?' executes with real academic sources",
        "[ ] 16. Evidence Matrix interactive click, detail explanation, and Back to Answer tested",
        "[ ] 17. IEEE Word document (.docx) downloaded and verified in Microsoft Word",
        "[ ] 18. Academic PDF (.pdf) downloaded and verified",
        "[ ] 19. Pytest backend test suite passes 100% (`python -m pytest backend/tests -v`)",
        "[ ] 20. Production build passes cleanly (`npm run build`)"
    ]
    for chk in checklist_items:
        doc.add_paragraph(chk)

    # --------------------------------------------------------------------------
    # CHAPTER 32: COMMAND CHEAT SHEET
    # --------------------------------------------------------------------------
    doc.add_heading("32. Unified Command Cheat Sheet", level=1)
    add_code_block_docx(
        doc,
        "# ==============================================================================\n"
        "# NEXUSRESEARCH COMMAND CHEAT SHEET\n"
        "# ==============================================================================\n"
        "# Clone Repository\n"
        "git clone https://github.com/TejeshwarDivekar/NexusAI.git\n\n"
        "# Install Frontend Dependencies\n"
        "npm install\n\n"
        "# Install Backend Dependencies\n"
        "cd backend && pip install -r requirements.txt && cd ..\n\n"
        "# Start Backend Server (Port 8000)\n"
        "python backend/run.py\n\n"
        "# Start Frontend Server (Port 3000)\n"
        "npm run dev\n\n"
        "# Run Backend Automated Tests (30 Tests)\n"
        "python -m pytest backend/tests -v\n\n"
        "# Run Frontend Production Build Validation\n"
        "npm run build\n\n"
        "# Docker Compose Multi-Service Start\n"
        "docker compose up -d --build"
    )

    # --------------------------------------------------------------------------
    # CHAPTER 33: ARCHITECTURE DIAGRAM
    # --------------------------------------------------------------------------
    doc.add_heading("33. System Architecture Topology Diagram", level=1)
    add_code_block_docx(
        doc,
        "User Browser (http://localhost:3000)\n"
        "  │\n"
        "  ▼\n"
        "Next.js 16 App Router (Port 3000)\n"
        "  ├── Launchpad & Query Composer\n"
        "  ├── Research Workspace & Tabbed Center Panel\n"
        "  ├── Interactive Evidence Matrix & Detail Inspector\n"
        "  └── Proxy Endpoint: /api/v1/[...path]/route.ts\n"
        "        │ (HTTP / SSE Proxy to BACKEND_INTERNAL_URL)\n"
        "        ▼\n"
        "FastAPI 0.115 Backend Core (Port 8000)\n"
        "  ├── Auth & OAuth Sync Router (/api/v1/auth)\n"
        "  ├── Research Pipeline Engine (/api/v1/research/run & /stream)\n"
        "  │     ├── 1. Query Classifier\n"
        "  │     ├── 2. Multi-Registry Search (OpenAlex, arXiv, PubMed, etc.)\n"
        "  │     ├── 3. Hard Relevance Gating Scorer\n"
        "  │     ├── 4. Sentence-Level Evidence Extractor\n"
        "  │     ├── 5. Contradiction & Conflict Auditor\n"
        "  │     ├── 6. Numerical & Tabular Data Analyzer\n"
        "  │     └── 7. Gemini 2.5 LLM Grounded Report Synthesizer\n"
        "  ├── Document Generation Engine (IEEE .docx & Academic .pdf)\n"
        "  └── SQLAlchemy Database Layer (PostgreSQL / SQLite)"
    )

    # --------------------------------------------------------------------------
    # CHAPTER 34: LOST LAPTOP RECOVERY
    # --------------------------------------------------------------------------
    doc.add_heading("34. The 18-Step 'Lost Laptop' Rapid Reconstruction", level=1)
    doc.add_paragraph("If setting up on a completely new machine with zero prior tools installed:")
    doc.add_paragraph("1. Install VS Code > 2. Install Git > 3. Install Node.js v20 > 4. Install Python 3.10+ (Add to PATH) > 5. `git clone https://github.com/TejeshwarDivekar/NexusAI.git` > 6. `cd NexusAI` > 7. `npm install` > 8. `cd backend && pip install -r requirements.txt && cd ..` > 9. Create `.env.local` > 10. Create `backend/.env` > 11. Add Gemini API Key > 12. In Terminal 1: `python backend/run.py` > 13. In Terminal 2: `npm run dev` > 14. Open `http://localhost:3000` > 15. Submit test research query > 16. Verify Evidence Matrix click > 17. Verify IEEE Word export > 18. Run `npm run build` & `pytest backend/tests`.")

    # --------------------------------------------------------------------------
    # CHAPTER 35: WHAT IS NOT STORED IN GITHUB
    # --------------------------------------------------------------------------
    doc.add_heading("35. What is (and is NOT) Stored in GitHub", level=1)
    doc.add_paragraph("• STORED IN GITHUB: All frontend and backend source code, Dockerfiles, package manifests, documentation guides, template environment files (`.env.example`), and automated tests.")
    doc.add_paragraph("• NEVER STORED IN GITHUB: Real API keys (`GOOGLE_API_KEY`), OAuth secrets (`AUTH_GOOGLE_SECRET`), session signing keys (`AUTH_SECRET`, `SECRET_KEY`), database passwords, production `.env` files, or binary database `.db` files.")

    # --------------------------------------------------------------------------
    # CHAPTER 36: SECURITY WARNING
    # --------------------------------------------------------------------------
    doc.add_heading("36. Critical Security & Secrets Governance Warning", level=1)
    add_callout_docx(
        doc,
        "CRITICAL SECURITY GOVERNANCE MANDATE:\n"
        "Never commit `.env.local`, `backend/.env`, or any file containing real API keys or database passwords to Git. "
        "Always verify your staged changes with `git status` and `git diff` before creating commits. "
        "Keep production API keys and database credentials strictly managed inside your cloud provider environment settings (e.g. Railway Variables).",
        title="SECURITY WARNING"
    )

    doc.save(filepath)


# ==============================================================================
# BUILD PDF GUIDE (REPORTLAB)
# ==============================================================================
def build_pdf_guide(filepath):
    doc = SimpleDocTemplate(
        filepath,
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54
    )

    styles = getSampleStyleSheet()
    
    # Custom Palette
    c_primary = colors.HexColor("#0F172A")
    c_accent = colors.HexColor("#0E7490")
    c_body = colors.HexColor("#334155")
    c_bg_light = colors.HexColor("#F8FAFC")
    c_border = colors.HexColor("#E2E8F0")

    # Typography styles
    styles.add(ParagraphStyle('DocTitle', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=22, leading=26, textColor=c_primary, spaceAfter=8))
    styles.add(ParagraphStyle('DocSubTitle', parent=styles['Normal'], fontName='Helvetica', fontSize=12, leading=16, textColor=c_accent, spaceAfter=18))
    styles.add(ParagraphStyle('SectionH1', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=14, leading=18, textColor=c_primary, spaceBefore=14, spaceAfter=6, keepWithNext=True))
    styles.add(ParagraphStyle('SectionH2', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=11, leading=15, textColor=c_accent, spaceBefore=10, spaceAfter=4, keepWithNext=True))
    styles.add(ParagraphStyle('BodyCustom', parent=styles['Normal'], fontName='Helvetica', fontSize=9, leading=13, textColor=c_body, spaceAfter=4))
    styles.add(ParagraphStyle('BulletCustom', parent=styles['Normal'], fontName='Helvetica', fontSize=8.5, leading=12, textColor=c_body, leftIndent=12, spaceAfter=3))
    styles.add(ParagraphStyle('CodeBlock', parent=styles['Normal'], fontName='Courier', fontSize=7.5, leading=10, textColor=c_primary))
    styles.add(ParagraphStyle('CalloutText', parent=styles['Normal'], fontName='Helvetica', fontSize=8.5, leading=12, textColor=c_primary))
    styles.add(ParagraphStyle('TableCell', parent=styles['Normal'], fontName='Helvetica', fontSize=7.5, leading=10, textColor=c_body))
    styles.add(ParagraphStyle('TableHead', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=8, leading=11, textColor=colors.white))

    def make_callout(text, title="NOTE"):
        content = [
            Paragraph(f"<b>[{title}]</b> {text}", styles['CalloutText'])
        ]
        t = Table([[content]], colWidths=[504])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor("#F0F4F8")),
            ('BOX', (0, 0), (-1, -1), 1.0, colors.HexColor("#0284C7")),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ]))
        return t

    def make_code(code_str):
        # Escape XML entities for ReportLab Paragraph
        escaped = code_str.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('\n', '<br/>').replace(' ', '&nbsp;')
        p = Paragraph(escaped, styles['CodeBlock'])
        t = Table([[p]], colWidths=[504])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), c_bg_light),
            ('BOX', (0, 0), (-1, -1), 0.5, c_border),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ]))
        return t

    story = []

    # --------------------------------------------------------------------------
    # COVER PAGE
    # --------------------------------------------------------------------------
    story.append(Spacer(1, 40))
    story.append(Paragraph("NEXUSRESEARCH ENTERPRISE PLATFORM", styles['DocSubTitle']))
    story.append(Paragraph("Complete Project Setup, Local Reconstruction & Deployment Guide", styles['DocTitle']))
    story.append(Paragraph("A Complete, Standalone Manual to Set Up, Run, Test, and Deploy NexusResearch on Any Windows Computer from Scratch Without AI Assistance", styles['BodyCustom']))
    story.append(Spacer(1, 15))
    
    story.append(make_callout(
        "This document is authored directly from the NexusResearch production repository. "
        "Every command, port, environment variable, database table, API endpoint, and deployment step in this guide has been verified against active code. "
        "No real API keys, passwords, or OAuth secrets are published in this guide; placeholder formats are provided throughout.",
        title="AUTHENTIC SOURCE OF TRUTH"
    ))
    story.append(Spacer(1, 20))

    meta_text = (
        "<b>Author:</b> NexusResearch Engineering & Architecture Group<br/>"
        "<b>Version:</b> 1.0.0 Production Release<br/>"
        "<b>Target Environment:</b> Windows 10 / Windows 11 (64-bit)<br/>"
        f"<b>Date of Publication:</b> {datetime.datetime.now().strftime('%B %d, %Y')}<br/>"
        "<b>Repository Source:</b> https://github.com/TejeshwarDivekar/NexusAI.git"
    )
    story.append(Paragraph(meta_text, styles['BodyCustom']))
    story.append(PageBreak())

    # --------------------------------------------------------------------------
    # CHAPTER 1: PROJECT OVERVIEW
    # --------------------------------------------------------------------------
    story.append(Paragraph("1. Project Overview & Capabilities", styles['SectionH1']))
    story.append(Paragraph(
        "NexusResearch is a full-stack, enterprise-grade AI research assistant engineered to automate grounded scholarly research, "
        "evidence synthesis, conflict auditing, and publication-ready academic document compilation. "
        "Unlike generic LLM wrappers that hallucinate facts or generate fabricated statistics, NexusResearch enforces a strict 7-stage "
        "deterministic pipeline that connects directly to real peer-reviewed registries and authoritative web databases.",
        styles['BodyCustom']
    ))
    story.append(Paragraph("Key Problems Solved:", styles['SectionH2']))
    story.append(Paragraph("• <b>Academic Hallucinations:</b> Every factual claim synthesized by the system is linked to a verified sentence-level quote with an exact bracketed citation ([1], [2]).", styles['BulletCustom']))
    story.append(Paragraph("• <b>Research Fragmentation:</b> Aggregates findings concurrently across OpenAlex, PubMed, Europe PMC, arXiv, Crossref, and Wikipedia into a single unified workspace.", styles['BulletCustom']))
    story.append(Paragraph("• <b>Manual Paper Formatting:</b> Automatically exports validated research findings into double-column IEEE standard Word documents (.docx) and formal academic PDFs (.pdf).", styles['BulletCustom']))
    story.append(Paragraph("• <b>Data Governance & Multi-Tenancy:</b> Ensures strict database-level isolation of user projects, research histories, and uploaded context documents.", styles['BulletCustom']))
    story.append(Spacer(1, 8))

    # --------------------------------------------------------------------------
    # CHAPTER 2: SYSTEM REQUIREMENTS
    # --------------------------------------------------------------------------
    story.append(Paragraph("2. System Requirements & Prerequisites Table", styles['SectionH1']))
    
    req_data = [
        [Paragraph("Software", styles['TableHead']), Paragraph("Required Version", styles['TableHead']), Paragraph("Purpose in NexusResearch", styles['TableHead']), Paragraph("Verification Command", styles['TableHead'])],
        [Paragraph("Windows OS", styles['TableCell']), Paragraph("10 / 11 (64-bit)", styles['TableCell']), Paragraph("Host operating system", styles['TableCell']), Paragraph("systeminfo", styles['TableCell'])],
        [Paragraph("VS Code", styles['TableCell']), Paragraph("Latest Stable", styles['TableCell']), Paragraph("Development IDE & terminal", styles['TableCell']), Paragraph("code --version", styles['TableCell'])],
        [Paragraph("Git for Windows", styles['TableCell']), Paragraph("v2.40.0+", styles['TableCell']), Paragraph("Version control & cloning", styles['TableCell']), Paragraph("git --version", styles['TableCell'])],
        [Paragraph("Node.js", styles['TableCell']), Paragraph("v18.17+ / v20+", styles['TableCell']), Paragraph("Next.js frontend runtime", styles['TableCell']), Paragraph("node --version", styles['TableCell'])],
        [Paragraph("npm", styles['TableCell']), Paragraph("v9.0+ / v10+", styles['TableCell']), Paragraph("Frontend package manager", styles['TableCell']), Paragraph("npm --version", styles['TableCell'])],
        [Paragraph("Python", styles['TableCell']), Paragraph("v3.10.x - v3.12.x", styles['TableCell']), Paragraph("FastAPI backend & doc gen", styles['TableCell']), Paragraph("python --version", styles['TableCell'])],
        [Paragraph("pip", styles['TableCell']), Paragraph("Latest", styles['TableCell']), Paragraph("Python package manager", styles['TableCell']), Paragraph("pip --version", styles['TableCell'])],
        [Paragraph("Google Gemini API", styles['TableCell']), Paragraph("API Access Key", styles['TableCell']), Paragraph("Synthesizing research answers", styles['TableCell']), Paragraph("N/A (Web Portal)", styles['TableCell'])],
    ]
    t_pdf_req = Table(req_data, colWidths=[90, 85, 180, 149])
    t_pdf_req.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), c_primary),
        ('GRID', (0, 0), (-1, -1), 0.5, c_border),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, c_bg_light]),
    ]))
    story.append(t_pdf_req)
    story.append(Spacer(1, 10))

    # --------------------------------------------------------------------------
    # CHAPTERS 3 & 4 & 5: INSTALLATION
    # --------------------------------------------------------------------------
    story.append(Paragraph("3. Tool Installation (VS Code, Git, Node.js, Python)", styles['SectionH1']))
    story.append(Paragraph("<b>1. VS Code:</b> Download from https://code.visualstudio.com. During setup, enable 'Add to PATH'. Open terminal via Ctrl + `.", styles['BulletCustom']))
    story.append(Paragraph("<b>2. Git for Windows:</b> Download from https://git-scm.com/download/win. Verify with `git --version`.", styles['BulletCustom']))
    story.append(Paragraph("<b>3. Node.js & npm:</b> Download v20 LTS from https://nodejs.org. Verify with `node --version` and `npm --version`.", styles['BulletCustom']))
    story.append(Paragraph("<b>4. Python 3.10+:</b> Download from https://python.org. CRITICAL: Check 'Add python.exe to PATH'. Verify with `python --version`.", styles['BulletCustom']))
    story.append(Spacer(1, 8))

    # --------------------------------------------------------------------------
    # CHAPTER 6: CLONING & STRUCTURE
    # --------------------------------------------------------------------------
    story.append(Paragraph("4. Obtaining the Code & Repository Structure", styles['SectionH1']))
    story.append(make_code(
        "# Clone repository from GitHub\n"
        "git clone https://github.com/TejeshwarDivekar/NexusAI.git\n"
        "cd NexusAI\n"
        "code ."
    ))
    story.append(Spacer(1, 6))

    # --------------------------------------------------------------------------
    # CHAPTER 8 & 9: DEPENDENCIES & ENVIRONMENT VARIABLES
    # --------------------------------------------------------------------------
    story.append(Paragraph("5. Installing Dependencies & Configuring Environment", styles['SectionH1']))
    story.append(Paragraph("Install frontend packages (npm) and backend packages (pip):", styles['BodyCustom']))
    story.append(make_code(
        "# 1. Install frontend packages in root\n"
        "npm install\n\n"
        "# 2. Install backend Python packages\n"
        "cd backend && pip install -r requirements.txt && cd .."
    ))
    story.append(Spacer(1, 8))

    story.append(Paragraph("Master Environment Variables Table:", styles['SectionH2']))
    env_pdf_data = [
        [Paragraph("Variable Name", styles['TableHead']), Paragraph("Req?", styles['TableHead']), Paragraph("Target", styles['TableHead']), Paragraph("Purpose", styles['TableHead']), Paragraph("Development Value", styles['TableHead'])],
        [Paragraph("AUTH_SECRET", styles['TableCell']), Paragraph("Yes", styles['TableCell']), Paragraph(".env.local", styles['TableCell']), Paragraph("NextAuth session signing secret", styles['TableCell']), Paragraph("Run: npx auth secret", styles['TableCell'])],
        [Paragraph("AUTH_TRUST_HOST", styles['TableCell']), Paragraph("Yes", styles['TableCell']), Paragraph(".env.local", styles['TableCell']), Paragraph("Trust proxy host headers", styles['TableCell']), Paragraph("true", styles['TableCell'])],
        [Paragraph("NEXTAUTH_URL", styles['TableCell']), Paragraph("Yes", styles['TableCell']), Paragraph(".env.local", styles['TableCell']), Paragraph("Public base URL", styles['TableCell']), Paragraph("http://localhost:3000", styles['TableCell'])],
        [Paragraph("BACKEND_INTERNAL_URL", styles['TableCell']), Paragraph("Yes", styles['TableCell']), Paragraph(".env.local", styles['TableCell']), Paragraph("Proxy target for FastAPI", styles['TableCell']), Paragraph("http://127.0.0.1:8000", styles['TableCell'])],
        [Paragraph("GOOGLE_GENERATIVE_AI_API_KEY", styles['TableCell']), Paragraph("Yes", styles['TableCell']), Paragraph("Both", styles['TableCell']), Paragraph("Gemini LLM API Key", styles['TableCell']), Paragraph("From aistudio.google.com", styles['TableCell'])],
        [Paragraph("SECRET_KEY", styles['TableCell']), Paragraph("Yes", styles['TableCell']), Paragraph("backend/.env", styles['TableCell']), Paragraph("FastAPI JWT signing key", styles['TableCell']), Paragraph("32+ character string", styles['TableCell'])],
        [Paragraph("DATABASE_URL", styles['TableCell']), Paragraph("Yes", styles['TableCell']), Paragraph("backend/.env", styles['TableCell']), Paragraph("SQLite connection string", styles['TableCell']), Paragraph("sqlite:///./nexusai_research.db", styles['TableCell'])],
    ]
    t_pdf_env = Table(env_pdf_data, colWidths=[130, 35, 75, 144, 120])
    t_pdf_env.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), c_primary),
        ('GRID', (0, 0), (-1, -1), 0.5, c_border),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, c_bg_light]),
    ]))
    story.append(t_pdf_env)
    story.append(Spacer(1, 10))

    # --------------------------------------------------------------------------
    # CHAPTER 15, 16, 17: RUNNING APPLICATION
    # --------------------------------------------------------------------------
    story.append(Paragraph("6. Running Backend & Frontend (Dual-Terminal Workflow)", styles['SectionH1']))
    story.append(Paragraph("NexusResearch requires two active terminal sessions in VS Code:", styles['BodyCustom']))
    story.append(make_code(
        "# TERMINAL 1 — Start FastAPI Backend Server (Port 8000)\n"
        "python backend/run.py\n"
        "# Expected: Uvicorn running on http://0.0.0.0:8000\n\n"
        "# TERMINAL 2 — Start Next.js Frontend Server (Port 3000)\n"
        "npm run dev\n"
        "# Expected: Ready on http://localhost:3000"
    ))
    story.append(Spacer(1, 8))

    # --------------------------------------------------------------------------
    # TESTING & FEATURES
    # --------------------------------------------------------------------------
    story.append(Paragraph("7. Verifying Features & Testing the Application", styles['SectionH1']))
    story.append(Paragraph("<b>1. Authentication:</b> Open `http://localhost:3000` and log in via Google OAuth or local credentials.", styles['BulletCustom']))
    story.append(Paragraph("<b>2. Research Inquiry:</b> Enter 'What is artificial intelligence?' and observe Stage 1 to 7 progress stream.", styles['BulletCustom']))
    story.append(Paragraph("<b>3. Evidence Matrix Interaction:</b> Click an evidence item in the right panel. Inspect the center detail explanation, then click '← Back to Answer' to verify instant scroll restoration.", styles['BulletCustom']))
    story.append(Paragraph("<b>4. Document Export:</b> Click 'Export IEEE (.docx)' and 'Export PDF (.pdf)' to download publication-grade reports.", styles['BulletCustom']))
    story.append(Paragraph("<b>5. Automated Test Suite:</b> Run `python -m pytest backend/tests -v` (30/30 tests pass).", styles['BulletCustom']))
    story.append(Paragraph("<b>6. Production Build Check:</b> Run `npm run build` (Next.js 16.3.1 compiles with 0 errors).", styles['BulletCustom']))
    story.append(Spacer(1, 10))

    # --------------------------------------------------------------------------
    # FRESH LAPTOP CHECKLIST
    # --------------------------------------------------------------------------
    story.append(Paragraph("8. Printable Fresh-Laptop Setup Checklist", styles['SectionH1']))
    chks = [
        "[ ] 1. Windows OS updated and VS Code installed with PowerShell terminal",
        "[ ] 2. Git for Windows installed (`git --version` verified)",
        "[ ] 3. Node.js v20.x and npm installed (`node --version` verified)",
        "[ ] 4. Python 3.10+ installed with 'Add to PATH' (`python --version` verified)",
        "[ ] 5. Repository cloned via `git clone https://github.com/TejeshwarDivekar/NexusAI.git`",
        "[ ] 6. Frontend packages installed (`npm install`)",
        "[ ] 7. Backend packages installed (`cd backend && pip install -r requirements.txt`)",
        "[ ] 8. Root `.env.local` created with valid `AUTH_SECRET` and `GOOGLE_GENERATIVE_AI_API_KEY`",
        "[ ] 9. Backend `backend/.env` created with `SECRET_KEY` and `GOOGLE_API_KEY`",
        "[ ] 10. Backend started (`python backend/run.py`) on port 8000",
        "[ ] 11. Frontend started (`npm run dev`) on port 3000",
        "[ ] 12. Health check verified at `http://127.0.0.1:8000/api/v1/health`",
        "[ ] 13. UI verified at `http://localhost:3000` with test research query",
        "[ ] 14. Evidence matrix inspection and Back to Answer verified",
        "[ ] 15. IEEE Word (.docx) and PDF (.pdf) downloads verified",
        "[ ] 16. Backend test suite passed (`python -m pytest backend/tests -v`)",
        "[ ] 17. Frontend production build verified (`npm run build`)"
    ]
    for c in chks:
        story.append(Paragraph(c, styles['BulletCustom']))
    story.append(Spacer(1, 10))

    # --------------------------------------------------------------------------
    # REBUILD RUNBOOK & SECURITY
    # --------------------------------------------------------------------------
    story.append(Paragraph("9. Rapid Reconstruction & Security Governance", styles['SectionH1']))
    story.append(make_callout(
        "CRITICAL SECURITY GOVERNANCE MANDATE:\n"
        "Never commit `.env.local`, `backend/.env`, or real API keys to GitHub. "
        "Always verify your staged changes with `git status` before committing. "
        "Keep production credentials safely managed inside cloud platform settings (e.g. Railway Variables).",
        title="SECURITY MANDATE"
    ))

    doc.build(story, canvasmaker=NumberedCanvas)


if __name__ == "__main__":
    generate_all_guides()
